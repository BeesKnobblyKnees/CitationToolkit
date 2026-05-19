#!/usr/bin/env python3
"""
Reference List Comparator — App 3
Compare two reference lists from EndNote XML, Word docs, or plain text.
Finds missing, extra, and duplicate references. Locates each in manuscript text.

Usage:
    python ref_comparator.py --a chapter.xml --b master.xml
    python ref_comparator.py --a chapter.xml --b master.xml --doc manuscript.docx
    python ref_comparator.py --a manuscript.docx --b library.xml --doc manuscript.docx
    python ref_comparator.py --a refs_v1.txt --b refs_v2.txt
"""

import argparse, re, sys, xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

MATCH_THRESHOLD = 0.30   # TF-IDF: confident same reference
FUZZY_THRESHOLD = 0.10   # TF-IDF: possible match, needs review

# ── Parsers ─────────────────────────────────────────────────────────────────

def _xml_text(elem, path):
    node = elem.find(path)
    if node is None: return ""
    parts = [node.text or ""] + [c.text or "" for c in node]
    return " ".join(p.strip() for p in parts if p.strip())

def load_endnote_xml(path):
    tree = ET.parse(path); root = tree.getroot(); refs = []
    for rec in root.iter("record"):
        authors = []
        for a in rec.findall(".//contributors/authors/author"):
            name = " ".join(p.strip() for p in ([a.text or ""]+[c.text or "" for c in a]) if p.strip())
            if name: authors.append(name)
        title   = _xml_text(rec, ".//titles/title")
        journal = _xml_text(rec, ".//periodical/full-title") or _xml_text(rec, ".//periodical/abbr-1")
        year    = _xml_text(rec, ".//dates/year")
        abstract= _xml_text(rec, ".//abstract")
        rec_num = _xml_text(rec, "rec-number")
        if not title: continue
        corpus = " ".join(filter(None,[title,abstract,journal,year," ".join(a.split(",")[0] for a in authors)]))
        refs.append(dict(id=rec_num, authors=authors, title=title,
                         journal=journal, year=year, corpus=corpus, source=path))
    return refs

def load_word_reflist(path):
    doc = Document(path); refs = []; in_refs = False
    ref_pattern = re.compile(r'^\s*\d+[\.\)]\s+(.+)')
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        if re.match(r'^(references?|bibliography|works cited)$', text, re.IGNORECASE):
            in_refs = True; continue
        if in_refs or ref_pattern.match(text):
            in_refs = True
            m = ref_pattern.match(text); ref_text = m.group(1) if m else text
            year_m = re.search(r'\b(19|20)\d{2}\b', ref_text)
            refs.append(dict(id=str(len(refs)+1), authors=[], title=ref_text[:200],
                             journal="", year=year_m.group(0) if year_m else "",
                             corpus=ref_text, source=path))
    if not refs:  # fallback: all substantive paragraphs
        for i,para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if len(text) > 40:
                year_m = re.search(r'\b(19|20)\d{2}\b', text)
                refs.append(dict(id=str(i), authors=[], title=text[:200], journal="",
                                 year=year_m.group(0) if year_m else "", corpus=text, source=path))
    return refs

def load_plain_text(path):
    refs = []
    with open(path, encoding="utf-8", errors="replace") as f: content = f.read()
    blocks = re.split(r'\n{2,}', content)
    for i,block in enumerate(blocks):
        block = block.strip()
        if not block or len(block) < 20: continue
        m = re.match(r'^\d+[\.\)]\s+(.*)', block, re.DOTALL)
        text = m.group(1) if m else block
        year_m = re.search(r'\b(19|20)\d{2}\b', text)
        refs.append(dict(id=str(i), authors=[], title=text[:200], journal="",
                         year=year_m.group(0) if year_m else "", corpus=text, source=path))
    return refs

def load_refs(path):
    p = Path(path); ext = p.suffix.lower()
    print(f"  Loading {p.name}…", end=" ", flush=True)
    if ext == ".xml":   refs = load_endnote_xml(path)
    elif ext == ".docx": refs = load_word_reflist(path)
    elif ext == ".txt":  refs = load_plain_text(path)
    else: print(f"Unsupported: {ext}"); sys.exit(1)
    print(f"{len(refs)} refs.")
    return refs

def fmt(ref, short=False):
    aa = ref.get("authors",[])
    a  = aa[0].split(",")[0] if len(aa)==1 else \
         f"{aa[0].split(',')[0]} & {aa[1].split(',')[0]}" if len(aa)==2 else \
         f"{aa[0].split(',')[0]} et al." if aa else "Unknown"
    y  = ref.get("year","n.d.")
    t  = ref["title"][:100] + ("…" if len(ref["title"])>100 else "")
    return f"{a} ({y}) — {t}" if short else f"{a} ({y}). {t}. {ref.get('journal','')}"

# ── Comparator Engine ────────────────────────────────────────────────────────

def compare(refs_a, refs_b):
    print("  Building TF-IDF comparison matrix…")
    all_corpora = [r["corpus"] for r in refs_a] + [r["corpus"] for r in refs_b]
    vec = TfidfVectorizer(ngram_range=(1,2), sublinear_tf=True, max_features=50000)
    vec.fit(all_corpora)
    emb_a = vec.transform([r["corpus"] for r in refs_a])
    emb_b = vec.transform([r["corpus"] for r in refs_b])
    matrix = cosine_similarity(emb_a, emb_b)

    matched, only_a, fuzzy, matched_b = [], [], [], set()
    for i, ref_a in enumerate(refs_a):
        best_j = int(matrix[i].argmax()); best_s = float(matrix[i][best_j])
        if best_s >= MATCH_THRESHOLD:
            matched.append((ref_a, refs_b[best_j], best_s)); matched_b.add(best_j)
        elif best_s >= FUZZY_THRESHOLD:
            fuzzy.append((ref_a, refs_b[best_j], best_s))
        else:
            only_a.append(ref_a)
    only_b = [refs_b[j] for j in range(len(refs_b)) if j not in matched_b]
    return dict(matched=matched, only_in_a=only_a, only_in_b=only_b, fuzzy=fuzzy)

# ── Usage Locator ────────────────────────────────────────────────────────────

def find_usage(refs, doc_path):
    doc   = Document(doc_path)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    usage = {}
    for ref in refs:
        words  = [w for w in ref["title"].split() if len(w) > 5][:5]
        aa     = ref.get("authors",[]); year = ref.get("year","")
        found  = []
        for para in paras:
            if aa and year:
                last = aa[0].split(",")[0].strip().split()[-1]
                if last.lower() in para.lower() and year in para:
                    found.append(para[:200]); continue
            hits = sum(1 for w in words if w.lower() in para.lower())
            if hits >= min(3, len(words)):
                found.append(para[:200])
        usage[ref["id"]] = found
    return usage

# ── Report Writer ────────────────────────────────────────────────────────────

def write_report(result, refs_a, refs_b, label_a, label_b, doc_path, out_path):
    doc = Document()
    doc.add_heading("Reference List Comparison Report", 0)
    doc.add_paragraph(f"Generated: {datetime.now():%Y-%m-%d %H:%M}")
    doc.add_paragraph(f"List A: {label_a}  ({len(refs_a)} references)")
    doc.add_paragraph(f"List B: {label_b}  ({len(refs_b)} references)")

    m = len(result["matched"]); a = len(result["only_in_a"])
    b = len(result["only_in_b"]); f = len(result["fuzzy"])
    doc.add_paragraph(f"Matched: {m} | Only in A: {a} | Only in B: {b} | Needs review: {f}")

    usage = {}
    if doc_path:
        print("  Locating references in manuscript…")
        all_miss = result["only_in_a"] + result["only_in_b"] + [x[0] for x in result["fuzzy"]]
        usage = find_usage(all_miss, doc_path)

    def add_usage(p, ref):
        locs = usage.get(ref["id"],[])
        if locs:
            p.add_run(f"\n  📍 Found in text ({len(locs)} location(s)):")
            for loc in locs[:2]: p.add_run(f"\n     [{loc[:150]}]")

    if result["only_in_a"]:
        doc.add_heading(f"Missing from B — Only in {label_a} ({len(result['only_in_a'])})", 1)
        doc.add_paragraph("Present in A, not found in B. May need to be added to master library.")
        for ref in result["only_in_a"]:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(fmt(ref)); r.font.color.rgb = RGBColor(0xC0,0,0)
            add_usage(p, ref)

    if result["only_in_b"]:
        doc.add_heading(f"Missing from A — Only in {label_b} ({len(result['only_in_b'])})", 1)
        doc.add_paragraph("Present in B, not found in A. May be unused or absent references.")
        for ref in result["only_in_b"]:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(fmt(ref)); r.font.color.rgb = RGBColor(0xC0,0,0)
            add_usage(p, ref)

    if result["fuzzy"]:
        doc.add_heading(f"Possible Matches — Review Required ({len(result['fuzzy'])})", 1)
        doc.add_paragraph("Similarity too low to confirm. May be same ref with different edition, journal abbrev, or formatting.")
        for ra,rb,score in result["fuzzy"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{score:.3f}] REVIEW\n").bold = True
            p.add_run(f"  A: {fmt(ra)}\n")
            p.add_run(f"  B: {fmt(rb)}")
            add_usage(p, ra)

    if result["matched"]:
        doc.add_heading(f"Confirmed Matches ({len(result['matched'])})", 1)
        for ra,rb,score in result["matched"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{score:.3f}] {fmt(ra,short=True)}")

    doc.save(out_path)
    print(f"  Report saved: {out_path}")

# ── Main ────────────────────────────────────────────────────────────────────

def main():
    ap = argparse.ArgumentParser(description="Reference List Comparator")
    ap.add_argument("--a",   required=True, help="List A (.xml, .docx, .txt)")
    ap.add_argument("--b",   required=True, help="List B (.xml, .docx, .txt)")
    ap.add_argument("--doc", default=None,  help="Manuscript .docx — locate refs in text")
    ap.add_argument("--out", default="comparison_report.docx")
    args = ap.parse_args()

    print(f"\n{'═'*65}\n  REFERENCE LIST COMPARATOR\n{'═'*65}\n")
    refs_a = load_refs(args.a); refs_b = load_refs(args.b)
    if not refs_a or not refs_b: print("ERROR: Could not parse one or both lists."); sys.exit(1)

    result = compare(refs_a, refs_b)
    print(f"\n  ✓ Matched:          {len(result['matched'])}")
    print(f"  ✗ Only in A:        {len(result['only_in_a'])}")
    print(f"  ✗ Only in B:        {len(result['only_in_b'])}")
    print(f"  ? Needs review:     {len(result['fuzzy'])}")

    write_report(result, refs_a, refs_b, Path(args.a).name, Path(args.b).name,
                 args.doc, args.out)
    print(f"\n{'═'*65}\n  DONE — {args.out}\n{'═'*65}\n")

if __name__ == "__main__":
    main()
