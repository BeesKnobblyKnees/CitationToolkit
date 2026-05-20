"""
Figure Extractor
Extracts figures, tables, boxes, and plates from PDF or Word source files.
Outputs: PDF (one figure per page) AND Word document (one figure per page).
Clean crop, caption below, no extra text, no paraphrasing.

Run: streamlit run figure_extractor.py
"""

import io
import re
import os
import tempfile
from pathlib import Path
from copy import deepcopy

import fitz          # PyMuPDF
import pdfplumber
import streamlit as st
from PIL import Image, ImageOps
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.platypus import (
    SimpleDocTemplate, Image as RLImage,
    Paragraph, Spacer, PageBreak, KeepTogether
)

# ─────────────────────────────────────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Figure Extractor",
    page_icon="🖼",
    layout="wide",
    initial_sidebar_state="collapsed"
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;600&family=IBM+Plex+Sans:wght@300;400;600&display=swap');
html,body,[class*="css"]{font-family:'IBM Plex Sans',sans-serif;}
.main .block-container{padding-top:2rem;max-width:1000px;}
.step-hd{font-size:0.75rem;font-weight:500;letter-spacing:0.12em;text-transform:uppercase;
  color:#607080;margin-bottom:0.6rem;border-bottom:1px solid #1e2530;padding-bottom:0.3rem;}
.fig-label{font-family:'IBM Plex Mono',monospace;font-size:0.75rem;
  background:#1e2a38;color:#4fc3f7;padding:0.15rem 0.6rem;border-radius:3px;
  display:inline-block;margin-bottom:0.5rem;}
.info-box{background:#0d1525;border:1px solid #1e3050;border-radius:6px;
  padding:1rem 1.2rem;margin:0.8rem 0;font-size:0.85rem;color:#90a8c8;line-height:1.8;}
.info-box b{color:#c8d0db;}
.info-box code{background:#1a2535;padding:0.1rem 0.4rem;border-radius:3px;
  font-family:'IBM Plex Mono',monospace;font-size:0.8rem;color:#4fc3f7;}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────

SCAN_DPI  = 200
OUT_DPI   = 300
PAGE_W    = 8.5 * inch
PAGE_H    = 11.0 * inch
MARGIN    = 0.75 * inch
CONTENT_W = PAGE_W - 2 * MARGIN
CONTENT_H = PAGE_H - 2 * MARGIN

CAPTION_LABEL_RE = re.compile(
    r'(?:^|\n)\s*((?:Fig(?:ure)?\.?|Table|Box|Plate|Video)\s*\d+[\w\-\.]*)',
    re.IGNORECASE
)

# ─────────────────────────────────────────────────────────────────────────────
# PDF HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def pdf_render(pdf_bytes, page_num, dpi=200):
    doc  = fitz.open(stream=pdf_bytes, filetype="pdf")
    page = doc[page_num]
    mat  = fitz.Matrix(dpi / 72, dpi / 72)
    pix  = page.get_pixmap(matrix=mat, alpha=False)
    img  = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    doc.close()
    return img


def pdf_image_boxes(pdf_bytes, page_num, dpi=200):
    doc   = fitz.open(stream=pdf_bytes, filetype="pdf")
    page  = doc[page_num]
    scale = dpi / 72
    boxes = []
    for img_info in page.get_images(full=True):
        for rect in page.get_image_rects(img_info[0]):
            boxes.append((int(rect.x0*scale), int(rect.y0*scale),
                          int(rect.x1*scale), int(rect.y1*scale)))
    doc.close()
    return boxes


def pdf_find_caption_y(pdf_bytes, page_num, search_text, dpi=200):
    doc   = fitz.open(stream=pdf_bytes, filetype="pdf")
    page  = doc[page_num]
    rects = page.search_for(search_text[:50].strip())
    doc.close()
    if rects:
        return int(rects[0].y0 * dpi / 72)
    return None


def pdf_extract_caption(pdf_bytes, page_num, label_text):
    """
    Extract full caption starting from label_text on page_num.
    Stops at a blank line or the next figure label.
    """
    doc  = fitz.open(stream=pdf_bytes, filetype="pdf")
    page = doc[page_num]
    text = page.get_text("text")
    doc.close()

    # Find where this label appears
    idx = text.lower().find(label_text.lower()[:30])
    if idx == -1:
        return label_text

    chunk = text[idx:]
    # Cut at double newline or next figure label (whichever comes first)
    end = re.search(
        r'\n\n|\n(?=(?:Fig(?:ure)?\.?\s*\d|Table\s*\d|Box\s*\d|Plate\s*\d|Video\s*\d))',
        chunk, re.IGNORECASE
    )
    if end:
        chunk = chunk[:end.start()]

    # Clean internal line breaks
    caption = re.sub(r'\n+', ' ', chunk).strip()
    caption = re.sub(r'\s{2,}', ' ', caption)
    return caption


def pdf_scan_figures(pdf_bytes, dpi=200):
    """Scan all pages of a PDF for figure/table/box/plate captions."""
    doc         = fitz.open(stream=pdf_bytes, filetype="pdf")
    total_pages = len(doc)
    doc.close()

    figures    = []
    seen_labels = set()
    bar         = st.progress(0, text="Scanning PDF pages...")

    for page_num in range(total_pages):
        bar.progress((page_num + 1) / total_pages,
                     text=f"Scanning page {page_num+1} of {total_pages}...")

        doc  = fitz.open(stream=pdf_bytes, filetype="pdf")
        page = doc[page_num]
        text = page.get_text("text")
        doc.close()

        for m in CAPTION_LABEL_RE.finditer(text):
            label = m.group(1).strip()
            # Normalise label for dedup
            norm  = re.sub(r'\s+', ' ', label.upper())
            if norm in seen_labels:
                continue
            seen_labels.add(norm)

            # Determine type / number
            lm     = re.match(r'(Fig(?:ure)?\.?|Table|Box|Plate|Video)\s*(\d+[\w\-\.]*)',
                              label, re.IGNORECASE)
            f_type = lm.group(1).rstrip('.').title() if lm else 'Figure'
            f_num  = lm.group(2) if lm else '?'

            # Render page + locate caption
            page_img  = pdf_render(pdf_bytes, page_num, dpi=dpi)
            img_boxes = pdf_image_boxes(pdf_bytes, page_num, dpi=dpi)
            caption_y = pdf_find_caption_y(pdf_bytes, page_num, label, dpi=dpi)
            full_cap  = pdf_extract_caption(pdf_bytes, page_num, label)

            figures.append(dict(
                source='pdf',
                page_num=page_num,
                fig_type=f_type,
                fig_num=f_num,
                label=f"{f_type} {f_num}",
                caption=full_cap,
                page_img=page_img,
                img_boxes=img_boxes,
                caption_y=caption_y,
                include=True,
                crop_top=0,
                crop_bot=0,
            ))

    bar.empty()
    return figures


# ─────────────────────────────────────────────────────────────────────────────
# WORD HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def docx_extract_images(docx_bytes):
    """Extract all embedded images from a Word document with their bytes."""
    images = {}
    with __import__('zipfile').ZipFile(io.BytesIO(docx_bytes)) as z:
        for name in z.namelist():
            if name.startswith('word/media/'):
                images[name] = z.read(name)
    return images


def docx_scan_figures(docx_bytes):
    """
    Scan a Word document for figures.
    Finds paragraphs containing embedded images, then looks for the
    caption paragraph(s) immediately below.
    """
    doc     = Document(io.BytesIO(docx_bytes))
    paras   = doc.paragraphs
    figures = []
    seen    = set()

    # Namespace for drawing elements
    DRAW_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    BLIP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
    REL_NS  = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

    # Extract relationship map: rId -> image path in zip
    import zipfile
    from lxml import etree
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        rels_xml = z.read('word/_rels/document.xml.rels')
        media    = {name for name in z.namelist() if name.startswith('word/media/')}

    rels_root = etree.fromstring(rels_xml)
    rel_map   = {}
    for rel in rels_root:
        rid    = rel.get('Id','')
        target = rel.get('Target','')
        if target.startswith('media/'):
            rel_map[rid] = 'word/' + target

    # Walk paragraphs
    for pi, para in enumerate(paras):
        # Check if this paragraph contains an image
        blips = para._p.findall(
            f'.//{{{DRAW_NS}}}blip',
            namespaces={}
        )
        # Also check with the actual blipFill namespace
        all_blips = para._p.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
        if not all_blips:
            # try alternate namespace
            all_blips = para._p.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing}blip')

        # Look for any r:embed attribute pointing to media
        has_image = False
        img_rids  = []
        xml_str   = para._p.xml if hasattr(para._p, 'xml') else str(para._p.tag)
        # Use string search on XML
        p_xml     = para._p.xml
        for rid, path in rel_map.items():
            if rid in p_xml and path in media:
                has_image = True
                img_rids.append(rid)

        if not has_image:
            continue

        # Found an image paragraph — look for caption below
        caption = ""
        cap_paras = []
        for j in range(pi + 1, min(pi + 8, len(paras))):
            next_text = paras[j].text.strip()
            if not next_text:
                continue
            # Check if this looks like a caption
            if CAPTION_LABEL_RE.match(next_text):
                cap_paras.append(next_text)
                # Keep reading until blank or next image
                for k in range(j + 1, min(j + 6, len(paras))):
                    kt = paras[k].text.strip()
                    if not kt or CAPTION_LABEL_RE.match(kt):
                        break
                    cap_paras.append(kt)
                break
            # If it looks like body text (long), stop
            if len(next_text) > 20 and not CAPTION_LABEL_RE.match(next_text):
                # Could still be a caption if it follows immediately
                if not cap_paras:
                    cap_paras.append(next_text)
                break

        if cap_paras:
            caption = ' '.join(cap_paras)
            caption = re.sub(r'\s{2,}', ' ', caption).strip()
        else:
            caption = f"Figure {len(figures)+1}"

        # Parse label from caption
        lm     = CAPTION_LABEL_RE.match(caption)
        f_type = 'Figure'
        f_num  = str(len(figures) + 1)
        if lm:
            inner = lm.group(1).strip()
            lm2   = re.match(r'(Fig(?:ure)?\.?|Table|Box|Plate|Video)\s*(\d+[\w\-\.]*)',
                             inner, re.IGNORECASE)
            if lm2:
                f_type = lm2.group(1).rstrip('.').title()
                f_num  = lm2.group(2)

        label = f"{f_type} {f_num}"
        norm  = re.sub(r'\s+', ' ', label.upper())
        if norm in seen:
            continue
        seen.add(norm)

        # Get image bytes
        img_bytes_list = []
        for rid in img_rids:
            path = rel_map.get(rid, '')
            if path in media:
                with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
                    raw = z.read(path)
                img_bytes_list.append(raw)

        if not img_bytes_list:
            continue

        # Convert first image to PIL
        try:
            pil_img = Image.open(io.BytesIO(img_bytes_list[0])).convert('RGB')
        except Exception:
            continue

        figures.append(dict(
            source='docx',
            page_num=pi,
            fig_type=f_type,
            fig_num=f_num,
            label=label,
            caption=caption,
            page_img=pil_img,
            img_bytes=img_bytes_list[0],
            img_boxes=[],
            caption_y=None,
            include=True,
            crop_top=0,
            crop_bot=0,
        ))

    return figures


# ─────────────────────────────────────────────────────────────────────────────
# IMAGE CROPPING
# ─────────────────────────────────────────────────────────────────────────────

def smart_crop(page_img, img_boxes, caption_y_px, crop_top_pct=0,
               crop_bot_pct=0, padding=20):
    """
    Crop the page image to just the figure content, excluding caption area.
    Applies manual top/bottom crop adjustments as percentage of page height.
    """
    w, h = page_img.size

    # Apply manual adjustments first
    manual_top = int(h * crop_top_pct / 100)
    manual_bot = h - int(h * crop_bot_pct / 100)

    if img_boxes:
        # Filter to boxes above the caption
        if caption_y_px:
            above = [b for b in img_boxes if b[3] <= caption_y_px + 30]
        else:
            above = img_boxes
        if not above:
            above = img_boxes

        x0 = max(0, min(b[0] for b in above) - padding)
        y0 = max(manual_top, min(b[1] for b in above) - padding)
        x1 = min(w, max(b[2] for b in above) + padding)
        y1 = min(manual_bot,
                 (caption_y_px - padding) if caption_y_px else max(b[3] for b in above) + padding)
        y1 = max(y0 + 10, y1)
        return page_img.crop((x0, y0, x1, y1))

    elif caption_y_px:
        # No image blocks detected — crop from manual_top to caption
        return page_img.crop((0, manual_top, w,
                              max(manual_top + 10, caption_y_px - padding)))
    else:
        return page_img.crop((0, manual_top, w, manual_bot))


def get_cropped_image(fig, pdf_bytes=None, high_dpi=300):
    """
    Produce the final high-quality cropped image for a figure.
    For PDF source: re-renders at high DPI then crops.
    For Word source: uses extracted image bytes directly.
    """
    if fig['source'] == 'docx':
        # Word: use extracted image, apply manual crop only
        img = Image.open(io.BytesIO(fig['img_bytes'])).convert('RGB')
        w, h = img.size
        top = int(h * fig['crop_top'] / 100)
        bot = h - int(h * fig['crop_bot'] / 100)
        if top > 0 or bot < h:
            img = img.crop((0, top, w, bot))
        img = ImageOps.autocontrast(img, cutoff=0.5)
        return img

    else:
        # PDF: render at high DPI and smart-crop
        page_img  = pdf_render(pdf_bytes, fig['page_num'], dpi=high_dpi)
        img_boxes = pdf_image_boxes(pdf_bytes, fig['page_num'], dpi=high_dpi)

        # Scale caption_y from scan DPI to output DPI
        cap_y = None
        if fig['caption_y'] is not None:
            cap_y = int(fig['caption_y'] * high_dpi / SCAN_DPI)

        cropped = smart_crop(
            page_img, img_boxes, cap_y,
            crop_top_pct=fig['crop_top'],
            crop_bot_pct=fig['crop_bot'],
        )
        cropped = ImageOps.autocontrast(cropped, cutoff=0.5)

        # Trim white border
        try:
            bbox = ImageOps.invert(cropped.convert('L')).getbbox()
            if bbox:
                pad = 15
                bw, bh = cropped.size
                cropped = cropped.crop((
                    max(0, bbox[0]-pad), max(0, bbox[1]-pad),
                    min(bw, bbox[2]+pad), min(bh, bbox[3]+pad)
                ))
        except Exception:
            pass

        return cropped


# ─────────────────────────────────────────────────────────────────────────────
# OUTPUT BUILDERS
# ─────────────────────────────────────────────────────────────────────────────

def img_to_bytes(img, fmt='PNG'):
    buf = io.BytesIO()
    img.save(buf, format=fmt, dpi=(OUT_DPI, OUT_DPI))
    buf.seek(0)
    return buf.read()


def build_pdf(figures, captions, pdf_bytes=None):
    """
    Build output PDF — one figure per page.
    Image centered, caption below in Times New Roman 10pt.
    """
    buf = io.BytesIO()

    # Caption paragraph style
    cap_style = ParagraphStyle(
        'Caption',
        fontName='Times-Roman',
        fontSize=10,
        leading=14,
        alignment=TA_LEFT,
        spaceAfter=6,
    )
    cap_bold_style = ParagraphStyle(
        'CaptionBold',
        fontName='Times-Bold',
        fontSize=10,
        leading=14,
        alignment=TA_LEFT,
    )

    story = []

    included = [f for f in figures if f.get('include', True)]

    for idx, fig in enumerate(included):
        if idx > 0:
            story.append(PageBreak())

        # Get cropped image
        try:
            cropped = get_cropped_image(fig, pdf_bytes=pdf_bytes, high_dpi=OUT_DPI)
        except Exception as e:
            continue

        img_bytes = img_to_bytes(cropped)
        img_w_px, img_h_px = cropped.size
        aspect = img_w_px / img_h_px

        # Estimate caption height
        caption_text = captions.get(fig['label'], fig['caption'])
        cap_lines    = max(1, len(caption_text) // 90 + 1)
        cap_h        = cap_lines * 14 / 72 * inch + 0.2 * inch

        # Scale image to fit
        max_img_h = CONTENT_H - cap_h - 0.3 * inch
        max_img_w = CONTENT_W

        display_w = min(max_img_w, img_w_px / OUT_DPI * inch)
        display_h = display_w / aspect
        if display_h > max_img_h:
            display_h = max_img_h
            display_w = display_h * aspect

        # Center image with spacer trick
        left_pad = (CONTENT_W - display_w) / 2
        img_obj  = RLImage(io.BytesIO(img_bytes),
                           width=display_w, height=display_h)

        # Vertical centering: push image down so image+caption fills page
        total_content = display_h + cap_h
        top_spacer    = (CONTENT_H - total_content) / 2
        if top_spacer > 0:
            story.append(Spacer(1, top_spacer))

        story.append(img_obj)
        story.append(Spacer(1, 0.15 * inch))

        # Caption: bold label + normal rest
        label_m = re.match(
            r'^((?:Fig(?:ure)?\.?\s*\d+[\w\-\.]*|Table\s*\d+[\w\-\.]*|'
            r'Box\s*\d+[\w\-\.]*|Plate\s*\d+[\w\-\.]*|Video\s*\d+[\w\-\.]*)'
            r'\.?\s*)',
            caption_text, re.IGNORECASE
        )
        if label_m:
            bold_part = label_m.group(1)
            rest_part = caption_text[len(bold_part):]
            cap_html  = f'<b>{bold_part}</b>{rest_part}'
        else:
            cap_html = caption_text

        story.append(Paragraph(cap_html, cap_style))

    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=MARGIN, rightMargin=MARGIN,
        topMargin=MARGIN, bottomMargin=MARGIN,
        title="Figures"
    )
    doc.build(story)
    buf.seek(0)
    return buf.read()


def build_word(figures, captions, pdf_bytes=None):
    """
    Build output Word document — one figure per page.
    Image centered, caption below in Times New Roman 10pt.
    """
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page setup: US Letter, 0.75" margins
    section = doc.sections[0]
    section.page_width   = int(8.5  * 914400)
    section.page_height  = int(11.0 * 914400)
    m = int(0.75 * 914400)
    section.top_margin = section.bottom_margin = m
    section.left_margin = section.right_margin = m

    content_w_in = 8.5  - 2 * 0.75
    content_h_in = 11.0 - 2 * 0.75

    included = [f for f in figures if f.get('include', True)]

    for idx, fig in enumerate(included):
        if idx > 0:
            doc.add_page_break()

        try:
            cropped = get_cropped_image(fig, pdf_bytes=pdf_bytes, high_dpi=OUT_DPI)
        except Exception:
            continue

        img_w_px, img_h_px = cropped.size
        aspect = img_w_px / img_h_px

        # Estimate caption height (rough: 1 line per 90 chars)
        caption_text = captions.get(fig['label'], fig['caption'])
        cap_lines    = max(1, len(caption_text) // 90 + 1)
        cap_h_in     = cap_lines * 14 / 72 + 0.2

        max_img_h = content_h_in - cap_h_in - 0.3
        max_img_w = content_w_in

        display_w = min(max_img_w, img_w_px / OUT_DPI)
        display_h = display_w / aspect
        if display_h > max_img_h:
            display_h = max_img_h
            display_w = display_h * aspect

        # Image paragraph — centered
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.paragraph_format.space_before = Pt(6)
        img_para.paragraph_format.space_after  = Pt(6)
        run = img_para.add_run()
        img_buf = io.BytesIO(img_to_bytes(cropped))
        run.add_picture(img_buf, width=Inches(display_w))

        # Caption paragraph
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap_para.paragraph_format.space_before = Pt(4)
        cap_para.paragraph_format.space_after  = Pt(4)

        label_m = re.match(
            r'^((?:Fig(?:ure)?\.?\s*\d+[\w\-\.]*|Table\s*\d+[\w\-\.]*|'
            r'Box\s*\d+[\w\-\.]*|Plate\s*\d+[\w\-\.]*|Video\s*\d+[\w\-\.]*)'
            r'\.?\s*)',
            caption_text, re.IGNORECASE
        )
        if label_m:
            bold_part = label_m.group(1)
            rest_part = caption_text[len(bold_part):]
            br = cap_para.add_run(bold_part)
            br.bold = True
            br.font.name = 'Times New Roman'
            br.font.size = Pt(10)
            if rest_part:
                rr = cap_para.add_run(rest_part)
                rr.font.name = 'Times New Roman'
                rr.font.size = Pt(10)
        else:
            r = cap_para.add_run(caption_text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────────────────────────────────────

for k, v in [
    ('figures', []),
    ('pdf_bytes', None),
    ('src_name', ""),
    ('src_type', ""),
    ('captions', {}),
    ('scanned', False),
]:
    if k not in st.session_state:
        st.session_state[k] = v

# ─────────────────────────────────────────────────────────────────────────────
# UI
# ─────────────────────────────────────────────────────────────────────────────

st.markdown("# 🖼 Figure Extractor")
st.markdown(
    "Extracts figures, tables, boxes, and plates from published PDF or Word files. "
    "Outputs a **PDF** and a **Word document** — one figure per page, "
    "clean crop, exact caption below."
)

st.markdown("""
<div class="info-box">
<b>How it works:</b><br>
1. Upload your PDF or Word file<br>
2. The app scans every page for captions starting with Figure / Table / Box / Plate<br>
3. Each figure is cropped cleanly — no surrounding text, no full pages<br>
4. Review crops and captions, adjust if needed<br>
5. Download both PDF and Word outputs — one figure per page
</div>
""", unsafe_allow_html=True)

# ── STEP 1: Upload ───────────────────────────────────────────────────────────
st.markdown('<div class="step-hd">Step 1 — Upload source file</div>',
            unsafe_allow_html=True)

uploaded = st.file_uploader(
    "PDF or Word document",
    type=["pdf", "docx"],
    key="src_upload",
    help="Published chapter or textbook PDF, or a Word document with embedded figures"
)

col1, col2 = st.columns(2)
with col1:
    quality = st.selectbox(
        "Scan quality",
        ["Fast (150 DPI)", "Standard (200 DPI)", "High (300 DPI)"],
        index=1
    )
with col2:
    margin_in = st.slider("Page margins (inches)", 0.5, 1.5, 0.75, 0.25)

dpi_map  = {"Fast (150 DPI)": 150, "Standard (200 DPI)": 200, "High (300 DPI)": 300}
scan_dpi = dpi_map[quality]

if uploaded:
    src_bytes = uploaded.read()
    src_ext   = Path(uploaded.name).suffix.lower()

    if src_bytes != st.session_state.pdf_bytes:
        st.session_state.pdf_bytes = src_bytes
        st.session_state.src_name  = uploaded.name
        st.session_state.src_type  = src_ext
        st.session_state.figures   = []
        st.session_state.captions  = {}
        st.session_state.scanned   = False

    if src_ext == '.pdf':
        doc = fitz.open(stream=src_bytes, filetype="pdf")
        st.caption(f"Loaded: {uploaded.name} — {len(doc.pages)} pages")
        doc.close()
    else:
        st.caption(f"Loaded: {uploaded.name} — Word document")

    if st.button("Scan for figures", type="primary"):
        with st.spinner("Scanning..."):
            if src_ext == '.pdf':
                figs = pdf_scan_figures(src_bytes, dpi=scan_dpi)
            else:
                figs = docx_scan_figures(src_bytes)

        st.session_state.figures = figs
        st.session_state.scanned = True
        st.session_state.captions = {}

        if not figs:
            st.warning(
                "No figures found. Make sure captions start with "
                "Figure / Fig / Table / Box / Plate followed by a number."
            )
        else:
            st.success(f"Found {len(figs)} figure(s). Review below.")

# ── STEP 2: Review ───────────────────────────────────────────────────────────
if st.session_state.scanned and st.session_state.figures:
    figs = st.session_state.figures
    st.divider()
    st.markdown(
        f'<div class="step-hd">Step 2 — Review {len(figs)} figure(s)</div>',
        unsafe_allow_html=True
    )

    # Type summary
    type_counts = {}
    for f in figs:
        type_counts[f['fig_type']] = type_counts.get(f['fig_type'], 0) + 1
    cols = st.columns(min(len(type_counts), 6))
    for i, (t, c) in enumerate(sorted(type_counts.items())):
        cols[i].metric(t + "s", c)

    gcol1, gcol2 = st.columns(2)
    with gcol1:
        if st.button("Select all"):
            for f in figs: f['include'] = True
            st.rerun()
    with gcol2:
        if st.button("Deselect all"):
            for f in figs: f['include'] = False
            st.rerun()

    st.markdown("")

    for i, fig in enumerate(figs):
        included = fig.get('include', True)
        border   = "2px solid #4fc3f7" if included else "1px solid #2a3a4a"
        opacity  = "1" if included else "0.45"

        with st.container():
            chk_col, content_col = st.columns([1, 11])

            with chk_col:
                inc = st.checkbox("", value=included, key=f"inc_{i}")
                if inc != included:
                    figs[i]['include'] = inc
                    st.rerun()

            with content_col:
                st.markdown(
                    f'<span class="fig-label">{fig["label"]}</span>'
                    f' &nbsp; <span style="font-size:0.8rem;color:#607080">'
                    f'{"Page " + str(fig["page_num"]+1) if fig["source"]=="pdf" else "Word doc"}'
                    f'</span>',
                    unsafe_allow_html=True
                )

                with st.expander("Preview & adjust crop"):
                    p1, p2 = st.columns(2)
                    with p1:
                        st.caption("Source")
                        thumb = fig['page_img'].copy()
                        thumb.thumbnail((280, 360))
                        st.image(thumb)
                    with p2:
                        st.caption("Auto-crop preview")
                        try:
                            if fig['source'] == 'pdf':
                                preview_crop = smart_crop(
                                    fig['page_img'],
                                    fig['img_boxes'],
                                    fig['caption_y'],
                                    crop_top_pct=fig['crop_top'],
                                    crop_bot_pct=fig['crop_bot'],
                                )
                            else:
                                preview_crop = fig['page_img'].copy()
                            preview_crop.thumbnail((280, 360))
                            st.image(preview_crop)
                        except Exception as e:
                            st.caption(f"Preview error: {e}")

                    if fig['source'] == 'pdf':
                        st.markdown("**Manual crop adjustment** (% of page)")
                        ac1, ac2 = st.columns(2)
                        with ac1:
                            top = st.slider("Remove from top %", 0, 60, fig['crop_top'],
                                            1, key=f"top_{i}")
                        with ac2:
                            bot = st.slider("Remove from bottom %", 0, 60, fig['crop_bot'],
                                            1, key=f"bot_{i}")
                        figs[i]['crop_top'] = top
                        figs[i]['crop_bot'] = bot
                    else:
                        st.markdown("**Manual crop** (Word source — crop top/bottom of image)")
                        ac1, ac2 = st.columns(2)
                        with ac1:
                            top = st.slider("Remove from top %", 0, 60, fig['crop_top'],
                                            1, key=f"top_{i}")
                        with ac2:
                            bot = st.slider("Remove from bottom %", 0, 60, fig['crop_bot'],
                                            1, key=f"bot_{i}")
                        figs[i]['crop_top'] = top
                        figs[i]['crop_bot'] = bot

                # Caption editor
                current = st.session_state.captions.get(fig['label'], fig['caption'])
                st.markdown(
                    "**Caption** *(exact text from source — only fix formatting, "
                    "do not paraphrase)*"
                )
                edited = st.text_area(
                    "Caption text",
                    value=current,
                    height=90,
                    key=f"cap_{i}",
                    label_visibility="collapsed",
                    help="Edit only if the extracted text has OCR errors or line break issues. Do not paraphrase."
                )
                if edited != current:
                    st.session_state.captions[fig['label']] = edited

                st.markdown("---")

    # ── STEP 3: Download ─────────────────────────────────────────────────────
    included_count = sum(1 for f in figs if f.get('include', True))
    st.divider()
    st.markdown(
        f'<div class="step-hd">Step 3 — Generate outputs ({included_count} figures)</div>',
        unsafe_allow_html=True
    )

    if included_count == 0:
        st.warning("No figures selected. Check at least one figure above.")
    else:
        st.markdown(f"""
        <div class="info-box">
        Generates <b>{included_count} page(s)</b> — one per figure.<br>
        Layout: figure image (cropped, centered) → caption below (bold label + normal text).<br>
        Font: Times New Roman 10pt &nbsp;|&nbsp; Images: {OUT_DPI} DPI &nbsp;|&nbsp;
        Margins: {margin_in}"
        </div>
        """, unsafe_allow_html=True)

        if st.button(f"Generate PDF + Word ({included_count} figures)", type="primary"):
            src_bytes = st.session_state.pdf_bytes
            src_type  = st.session_state.src_type
            pdf_bytes_for_crop = src_bytes if src_type == '.pdf' else None

            with st.spinner("Cropping figures and building documents..."):
                try:
                    captions_map = st.session_state.captions

                    pdf_out  = build_pdf(figs, captions_map,
                                         pdf_bytes=pdf_bytes_for_crop)
                    word_out = build_word(figs, captions_map,
                                          pdf_bytes=pdf_bytes_for_crop)

                    stem = Path(st.session_state.src_name).stem
                    st.success("✓ Both files generated.")

                    dl1, dl2 = st.columns(2)
                    with dl1:
                        st.download_button(
                            "⬇ Download PDF",
                            data=pdf_out,
                            file_name=f"{stem}_figures.pdf",
                            mime="application/pdf",
                            type="primary",
                            use_container_width=True
                        )
                    with dl2:
                        st.download_button(
                            "⬇ Download Word (.docx)",
                            data=word_out,
                            file_name=f"{stem}_figures.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary",
                            use_container_width=True
                        )

                except Exception as e:
                    st.error(f"Generation failed: {e}")
                    st.exception(e)

        # Individual PNG downloads
        with st.expander("Download individual figures as PNG"):
            src_bytes = st.session_state.pdf_bytes
            src_type  = st.session_state.src_type
            pdf_for_crop = src_bytes if src_type == '.pdf' else None
            for i, fig in enumerate(figs):
                if not fig.get('include', True):
                    continue
                c1, c2 = st.columns([4, 1])
                with c1:
                    st.markdown(f"**{fig['label']}**")
                with c2:
                    try:
                        crop = get_cropped_image(fig, pdf_bytes=pdf_for_crop, high_dpi=OUT_DPI)
                        png  = img_to_bytes(crop)
                        st.download_button(
                            "⬇ PNG",
                            data=png,
                            file_name=f"{fig['label'].replace(' ','_')}.png",
                            mime="image/png",
                            key=f"png_{i}"
                        )
                    except Exception as e:
                        st.caption(f"Error: {e}")

elif not uploaded:
    st.markdown("""
    <div class="info-box">
    <b>Supported input formats:</b><br>
    • <b>PDF</b> — digital or scanned published text (chapters, textbook pages)<br>
    • <b>Word (.docx)</b> — documents with embedded images and captions below<br><br>
    <b>Detected caption types:</b>
    Figure / Fig. / Table / Box / Plate / Video — followed by a number<br><br>
    <b>Output:</b> PDF + Word document, one figure per page,
    clean crop, exact caption text, Times New Roman 10pt.
    </div>
    """, unsafe_allow_html=True)
