"""
Citation Toolkit v3 — Complete
All 8 apps in one file.
Run: streamlit run app_v2_new.py
"""
import base64
import html as html_module
import io
import re
import sqlite3
import tempfile
import os
import xml.etree.ElementTree as ET
import zipfile
from copy import deepcopy
from datetime import datetime
from pathlib import Path

import requests
import streamlit as st
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

st.set_page_config(page_title="Citation Toolkit", page_icon="📚",
                   layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Libre+Baskerville:ital,wght@0,400;0,700;1,400&family=Source+Sans+3:wght@300;400;500;600&family=Source+Serif+4:ital,opsz,wght@0,8..60,300;0,8..60,400;1,8..60,300;1,8..60,400&display=swap');

:root {
  --bg:         #f5f0e8;
  --bg2:        #ede8dc;
  --surface:    #faf7f2;
  --border:     #d4c9b5;
  --border-soft:#e2dace;
  --ink:        #1a1510;
  --ink-mid:    #4a3f30;
  --ink-dim:    #8a7a65;
  --ink-faint:  #b5a890;
  --accent:     #8b1a1a;
  --accent2:    #c0392b;
  --green:      #2d6a4f;
  --amber:      #b5541c;
  --radius:     4px;
  --shadow:     0 1px 4px rgba(0,0,0,0.10);
}

html, body, [class*="css"] {
  font-family: 'Source Sans 3', Georgia, sans-serif;
  background: var(--bg) !important;
  color: var(--ink);
}
h1,h2,h3,h4 {
  font-family: 'Libre Baskerville', Georgia, serif;
  color: var(--ink);
  font-weight: 400;
}
h2 {
  font-size: 2.1rem;
  letter-spacing: -0.02em;
  line-height: 1.2;
  margin-bottom: 0.3rem;
  border: none;
}
h2 em, h2 i {
  color: var(--accent);
  font-style: italic;
}
h3 { font-size: 1.25rem; font-weight: 400; color: var(--ink-mid); }
p  { color: var(--ink-mid); line-height: 1.75; font-size: 1.02rem; }

/* Main layout */
.main { background: var(--bg) !important; }
.main .block-container {
  padding: 2.8rem 3.5rem 5rem;
  max-width: 980px;
  background: var(--bg);
}

/* Sidebar */
section[data-testid="stSidebar"] {
  background: var(--surface) !important;
  border-right: 1px solid var(--border) !important;
}
section[data-testid="stSidebar"] * { color: var(--ink-mid) !important; }
section[data-testid="stSidebar"] .stRadio label {
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.9rem !important;
  padding: 6px 4px !important;
  color: var(--ink-mid) !important;
  cursor: pointer;
}
section[data-testid="stSidebar"] .stRadio > div { gap: 2px !important; }

/* Sidebar logo area */
.sidebar-logo {
  font-family: 'Libre Baskerville', Georgia, serif;
  font-size: 1.25rem;
  color: var(--ink) !important;
  font-weight: 400;
  line-height: 1.2;
}
.sidebar-logo em { color: var(--accent) !important; font-style: italic; }
.sidebar-sub {
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.72rem;
  letter-spacing: 0.14em;
  text-transform: uppercase;
  color: var(--ink-faint) !important;
  margin-top: 4px;
  margin-bottom: 1.2rem;
}
.sidebar-rule {
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.65rem;
  letter-spacing: 0.18em;
  text-transform: uppercase;
  color: var(--ink-faint) !important;
  display: flex;
  align-items: center;
  gap: 8px;
  margin: 1.2rem 0 0.5rem;
}
.sidebar-rule::after {
  content: '';
  flex: 1;
  height: 1px;
  background: var(--border);
}

/* App pill badge */
.app-label {
  display: inline-block;
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.68rem;
  letter-spacing: 0.18em;
  text-transform: uppercase;
  color: var(--ink-faint);
  border-top: 2px solid var(--border);
  border-bottom: 2px solid var(--border);
  padding: 4px 0;
  margin-bottom: 0.6rem;
  display: flex;
  align-items: center;
  gap: 10px;
}
.app-label::before, .app-label::after {
  content: '—';
  color: var(--ink-faint);
  font-size: 0.8rem;
}

/* Instruction box */
.instruction-box {
  background: var(--surface);
  border: 1px solid var(--border-soft);
  border-left: 3px solid var(--accent);
  border-radius: var(--radius);
  padding: 1rem 1.3rem;
  margin: 0.8rem 0 1.6rem;
  font-size: 0.9rem;
  color: var(--ink-mid);
  line-height: 1.8;
}
.instruction-box b  { color: var(--ink); font-weight: 600; }
.instruction-box code {
  background: var(--bg2);
  padding: 1px 5px;
  border-radius: 3px;
  font-family: 'Source Code Pro', monospace;
  font-size: 0.82rem;
  color: var(--accent);
  border: 1px solid var(--border-soft);
}
.instruction-box ul { margin: 0.3rem 0 0; padding-left: 1.2rem; }
.instruction-box li { margin-bottom: 0.25rem; }

/* Cards */
.card {
  background: var(--surface);
  border: 1px solid var(--border);
  border-radius: var(--radius);
  padding: 1.2rem 1.5rem;
  margin-bottom: 0.9rem;
  box-shadow: var(--shadow);
}
.card-accent { border-left: 3px solid var(--accent); }
.card-green  { border-left: 3px solid var(--green); background: #f2f7f4; }
.card-red    { border-left: 3px solid #c0392b; background: #fdf5f5; }
.card-amber  { border-left: 3px solid var(--amber); background: #fdf8f2; }

/* Reference items */
.ref-item {
  padding: 0.5rem 0.85rem;
  border-radius: var(--radius);
  margin-bottom: 5px;
  font-size: 0.85rem;
  border-left: 3px solid;
  line-height: 1.5;
  color: var(--ink-mid);
}
.ref-item.error   { border-color: #c0392b; background: #fdf5f5; }
.ref-item.warning { border-color: var(--amber); background: #fdf8f2; }
.ref-item.ok      { border-color: var(--green); background: #f2f7f4; }

/* Section divider */
.section-rule {
  display: flex;
  align-items: center;
  gap: 10px;
  margin: 2rem 0 1.2rem;
  font-family: 'Source Sans 3', sans-serif;
  font-size: 0.68rem;
  letter-spacing: 0.18em;
  text-transform: uppercase;
  color: var(--ink-faint);
}
.section-rule::before, .section-rule::after {
  content: ''; flex: 1; height: 1px; background: var(--border);
}

/* Match cards */
.match-card {
  background: var(--surface);
  border: 1px solid var(--border);
  border-left: 3px solid var(--accent);
  border-radius: var(--radius);
  padding: 1rem 1.2rem;
  margin-bottom: 0.7rem;
  box-shadow: var(--shadow);
}
.match-card.accepted { border-left-color: var(--green); background: #f2f7f4; }
.match-card.skipped  { border-left-color: #c0392b; background: #fdf5f5; }
.match-sentence { font-size: 0.87rem; color: var(--ink-dim); font-style: italic; margin-bottom: 0.5rem; line-height: 1.7; }
.match-marker { font-family: monospace; font-size: 0.75rem; background: var(--bg2); color: var(--accent); padding: 2px 7px; border-radius: 3px; border: 1px solid var(--border-soft); display: inline-block; margin-bottom: 6px; }
.score-pill { font-family: monospace; font-size: 0.7rem; padding: 2px 8px; border-radius: 3px; border: 1px solid; }
.score-high { background: #f2f7f4; color: var(--green); border-color: #a8d4be; }
.score-mid  { background: #fdf8f2; color: var(--amber); border-color: #e0c8a8; }
.score-low  { background: #fdf5f5; color: #c0392b; border-color: #e0b8b8; }

/* Streamlit overrides */
.stApp { background: var(--bg) !important; }
div[data-testid="stMetric"] {
  background: var(--surface);
  border: 1px solid var(--border);
  border-radius: var(--radius);
  padding: 0.8rem 1rem;
  box-shadow: var(--shadow);
}
div[data-testid="stMetric"] label { color: var(--ink-dim) !important; font-size: 0.78rem !important; font-family: 'Source Sans 3', sans-serif !important; letter-spacing: 0.05em !important; text-transform: uppercase !important; }
div[data-testid="stMetric"] div[data-testid="stMetricValue"] { font-family: 'Libre Baskerville', serif !important; font-size: 2rem !important; font-weight: 400 !important; color: var(--ink) !important; }

button[kind="primary"] {
  background: var(--ink) !important;
  color: var(--bg) !important;
  border: none !important;
  border-radius: 0 !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.82rem !important;
  font-weight: 600 !important;
  letter-spacing: 0.1em !important;
  text-transform: uppercase !important;
  padding: 0.5rem 1.4rem !important;
  transition: background 0.15s !important;
}
button[kind="primary"]:hover { background: var(--accent) !important; }
button[kind="secondary"] {
  background: transparent !important;
  border: 1px solid var(--border) !important;
  border-radius: 0 !important;
  color: var(--ink-mid) !important;
  font-family: 'Source Sans 3', sans-serif !important;
  font-size: 0.82rem !important;
  letter-spacing: 0.06em !important;
}
button[kind="secondary"]:hover { border-color: var(--ink-mid) !important; }

div[data-testid="stFileUploader"] {
  border: 1px dashed var(--border) !important;
  border-radius: var(--radius) !important;
  background: var(--surface) !important;
}
div[data-testid="stFileUploader"]:hover { border-color: var(--accent) !important; }

div[data-testid="stExpander"] {
  background: var(--surface) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  box-shadow: var(--shadow) !important;
  margin-bottom: 0.6rem !important;
}
div[data-testid="stExpander"] summary {
  font-size: 0.92rem !important;
  color: var(--ink) !important;
  font-family: 'Source Sans 3', sans-serif !important;
}

div[data-testid="stTextInput"] input,
div[data-testid="stTextArea"] textarea,
div[data-testid="stNumberInput"] input {
  background: var(--surface) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  color: var(--ink) !important;
  font-family: 'Source Sans 3', sans-serif !important;
}
div[data-testid="stSelectbox"] > div {
  background: var(--surface) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
}
.stDataFrame { border: 1px solid var(--border) !important; border-radius: var(--radius) !important; }
.stCaption { color: var(--ink-dim) !important; font-size: 0.8rem !important; }
div[data-testid="stProgress"] > div { background: var(--bg2) !important; }
div[data-testid="stProgress"] > div > div { background: var(--accent) !important; }
div[data-testid="stAlert"] { border-radius: var(--radius) !important; }
div[data-testid="stTabs"] button { font-family: 'Source Sans 3', sans-serif !important; font-size: 0.87rem !important; color: var(--ink-mid) !important; }
div[data-testid="stTabs"] button[aria-selected="true"] { color: var(--accent) !important; border-bottom-color: var(--accent) !important; }
div[data-testid="stRadio"] label { color: var(--ink-mid) !important; font-family: 'Source Sans 3', sans-serif !important; font-size: 0.9rem !important; }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────
TOP_N = 5
TFIDF_THRESHOLD = 0.12
PUBMED_MAX = 5
MISSING_PATTERNS = [
    r'\[CITATION\]', r'\[REF\]', r'\[ref\]', r'\[\?\]',
    r'\[citation needed\]', r'\bXXX\b', r'\[#\]', r'<citation>', r'\[ *\]',
]
CITATION_MARKERS = re.compile('|'.join(MISSING_PATTERNS), re.IGNORECASE)
MATCH_THRESHOLD = 0.28
FUZZY_THRESHOLD = 0.08
PUBMED_ESEARCH = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
PUBMED_ESUM    = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi"

# ─────────────────────────────────────────────────────────────────────────────
# SHARED HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def xml_text(elem, path):
    node = elem.find(path)
    if node is None: return ""
    return " ".join(p.strip() for p in ([node.text or ""] + [c.text or "" for c in node]) if p.strip())

def fmt_ref(ref, short=False):
    aa = ref.get("authors", [])
    if aa:
        a = (aa[0].split(",")[0] if len(aa)==1 else
             f"{aa[0].split(',')[0]} & {aa[1].split(',')[0]}" if len(aa)==2 else
             f"{aa[0].split(',')[0]} et al.")
    else: a = "Unknown"
    y = ref.get("year","n.d.")
    t = ref.get("title","")[:90] + ("..." if len(ref.get("title",""))>90 else "")
    return f"{a} ({y}) — {t}" if short else f"{a} ({y}). {t}. {ref.get('journal','')}"

def score_class(s):
    return "score-high" if s>=0.20 else "score-mid" if s>=0.10 else "score-low"

def doc_to_bytes(doc):
    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()

def _enl_table(cursor):
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
    tables = [t[0] for t in cursor.fetchall()]
    return 'refs' if 'refs' in tables else 'enl_refs'

# ─────────────────────────────────────────────────────────────────────────────
# APP 2 LOGIC — BROKEN CITATION FIXER
# ─────────────────────────────────────────────────────────────────────────────
def analyze_docx_citations(docx_bytes):
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        raw = z.read('word/document.xml').decode('utf-8')
    total    = raw.count('ADDIN EN.CITE')
    working  = len(re.findall(r'&lt;EndNote&gt;', raw))
    empty    = len(re.findall(r'<w:instrText[^>]*> ADDIN EN\.CITE </w:instrText>', raw))
    flddata  = raw.count('<w:fldData')
    db_ids   = list(set(re.findall(r'&lt;key[^&]*db-id=&quot;([^&]+)&quot;', raw)))

    # Count unique RecNums EndNote can currently see (working fields only)
    working_rns = set(re.findall(r'&lt;RecNum&gt;(\d+)&lt;/RecNum&gt;', raw))

    # Count unique RecNums locked in broken fldData (invisible to EndNote)
    import base64 as _b64
    fld_pat = re.compile(r'<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>')
    fld_rns = set()
    for b64r in fld_pat.findall(raw):
        b64 = b64r.replace('\r','').replace('\n','').replace(' ','')
        pad = (4-len(b64)%4)%4
        try:
            dec = _b64.b64decode(b64+'='*pad).decode('utf-8',errors='replace').replace('\x00','')
            for rn in re.findall(r'<RecNum>(\d+)</RecNum>', dec):
                fld_rns.add(rn)
        except: pass
    hidden_rns = fld_rns - working_rns   # RecNums only in fldData, invisible to EndNote

    # Count bibliography entries
    try:
        _doc = Document(io.BytesIO(docx_bytes))
        _ref_pat = re.compile(r'^\s*(\d+)[\.)\s]\s+')
        bib_count = sum(1 for p in _doc.paragraphs if _ref_pat.match(p.text.strip()))
    except: bib_count = 0

    return dict(raw=raw, total_fields=total, working=working,
                broken_empty=empty, flddata_count=flddata, db_ids=db_ids,
                working_rns=working_rns, hidden_rns=hidden_rns,
                endnote_sees=len(working_rns),
                endnote_misses=len(hidden_rns),
                bib_count=bib_count)

def fix_broken_fields(raw_xml):
    """
    Fixes three patterns of broken EndNote citation field codes:
    A: instrText says 'ADDIN EN.CITE' but is empty — data only in fldData
    B: instrText says 'ADDIN EN.CITE.DATA' — data only in fldData
    C: no instrText at all — fldChar begin immediately before fldData
    All three cause EndNote to undercount references in the bibliography.
    """
    fixes = 0

    # Pattern A: empty ADDIN EN.CITE instrText with fldData nearby
    pat_a = re.compile(
        r'(<w:instrText[^>]*>) ADDIN EN\.CITE (</w:instrText>)'
        r'([\s\S]{0,2000}?)<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>', re.DOTALL)
    def rep_a(m):
        nonlocal fixes
        io_, ic, between, b64r = m.group(1), m.group(2), m.group(3), m.group(4)
        b64 = b64r.replace('\r','').replace('\n','').replace(' ','')
        pad = (4-len(b64)%4)%4
        try:
            dec = base64.b64decode(b64+'='*pad).decode('utf-8', errors='replace')
            dec = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', dec)
            esc = dec.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')
            fixes += 1
            return f'{io_} ADDIN EN.CITE {esc}{ic}{between}<w:fldData xml:space="preserve">{b64r}</w:fldData>'
        except: return m.group(0)
    result = pat_a.sub(rep_a, raw_xml)

    # Pattern B: instrText says "ADDIN EN.CITE.DATA" instead of full XML
    pat_b = re.compile(
        r'(<w:instrText[^>]*>)\s*ADDIN EN\.CITE\.DATA\s*(</w:instrText>)'
        r'([\s\S]{0,2000}?)<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>', re.DOTALL)
    def rep_b(m):
        nonlocal fixes
        io_, ic, between, b64r = m.group(1), m.group(2), m.group(3), m.group(4)
        b64 = b64r.replace('\r','').replace('\n','').replace(' ','')
        pad = (4-len(b64)%4)%4
        try:
            dec = base64.b64decode(b64+'='*pad).decode('utf-8', errors='replace')
            dec = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', dec)
            esc = dec.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')
            fixes += 1
            return f'{io_} ADDIN EN.CITE {esc}{ic}{between}<w:fldData xml:space="preserve">{b64r}</w:fldData>'
        except: return m.group(0)
    result = pat_b.sub(rep_b, result)

    # Pattern B (reverse): fldData appears BEFORE the instrText (nested field variant)
    # Word sometimes stores: fldChar(begin) + fldData(data) + instrText(ADDIN EN.CITE.DATA) + fldChar(end)
    # The cite data is in the fldData block that precedes the marker instrText.
    def _rep_b_rev(m):
        nonlocal fixes
        b64r = m.group(1)
        b64  = b64r.replace('\r','').replace('\n','').replace(' ','')
        pad  = (4 - len(b64) % 4) % 4
        try:
            dec = base64.b64decode(b64 + '=' * pad).decode('utf-8', errors='replace')
            dec = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', dec)
            if '<EndNote>' not in dec: return m.group(0)
            esc = dec.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')
            fixes += 1
            return (f'<w:fldData xml:space="preserve">{b64r}</w:fldData>'
                    f'{m.group(2)}'
                    f'{m.group(3)} ADDIN EN.CITE {esc}{m.group(4)}')
        except: return m.group(0)
    _pat_b_rev = re.compile(
        r'<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>'
        r'([\s\S]{0,500}?)'
        r'(<w:instrText[^>]*>)\s*ADDIN EN\.CITE\.DATA\s*(</w:instrText>)',
        re.DOTALL)
    result = _pat_b_rev.sub(_rep_b_rev, result)


    # Pattern C: no instrText at all — insert clean field before the malformed structure
    fld_pat = re.compile(r'<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>')
    for _ in range(50):
        inserted = False
        for m in fld_pat.finditer(result):
            b64r = m.group(1)
            b64  = b64r.replace('\r','').replace('\n','').replace(' ','')
            pad  = (4-len(b64)%4)%4
            try:
                dec = base64.b64decode(b64+'='*pad).decode('utf-8', errors='replace').replace('\x00','')
            except: continue
            if '<EndNote>' not in dec: continue
            rns = set(re.findall(r'<RecNum>(\d+)</RecNum>', dec))
            if not any(f'&lt;RecNum&gt;{rn}&lt;/RecNum&gt;' not in result for rn in rns):
                continue
            begin_pos = result.rfind('<w:fldChar w:fldCharType="begin"', 0, m.start())
            if begin_pos < 0: continue
            rs1 = result.rfind('<w:r ', 0, begin_pos)
            rs2 = result.rfind('<w:r>', 0, begin_pos)
            run_start = max(rs1, rs2)
            if run_start < 0: continue
            dec_clean = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', dec)
            esc = dec_clean.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')
            clean = (
                f'<w:r><w:rPr><w:noProof/><w:vertAlign w:val="superscript"/></w:rPr>'
                f'<w:fldChar w:fldCharType="begin"/></w:r>'
                f'<w:r><w:instrText xml:space="preserve"> ADDIN EN.CITE {esc}</w:instrText></w:r>'
                f'<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
                f'<w:r><w:rPr><w:noProof/><w:vertAlign w:val="superscript"/></w:rPr>'
                f'<w:t></w:t></w:r>'
                f'<w:r><w:fldChar w:fldCharType="end"/></w:r>'
            )
            result = result[:run_start] + clean + result[run_start:]
            fixes += 1
            inserted = True
            break
        if not inserted: break

    result = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', result)

    # Fix fields missing their separate marker (begin->instrText->end without separate)
    # Word requires: begin -> instrText -> separate -> display -> end
    _SEP_RUN = ('<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
                '<w:r><w:rPr><w:noProof/><w:vertAlign w:val="superscript"/></w:rPr>'
                '<w:t></w:t></w:r>')
    result = re.sub(
        r'(</w:instrText></w:r>)'
        r'((?:(?!fldCharType="separate"|fldCharType="end"|</w:instrText>).)*?)'
        r'(<w:r[^>]*>(?:[^<]|<(?!w:fldChar))*?<w:fldChar w:fldCharType="end"/>)',
        lambda _m: (_m.group(1) + _m.group(2) + _SEP_RUN + _m.group(3)
                    if 'fldCharType="separate"' not in _m.group(2) else _m.group(0)),
        result, flags=re.DOTALL
    )

    # Deduplicate w:id attributes — paragraph copying causes ID clashes
    # that make Word refuse to open the file ("unreadable content")
    from collections import Counter as _Counter
    _all_ids = re.findall(r'\bw:id="(\d+)"', result)
    if _all_ids:
        _max_id  = max(int(x) for x in _all_ids)
        _next_id = [_max_id + 1]; _seen_ids = set()
        def _fix_wid(m):
            v = m.group(2)
            if v in _seen_ids:
                nw = str(_next_id[0]); _next_id[0] += 1
                return f'{m.group(1)}{nw}{m.group(3)}'
            _seen_ids.add(v); return m.group(0)
        result = re.sub(r'(w:id=")(\d+)(")', _fix_wid, result)

    # Remove orphaned plain-text citation numbers left from merges
    # (depth-0 pure-digit/comma runs in paragraphs that have EN.CITE fields)
    # We do this at XML level to avoid needing python-docx here
    import re as _re
    _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    try:
        from lxml import etree as _etree
        from docx import Document as _DocX
        import io as _io
        # Quick pass: remove depth-0 digit-only runs in cite paragraphs
        _root = _etree.fromstring(result.encode('utf-8'))
        _orphans_removed = 0
        for _para in _root.findall(f'.//{{{_W}}}p'):
            _has_cite = any('EN.CITE' in (_i.text or '')
                            for _i in _para.findall(f'.//{{{_W}}}instrText'))
            if not _has_cite: continue
            _fd = 0; _to_del = []
            for _run in list(_para):
                _tag = _run.tag.split('}')[-1] if '}' in _run.tag else _run.tag
                if _tag != 'r': continue
                for _fc in _run.findall(f'.//{{{_W}}}fldChar'):
                    _ft = _fc.get(f'{{{_W}}}fldCharType', '')
                    if _ft == 'begin': _fd += 1
                    elif _ft == 'end': _fd = max(0, _fd - 1)
                if _fd > 0: continue
                _t = _run.find(f'{{{_W}}}t')
                _txt = (_t.text or '') if _t is not None else ''
                if _txt.strip() and _re.match(r'^[\d,;\s]+$', _txt.strip()):
                    _to_del.append(_run)
            for _run in _to_del:
                _para.remove(_run); _orphans_removed += 1
        if _orphans_removed:
            result = _etree.tostring(_root, xml_declaration=True,
                                     encoding='UTF-8', standalone=True).decode('utf-8')
            fixes += _orphans_removed
    except Exception: pass  # non-critical; don't break the main fix

    # Fix garbled figure numbers: merge sometimes concatenates chapter numbers
    # resulting in "33" + "34" across runs -> shows as "Fig. 3334.X"
    # (This is handled in safe_merge_documents for full merges,
    #  but also apply here for standalone repairs)

    return result, fixes


def remove_orphan_superscripts(docx_bytes):
    """
    Remove orphaned plain-text citation numbers left over after merges.

    Merges leave behind three types of plain-text citation remnants in paragraphs
    that also have working EndNote field codes:

    1. Superscript runs (vertAlign=superscript) outside field codes — the most
       obvious: numbers displayed as superscripts next to field codes.

    2. Style-based superscript runs (rStyle=citsup or rStyle=sup) outside field
       codes — same problem but using Word character styles instead of direct
       formatting.

    3. Depth-0 plain digit runs anywhere in the paragraph — numbers that are not
       superscripted but are purely citation digits/commas sitting outside field
       codes in a paragraph that has working EN.CITE fields. These are harder to
       detect visually but still produce duplicate citation display.

    All three patterns: only removes runs containing ONLY digits, commas,
    semicolons, and spaces. Never removes runs with other content.

    Returns (fixed_bytes, n_removed)
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    doc = Document(io.BytesIO(docx_bytes))
    removed = 0

    for para in doc.paragraphs:
        has_field = any(
            'EN.CITE' in (instr.text or '')
            for instr in para._p.findall(f'.//{{{W}}}instrText')
        )
        if not has_field:
            continue

        field_depth = 0
        to_remove   = []

        for run in list(para._p):
            tag = run.tag.split('}')[-1] if '}' in run.tag else run.tag
            if tag != 'r': continue

            for fc in run.findall(f'.//{{{W}}}fldChar'):
                ft = fc.get(f'{{{W}}}fldCharType', '')
                if ft == 'begin':   field_depth += 1
                elif ft == 'end':   field_depth = max(0, field_depth - 1)

            if field_depth > 0:
                continue  # inside a field — leave alone

            t    = run.find(f'{{{W}}}t')
            text = (t.text or '') if t is not None else ''
            if not text.strip():
                continue
            # Must be purely citation-like: digits, commas, semicolons, spaces
            if not re.match(r'^[\d,;\s]+$', text.strip()):
                continue

            # Pattern 1 & 2: superscript by vertAlign or character style
            rpr = run.find(f'{{{W}}}rPr')
            is_super = False
            if rpr is not None:
                va = rpr.find(f'{{{W}}}vertAlign')
                rs = rpr.find(f'{{{W}}}rStyle')
                if va is not None and va.get(f'{{{W}}}val') == 'superscript':
                    is_super = True
                if rs is not None and rs.get(f'{{{W}}}val', '') in ('citsup', 'sup', 'superscript'):
                    is_super = True

            # Pattern 3: plain depth-0 digit run in a para that has cite fields
            # (catches runs that survived the merge without any superscript formatting)
            # Always remove any depth-0 pure-digit run in a cite paragraph.
            to_remove.append(run)

        for run in to_remove:
            para._p.remove(run)
            removed += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read(), removed


def extract_karol_db_id(xml_bytes):
    try:
        content = xml_bytes.decode('utf-8',errors='replace')
        m = re.search(r'db-id="([a-z0-9]{20,45})"',content,re.IGNORECASE)
        return m.group(1) if m else None
    except: return None

def get_karol_rec_nums(enl_bytes):
    with tempfile.NamedTemporaryFile(suffix='.enl',delete=False) as f:
        f.write(enl_bytes); tmp=f.name
    try:
        conn=sqlite3.connect(tmp); cursor=conn.cursor()
        tbl=_enl_table(cursor)
        cursor.execute(f'SELECT id,author,year,title FROM {tbl} WHERE trash_state=0 OR trash_state IS NULL')
        rows=cursor.fetchall(); conn.close()
        return {str(r[0]):{'id':str(r[0]),'author':r[1] or '','year':str(r[2] or ''),'title':r[3] or ''} for r in rows}
    except: return {}
    finally: os.unlink(tmp)

def check_missing_from_karol(raw_xml, karol_ids):
    all_rns = set(re.findall(r'&lt;RecNum&gt;(\d+)&lt;/RecNum&gt;', raw_xml))
    return [rn for rn in all_rns if rn not in karol_ids]

def patch_db_ids(raw_xml, old_ids, new_id):
    result=raw_xml; replaced=0
    for old in old_ids:
        if old!=new_id:
            count=result.count(old); result=result.replace(old,new_id); replaced+=count
    return result, replaced

def build_fixed_docx(original_bytes, fixed_xml):
    with zipfile.ZipFile(io.BytesIO(original_bytes)) as z:
        all_files={n:z.read(n) for n in z.namelist()}
    all_files['word/document.xml']=fixed_xml.encode('utf-8')
    buf=io.BytesIO()
    with zipfile.ZipFile(buf,'w',zipfile.ZIP_DEFLATED) as zout:
        for n,d in all_files.items(): zout.writestr(n,d)
    buf.seek(0); return buf.read()

def extract_traveling_library_xml(docx_bytes):
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        raw=z.read('word/document.xml').decode('utf-8')
    all_cx=[]
    for m in re.findall(r'ADDIN EN\.CITE &lt;EndNote&gt;([\s\S]+?)&lt;/EndNote&gt;',raw):
        all_cx.append(html_module.unescape(f'<EndNote>{m}</EndNote>'))
    for b64r in re.findall(r'<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>',raw):
        b64=b64r.replace('\r','').replace('\n','').replace(' ','')
        pad=(4-len(b64)%4)%4
        try:
            dec=base64.b64decode(b64+'='*pad).decode('utf-8',errors='replace').replace('\x00','')
            if '<EndNote>' in dec: all_cx.append(dec)
        except: pass
    from lxml import etree
    import html as hmod
    traveling={}
    for cx in all_cx:
        try:
            if not cx.startswith('<EndNote>'): cx=f'<EndNote>{cx}</EndNote>'
            root=etree.fromstring(cx.encode('utf-8'))
            for cite in root.findall('.//Cite'):
                rn=cite.findtext('RecNum') or ''
                rec=cite.find('record')
                if rn and rec is not None and rn not in traveling:
                    traveling[rn]=rec
        except: pass
    def gt(elem,tag):
        n=elem.find(f'.//{tag}')
        return ''.join(n.itertext()).strip() if n is not None else ''
    output='<?xml version="1.0" encoding="UTF-8"?>\n<xml>\n  <records>\n'
    for rn,record in sorted(traveling.items(),key=lambda x:int(x[0]) if x[0].isdigit() else 9999):
        rte=record.find('.//ref-type')
        rtn=rte.get('name','Journal Article') if rte is not None else 'Journal Article'
        rtv=rte.text if rte is not None else '17'
        authors=[''.join(a.itertext()).strip() for a in record.findall('.//contributors/authors/author')]
        secauths=[''.join(a.itertext()).strip() for a in record.findall('.//contributors/secondary-authors/author')]
        r=f'    <record>\n      <rec-number>{rn}</rec-number>\n      <ref-type name="{hmod.escape(rtn)}">{rtv}</ref-type>\n'
        if authors:
            r+='      <contributors>\n        <authors>\n'
            for a in authors:
                if a: r+=f'          <author>{hmod.escape(a)}</author>\n'
            r+='        </authors>\n'
            if secauths:
                r+='        <secondary-authors>\n'
                for a in secauths:
                    if a: r+=f'          <author>{hmod.escape(a)}</author>\n'
                r+='        </secondary-authors>\n'
            r+='      </contributors>\n'
        r+='      <titles>\n'
        for tag,xmltag in [('title','title'),('secondary-title','secondary-title'),('tertiary-title','tertiary-title')]:
            v=gt(record,tag)
            if v: r+=f'        <{xmltag}>{hmod.escape(v)}</{xmltag}>\n'
        r+='      </titles>\n'
        for tag,xmltag in [('year','dates><year'),('volume','volume'),('number','number'),
                           ('pages','pages'),('edition','edition'),('publisher','publisher'),
                           ('pub-location','pub-location'),('abstract','abstract')]:
            v=gt(record,tag)
            if v:
                if xmltag=='dates><year': r+=f'      <dates><year>{hmod.escape(v)}</year></dates>\n'
                else: r+=f'      <{xmltag}>{hmod.escape(v)}</{xmltag}>\n'
        kws=[(''.join(k.itertext()).strip()) for k in record.findall('.//keyword')]
        if any(kws):
            r+='      <keywords>\n'
            for kw in kws:
                if kw: r+=f'        <keyword>{hmod.escape(kw)}</keyword>\n'
            r+='      </keywords>\n'
        r+='    </record>\n'
        output+=r
    output+='  </records>\n</xml>\n'
    return output, len(traveling)

def remap_traveling_citations(docx_bytes, enl_bytes):
    from lxml import etree
    import html as hm
    with tempfile.NamedTemporaryFile(suffix='.enl',delete=False) as f:
        f.write(enl_bytes); tmp=f.name
    try:
        conn=sqlite3.connect(tmp); cursor=conn.cursor()
        tbl=_enl_table(cursor)
        cursor.execute(f'SELECT id,author,year,title FROM {tbl} WHERE trash_state=0 OR trash_state IS NULL')
        karol_rows=cursor.fetchall(); conn.close()
    finally: os.unlink(tmp)
    def norm(s): return re.sub(r'[^a-z0-9]','',s.lower()) if s else ''
    def alast(s): return norm(s.split('\r')[0].split('\n')[0].strip().split(',')[0]) if s else ''
    karol_ids=set(str(r[0]) for r in karol_rows)
    by_ay={}
    for row in karol_rows:
        k=(alast(row[1] or ''),norm(str(row[2] or '')))
        by_ay.setdefault(k,[]).append(row)
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        raw=z.read('word/document.xml').decode('utf-8')
        all_files={n:z.read(n) for n in z.namelist()}
    report=[]; remap={}
    for cx_esc in re.findall(r'ADDIN EN\.CITE &lt;EndNote&gt;([\s\S]+?)&lt;/EndNote&gt;',raw):
        cx=hm.unescape(f'<EndNote>{cx_esc}</EndNote>')
        try:
            root=etree.fromstring(cx.encode('utf-8'))
            for cite in root.findall('.//Cite'):
                rn=cite.findtext('RecNum') or ''
                au=cite.findtext('Author') or ''
                yr=cite.findtext('Year') or ''
                if rn in karol_ids or rn in remap: continue
                title=''
                rec=cite.find('record')
                if rec is not None:
                    te=rec.find('.//title')
                    if te is not None: title=''.join(te.itertext())
                key=(alast(au),norm(yr)); cands=by_ay.get(key,[])
                matched=None
                if len(cands)==1: matched=str(cands[0][0])
                elif len(cands)>1:
                    for c in cands:
                        if norm(title)[:25] and norm(str(c[3] or ''))[:25]==norm(title)[:25]:
                            matched=str(c[0]); break
                    if not matched: matched=str(cands[0][0])
                remap[rn]=matched
                report.append({'status':'remapped' if matched else 'not_found',
                               'old_rec_num':rn,'new_rec_num':matched,
                               'author':au,'year':yr,'title':title[:80]})
        except: pass
    if not remap: return docx_bytes, report
    fixed=raw
    for old,new in remap.items():
        if new:
            fixed=fixed.replace(f'&lt;RecNum&gt;{old}&lt;/RecNum&gt;',
                                 f'&lt;RecNum&gt;{new}&lt;/RecNum&gt;')
    def fix_fld(m):
        b64r=m.group(1)
        b64=b64r.replace('\r','').replace('\n','').replace(' ','')
        pad=(4-len(b64)%4)%4
        try:
            dec=base64.b64decode(b64+'='*pad).decode('utf-8',errors='replace').replace('\x00','')
            mod=dec
            for old,new in remap.items():
                if new: mod=mod.replace(f'<RecNum>{old}</RecNum>',f'<RecNum>{new}</RecNum>')
            if mod!=dec:
                nb64=base64.b64encode(mod.encode('utf-8')).decode('ascii')
                wrapped='\r\n'.join(nb64[i:i+76] for i in range(0,len(nb64),76))
                return f'<w:fldData xml:space="preserve">{wrapped}</w:fldData>'
        except: pass
        return m.group(0)
    fixed=re.compile(r'<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>').sub(fix_fld,fixed)
    all_files['word/document.xml']=fixed.encode('utf-8')
    buf=io.BytesIO()
    with zipfile.ZipFile(buf,'w',zipfile.ZIP_DEFLATED) as zout:
        for n,d in all_files.items(): zout.writestr(n,d)
    buf.seek(0)
    return buf.read(), report

def generate_vba_macro(doc_name):
    return f"""' EndNote Citation Relinker — Word VBA Macro
' Generated: {datetime.now():%Y-%m-%d %H:%M}
' HOW TO USE:
'  1. Open "{doc_name}" in Word with your EndNote library open
'  2. Press Alt+F11 → Insert → Module → paste this macro
'  3. Press Alt+F8 → select RelinkAllCitations → Run
Sub RelinkAllCitations()
    Dim oDoc As Document
    Set oDoc = ActiveDocument
    If oDoc Is Nothing Then MsgBox "No document open.", vbCritical: Exit Sub
    Dim sBackup As String
    sBackup = oDoc.Path & "\\{Path(doc_name).stem}_BACKUP_" & Format(Now,"YYYYMMDD_HHMMSS") & ".docx"
    oDoc.SaveAs2 sBackup, wdFormatXMLDocument
    MsgBox "Backup saved: " & sBackup, vbInformation
    On Error Resume Next
    Application.Run "EndNote.UnformatAll"
    If Err.Number <> 0 Then Err.Clear: Application.Run "EndNote.FormatAll": GoTo done
    Err.Clear
    Application.Run "EndNote.FormatAll"
    If Err.Number <> 0 Then Err.Clear: Application.Run "EndNote.UpdateAll"
    done:
    On Error GoTo 0
    oDoc.Save
    MsgBox "Done! Citations re-linked to your library." & vbCrLf & _
           "If any citations are yellow, right-click → Edit Citation → Find.", vbInformation
End Sub
"""

# ─────────────────────────────────────────────────────────────────────────────
# APP 1 LOGIC — CITATION REPAIR
# ─────────────────────────────────────────────────────────────────────────────
@st.cache_data(show_spinner=False)
def parse_endnote_xml_bytes(xml_bytes):
    root=ET.fromstring(xml_bytes); refs=[]
    for rec in root.iter("record"):
        authors=[]
        for a in rec.findall(".//contributors/authors/author"):
            name=" ".join(p.strip() for p in ([a.text or ""]+[c.text or "" for c in a]) if p.strip())
            if name: authors.append(name)
        title=xml_text(rec,".//titles/title")
        journal=xml_text(rec,".//periodical/full-title") or xml_text(rec,".//periodical/abbr-1")
        year=xml_text(rec,".//dates/year"); abstract=xml_text(rec,".//abstract")
        if not title: continue
        corpus=" ".join(filter(None,[title,abstract,journal,year]))
        refs.append(dict(authors=authors,title=title,journal=journal,year=year,corpus=corpus))
    return refs

@st.cache_resource(show_spinner=False)
def build_tfidf(corpora_tuple):
    vec=TfidfVectorizer(ngram_range=(1,2),sublinear_tf=True,max_features=50000)
    mat=vec.fit_transform(list(corpora_tuple)); return vec,mat

def match_sentence(sentence,vec,mat,refs,top_n=TOP_N):
    sv=vec.transform([sentence]); sims=cosine_similarity(sv,mat)[0]
    idx=sims.argsort()[::-1][:top_n]
    return [dict(ref=refs[i],score=float(sims[i])) for i in idx]

def extract_flagged(docx_bytes):
    doc=Document(io.BytesIO(docx_bytes)); flagged=[]
    for pi,para in enumerate(doc.paragraphs):
        text=para.text
        if not text.strip(): continue
        for m in CITATION_MARKERS.finditer(text):
            sents=re.split(r'(?<=[.!?])\s+',text); cum,target=0,text
            for s in sents:
                if cum+len(s)>=m.start(): target=s; break
                cum+=len(s)+1
            flagged.append(dict(para_idx=pi,sentence=target,marker=m.group(),para_text=text))
    return flagged,doc

def author_label(ref):
    aa=ref.get("authors",[]); last=aa[0].split(",")[0].strip().split()[-1] if aa else "Ref"
    return (last+" "+ref.get("year","")).strip()

def insert_superscript(para,marker,label):
    if marker not in para.text: return False
    combined,run_map="",[]
    for run in para.runs:
        s=len(combined); combined+=run.text; run_map.append((s,s+len(run.text),run))
    pos=combined.find(marker)
    if pos==-1: return False
    for (s,e,run) in run_map:
        if s<=pos<e:
            before=run.text[:pos-s]; after=run.text[pos-s+len(marker):]
            run.text=before
            rPr=OxmlElement("w:rPr"); va=OxmlElement("w:vertAlign")
            va.set(qn("w:val"),"superscript"); rPr.append(va)
            nr=OxmlElement("w:r"); nr.append(deepcopy(rPr))
            t=OxmlElement("w:t"); t.text=f"[{label}]"
            t.set("{http://www.w3.org/XML/1998/namespace}space","preserve"); nr.append(t)
            run._r.addnext(nr)
            if after:
                tr=OxmlElement("w:r"); tt=OxmlElement("w:t")
                tt.text=after; tt.set("{http://www.w3.org/XML/1998/namespace}space","preserve")
                tr.append(tt); nr.addnext(tr)
            return True
    return False

def write_repair_report(decisions):
    doc=Document(); doc.add_heading("Citation Repair Report",0)
    doc.add_paragraph(f"Generated: {datetime.now():%Y-%m-%d %H:%M}")
    accepted=[d for d in decisions if d["action"]=="accepted"]
    skipped=[d for d in decisions if d["action"]=="skipped"]
    doc.add_paragraph(f"Total:{len(decisions)} | Accepted:{len(accepted)} | Skipped:{len(skipped)}")
    if accepted:
        doc.add_heading("Accepted",1)
        for d in accepted:
            p=doc.add_paragraph(style="List Bullet")
            p.add_run(f"Marker:{d['marker']}\n").bold=True
            p.add_run(f"Context:{d['sentence'][:200]}\n")
            p.add_run(f"Inserted:{fmt_ref(d['ref'])}\nScore:{d['score']:.3f}")
    if skipped:
        doc.add_heading("Skipped",1)
        for d in skipped:
            p=doc.add_paragraph(style="List Bullet")
            r=p.add_run("NEEDS REVIEW\n"); r.font.color.rgb=RGBColor(0xC0,0,0)
            p.add_run(f"Marker:{d['marker']}\nContext:{d['sentence'][:200]}\n")
    return doc_to_bytes(doc)

# ─────────────────────────────────────────────────────────────────────────────
# APP 3 LOGIC — REFERENCE COMPARATOR
# ─────────────────────────────────────────────────────────────────────────────
def load_ref_file(f):
    name=f.name; data=f.read()
    if name.endswith(".xml"):
        refs=parse_endnote_xml_bytes(data); return refs,name
    elif name.endswith(".docx"):
        doc=Document(io.BytesIO(data)); refs=[]; in_refs=False
        pat=re.compile(r'^\s*\d+[\.\)]\s+(.+)')
        for para in doc.paragraphs:
            text=para.text.strip()
            if not text: continue
            if re.match(r'^(references?|bibliography)$',text,re.IGNORECASE): in_refs=True; continue
            if in_refs or pat.match(text):
                in_refs=True; m=pat.match(text); rt=m.group(1) if m else text
                ym=re.search(r'\b(19|20)\d{2}\b',rt)
                refs.append(dict(authors=[],title=rt[:200],journal="",
                                 year=ym.group(0) if ym else "",corpus=rt,id=str(len(refs)+1)))
        return refs,name
    elif name.endswith(".txt"):
        content=data.decode("utf-8",errors="replace"); refs=[]
        for i,block in enumerate(re.split(r'\n{2,}',content)):
            block=block.strip()
            if len(block)<20: continue
            m=re.match(r'^\d+[\.\)]\s+(.*)',block,re.DOTALL); text=m.group(1) if m else block
            ym=re.search(r'\b(19|20)\d{2}\b',text)
            refs.append(dict(authors=[],title=text[:200],journal="",
                             year=ym.group(0) if ym else "",corpus=text,id=str(i)))
        return refs,name
    return [],name

# ─────────────────────────────────────────────────────────────────────────────
# APP 4 LOGIC — DOCUMENT MERGER
# ─────────────────────────────────────────────────────────────────────────────
def analyze_merge_damage(merged_bytes):
    from lxml import etree as _etree
    import base64 as _b64
    with zipfile.ZipFile(io.BytesIO(merged_bytes)) as z:
        raw=z.read('word/document.xml').decode('utf-8')
    total=raw.count('ADDIN EN.CITE')
    working=len(re.findall(r'&lt;EndNote&gt;',raw))
    empty=len(re.findall(r'<w:instrText[^>]*> ADDIN EN\.CITE </w:instrText>',raw))
    begins=raw.count('fldCharType="begin"')
    separates=raw.count('fldCharType="separate"')
    ends=raw.count('fldCharType="end"')
    has_tracked='<w:ins ' in raw or '<w:del ' in raw
    ins_count=raw.count('<w:ins ')
    del_count=raw.count('<w:del ')
    db_ids=list(set(re.findall(r'&lt;key[^&]*db-id=&quot;([^&]+)&quot;',raw)))

    # RecNums in working field codes (what EndNote currently sees)
    working_rns = set(re.findall(r'&lt;RecNum&gt;(\d+)&lt;/RecNum&gt;', raw))

    # RecNums locked in broken fldData (invisible to EndNote)
    fld_pat = re.compile(r'<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>')
    fld_rns = set()
    for b64r in fld_pat.findall(raw):
        b64 = b64r.replace('\r','').replace('\n','').replace(' ','')
        pad = (4-len(b64)%4)%4
        try:
            dec = _b64.b64decode(b64+'='*pad).decode('utf-8',errors='replace').replace('\x00','')
            for rn in re.findall(r'<RecNum>(\d+)</RecNum>', dec):
                fld_rns.add(rn)
        except: pass
    hidden_rns = fld_rns - working_rns

    # Count ALL superscript number runs
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    try:
        root = _etree.fromstring(raw.encode('utf-8'))
        super_nums = set()
        for r in root.iter(f'{{{W}}}r'):
            rpr = r.find(f'{{{W}}}rPr')
            if rpr is None: continue
            va = rpr.find(f'{{{W}}}vertAlign')
            if va is None: continue
            if va.get(f'{{{W}}}val') != 'superscript': continue
            texts = []
            for t in r.iter(f'{{{W}}}t'):
                if t.text: texts.append(t.text)
            for t in r.iter(f'{{{W}}}delText'):
                if t.text: texts.append(t.text)
            text = ''.join(texts).strip()
            for part in re.split(r'[,;\s]+', text):
                p = part.strip().rstrip('.')
                if p.isdigit(): super_nums.add(int(p))
    except: super_nums = set()

    # Bibliography entries
    ref_pat = re.compile(r'^\s*(\d+)[\.)\s]\s+')
    try:
        doc = Document(io.BytesIO(merged_bytes))
        bib_nums = set()
        for p in doc.paragraphs:
            m = ref_pat.match(p.text.strip())
            if m: bib_nums.add(int(m.group(1)))
    except: bib_nums = set()

    lost_count  = len(bib_nums - super_nums) if bib_nums else 0
    bib_count   = len(bib_nums)
    cited_count = len(super_nums)

    # Check for w:id duplicates (cause Word to reject file with "unreadable content")
    from collections import Counter as _Ctr
    all_wids = re.findall(r'\bw:id="(\d+)"', raw)
    dup_ids  = len({k for k,v in _Ctr(all_wids).items() if v > 1})

    # Check footnote/comment reference integrity
    fn_refs  = set(re.findall(r'<w:footnoteReference\b[^>]*w:id="(\d+)"', raw))
    com_refs = len(re.findall(r'<w:comment(?:RangeStart|RangeEnd|Reference)\b', raw))

    return dict(raw=raw, total_en=total, with_data=working, empty_cite=empty,
                begins=begins, separates=separates, ends=ends,
                balanced=(begins==separates==ends),
                has_tracked=has_tracked, ins_count=ins_count, del_count=del_count,
                db_ids=db_ids,
                working_rns=working_rns, hidden_rns=hidden_rns,
                endnote_sees=len(working_rns),
                endnote_misses=len(hidden_rns),
                super_nums=super_nums, cited_count=cited_count,
                bib_count=bib_count, lost_in_merge=lost_count,
                dup_ids=dup_ids, fn_refs=fn_refs, orphan_comments=com_refs)

def repair_post_merge_citations(merged_bytes,original_bytes=None):
    analysis=analyze_merge_damage(merged_bytes)
    raw=analysis['raw']
    with zipfile.ZipFile(io.BytesIO(merged_bytes)) as z:
        all_files={n:z.read(n) for n in z.namelist()}
    report=dict(analysis=analysis,steps=[],citations_before=analysis['with_data'])
    fixed=raw
    # Accept tracked changes safely
    if analysis['has_tracked']:
        def rescue_del(m):
            dc=m.group(0)
            if 'ADDIN EN.CITE' in dc:
                runs=re.findall(r'<w:r[^>]*>.*?</w:r>',dc,re.DOTALL)
                cite_runs=[r.replace('<w:delText','<w:t').replace('</w:delText>','</w:t>')
                           for r in runs if any(x in r for x in ['fldChar','instrText','fldData','ADDIN'])]
                if cite_runs: return ''.join(cite_runs)
            return ''
        fixed=re.compile(r'<w:del\b[^>]*>[\s\S]*?</w:del>',re.DOTALL).sub(rescue_del,fixed)
        def accept_ins(m):
            inner=re.sub(r'^<w:ins[^>]*>','',m.group(0)); inner=re.sub(r'</w:ins>$','',inner)
            return inner
        fixed=re.compile(r'<w:ins\b[^>]*>[\s\S]*?</w:ins>',re.DOTALL).sub(accept_ins,fixed)
        report['steps'].append('track_changes_accepted')
    # Restore broken fields
    if analysis['empty_cite']>0:
        fixed,n_fixed=fix_broken_fields(fixed)
        if n_fixed>0: report['steps'].append(f'restored_{n_fixed}_fields')
    # Compare against original
    if original_bytes:
        with zipfile.ZipFile(io.BytesIO(original_bytes)) as z:
            orig_raw=z.read('word/document.xml').decode('utf-8')
        orig_rns=set(re.findall(r'&lt;RecNum&gt;(\d+)&lt;/RecNum&gt;',orig_raw))
        merged_rns=set(re.findall(r'&lt;RecNum&gt;(\d+)&lt;/RecNum&gt;',fixed))
        lost=orig_rns-merged_rns
        report['lost_rec_nums']=list(lost)
        if lost: report['steps'].append(f'{len(lost)}_citations_lost')
    report['citations_after']=len(re.findall(r'&lt;EndNote&gt;',fixed))
    fixed=fixed.replace('\x00','')
    all_files['word/document.xml']=fixed.encode('utf-8')
    buf=io.BytesIO()
    with zipfile.ZipFile(buf,'w',zipfile.ZIP_DEFLATED) as zout:
        for n,d in all_files.items(): zout.writestr(n,d)
    buf.seek(0)
    return buf.read(),report

# ─────────────────────────────────────────────────────────────────────────────
# APP 5 LOGIC
# ─────────────────────────────────────────────────────────────────────────────
# APP 4 LOGIC — SAFE DOCUMENT MERGE
# ─────────────────────────────────────────────────────────────────────────────

def _para_sig(text):
    """Normalised fingerprint for fuzzy paragraph matching."""
    return re.sub(r'[^a-z0-9]', '', text.lower())[:120]


def _para_has_citations(para, W):
    """Return True if a paragraph has any citation content."""
    # EndNote field codes
    for f in para._p.findall(f'.//{{{W}}}instrText'):
        if 'EN.CITE' in (f.text or ''):
            return True
    # Superscript number runs
    for run in para.runs:
        rpr = run._r.find(f'{{{W}}}rPr')
        if rpr is None: continue
        va = rpr.find(f'{{{W}}}vertAlign')
        if va is not None and va.get(f'{{{W}}}val') == 'superscript':
            if any(c.isdigit() for c in run.text):
                return True
    return False


def _extract_cite_runs(para, W):
    """
    Extract citation-related XML elements from a paragraph:
    field begin/instrText/separate/end sequences and superscript runs.
    Returns list of lxml elements.
    """
    cite_elems = []
    p_elem = para._p
    # Walk runs looking for field chars and superscript runs
    in_field = False
    for child in p_elem:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            # Check for fldChar begin
            for fc in child.findall(f'.//{{{W}}}fldChar'):
                ft = fc.get(f'{{{W}}}fldCharType', '')
                if ft == 'begin':
                    in_field = True
                elif ft == 'end':
                    in_field = False
            if in_field:
                cite_elems.append(child)
            else:
                # Check for superscript
                rpr = child.find(f'{{{W}}}rPr')
                if rpr is not None:
                    va = rpr.find(f'{{{W}}}vertAlign')
                    if va is not None and va.get(f'{{{W}}}val') == 'superscript':
                        cite_elems.append(child)
        elif tag in ('bookmarkStart', 'bookmarkEnd'):
            pass  # skip
    return cite_elems


def safe_merge_documents(new_bytes, old_bytes):
    """
    Safely merge a citation-damaged new document with a citation-intact old document.

    1. Fixes any broken citation fields in the old document first (fldData recovery)
    2. Builds a paragraph signature index of the old document
    3. For each paragraph in the new document that lacks citations:
       - Finds the matching paragraph in the old document
       - Copies all citation field code runs and superscript runs across
    4. New-only paragraphs (no match in old) are kept unchanged
    Returns (merged_bytes, report_dict)
    """
    from copy import deepcopy

    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # Step 1: Fix broken fields in old doc before using it as source
    with zipfile.ZipFile(io.BytesIO(old_bytes)) as z:
        old_raw   = z.read('word/document.xml').decode('utf-8')
        old_files = {n: z.read(n) for n in z.namelist()}
    old_fixed, n_fixed = fix_broken_fields(old_raw)
    old_fixed = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', old_fixed)
    old_files['word/document.xml'] = old_fixed.encode('utf-8')
    old_buf = io.BytesIO()
    with zipfile.ZipFile(old_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
        for n, d in old_files.items(): zout.writestr(n, d)
    old_buf.seek(0)
    doc_old = Document(old_buf)
    doc_new = Document(io.BytesIO(new_bytes))

    def _has_cites(para):
        for f in para._p.findall(f'.//{{{W}}}instrText'):
            if 'EN.CITE' in (f.text or ''): return True
        for run in para.runs:
            rpr = run._r.find(f'{{{W}}}rPr')
            if rpr is None: continue
            va = rpr.find(f'{{{W}}}vertAlign')
            if va is not None and va.get(f'{{{W}}}val') == 'superscript':
                if any(c.isdigit() for c in run.text): return True
        return False

    def _cite_runs(para):
        """Return all citation-related run elements from a paragraph."""
        runs = para._p.findall(f'{{{W}}}r')
        result = []; in_field = False
        for run in runs:
            is_cite = False
            for fc in run.findall(f'.//{{{W}}}fldChar'):
                ft = fc.get(f'{{{W}}}fldCharType', '')
                if ft == 'begin':   in_field = True;  is_cite = True
                elif ft == 'end':   is_cite = True;   in_field = False
                else:                is_cite = True
            if in_field and not is_cite: is_cite = True
            rpr = run.find(f'{{{W}}}rPr')
            if not is_cite and rpr is not None:
                va = rpr.find(f'{{{W}}}vertAlign')
                if va is not None and va.get(f'{{{W}}}val') == 'superscript':
                    t = run.find(f'{{{W}}}t')
                    if t is not None and any(c.isdigit() for c in (t.text or '')):
                        is_cite = True
            if is_cite: result.append(run)
        return result

    # Build old paragraph index
    old_sig_map = {}
    for pi, para in enumerate(doc_old.paragraphs):
        text = para.text.strip()
        if len(text) < 15: continue
        sig = re.sub(r'[^a-z0-9]', '', text.lower())[:80]
        old_sig_map.setdefault(sig, []).append((pi, para))

    report = dict(total_new_paras=len(doc_new.paragraphs),
                  matched=0, citations_restored=0,
                  already_had_cites=0, unmatched=0,
                  old_fields_fixed=n_fixed, details=[])

    for pi, para_new in enumerate(doc_new.paragraphs):
        text = para_new.text.strip()
        if len(text) < 15: continue

        sig = re.sub(r'[^a-z0-9]', '', text.lower())[:80]
        candidates = old_sig_map.get(sig, [])

        # Partial match fallback (first 50 chars)
        if not candidates and len(sig) > 40:
            for osig, olist in old_sig_map.items():
                if sig[:50] == osig[:50]:
                    candidates = olist; break

        if not candidates:
            report['unmatched'] += 1; continue

        report['matched'] += 1
        _, para_old = candidates[0]

        new_has = _has_cites(para_new)
        old_has = _has_cites(para_old)

        if new_has:
            report['already_had_cites'] += 1; continue
        if not old_has:
            continue

        # Copy citation runs from old → new
        cite_elems = _cite_runs(para_old)
        if not cite_elems: continue

        new_runs = para_new._p.findall(f'{{{W}}}r')
        if not new_runs: continue

        last = new_runs[-1]
        for elem in cite_elems:
            new_elem = deepcopy(elem)
            last.addnext(new_elem)
            last = new_elem

        report['citations_restored'] += 1
        report['details'].append({
            'para_idx':     pi,
            'text_preview': text[:80],
            'cites_added':  len(cite_elems),
        })

    # Fix garbled figure numbers: merge sometimes concatenates chapter numbers
    # e.g. "Fig. 33" + "34" across adjacent runs creates "Fig. 3334.X"
    # Fix by removing the extra chapter number from the second run
    fig_fixes = 0
    for para in doc_new.paragraphs:
        runs = list(para.runs)
        i = 0
        while i < len(runs) - 1:
            t1 = runs[i].text or ''
            t2 = runs[i+1].text or ''
            if t1.endswith('33') and t2.startswith('34'):
                new_t2 = t2[2:]
                t_elem = runs[i+1]._r.find(f'{{{W}}}t')
                if t_elem is not None:
                    t_elem.text = new_t2
                    if not new_t2:
                        runs[i+1]._r.getparent().remove(runs[i+1]._r)
                    fig_fixes += 1
            i += 1
    if fig_fixes:
        report['fig_number_fixes'] = fig_fixes

    # Remove orphaned plain superscripts left by the merge
    buf_check = io.BytesIO(); doc_new.save(buf_check); buf_check.seek(0)
    _, orphans_removed = remove_orphan_superscripts(buf_check.read())
    if orphans_removed:
        # Re-load to apply removal
        buf_check.seek(0)
        clean_bytes, _ = remove_orphan_superscripts(buf_check.read())
        doc_new = Document(io.BytesIO(clean_bytes))
        report['orphan_superscripts_removed'] = orphans_removed

    buf = io.BytesIO()
    doc_new.save(buf)
    buf.seek(0)

    # Post-save: fix duplicate w:id attributes caused by copying paragraph XML
    import zipfile as _zf
    raw_buf = buf.read()
    with _zf.ZipFile(io.BytesIO(raw_buf)) as _z:
        _doc_xml  = _z.read('word/document.xml').decode('utf-8')
        _all_files = {n: _z.read(n) for n in _z.namelist()}
    _all_ids = re.findall(r'\bw:id="(\d+)"', _doc_xml)
    if _all_ids:
        _max_id  = max(int(x) for x in _all_ids)
        _next    = [_max_id + 1]; _seen = set()
        def _fix_id2(m):
            v = m.group(2)
            if v in _seen:
                nw = str(_next[0]); _next[0] += 1
                return f'{m.group(1)}{nw}{m.group(3)}'
            _seen.add(v); return m.group(0)
        _doc_xml = re.sub(r'(w:id=")(\d+)(")', _fix_id2, _doc_xml)
        _all_files['word/document.xml'] = _doc_xml.encode('utf-8')
        _out = io.BytesIO()
        with _zf.ZipFile(_out, 'w', _zf.ZIP_DEFLATED) as _zout:
            for _n, _d in _all_files.items(): _zout.writestr(_n, _d)
        _out.seek(0)
        raw_buf = _out.read()

    return raw_buf, report

# APP 5 — CITATION RENUMBERING
# ─────────────────────────────────────────────────────────────────────────────
def _apply_superscript_mapping(doc, mapping):
    """Apply old->new number mapping to all superscript runs in a document."""
    for para in doc.paragraphs:
        for run in para.runs:
            rpr = run._r.find(qn('w:rPr'))
            if rpr is None: continue
            va = rpr.find(qn('w:vertAlign'))
            if va is None or va.get(qn('w:val')) != 'superscript': continue
            text  = run.text
            sep   = ';' if ';' in text else ','
            parts = re.split(r'[,;]', text)
            new_parts = []
            for p in parts:
                ps = p.strip()
                if ps.isdigit():
                    new_parts.append(str(mapping.get(int(ps), int(ps))))
                else:
                    new_parts.append(p)
            run.text = sep.join(new_parts)


def _renumber_bib_paras(doc, bib, mapping, bib_sorted=None):
    """
    Update leading numbers in bibliography paragraphs.
    If bib_sorted is provided, also physically re-sorts them in that order.
    """
    W     = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body  = doc.element.body
    all_p = list(body.findall(f'{{{W}}}p'))

    # Update leading number text
    for pi, old_num, _ in bib:
        new_num = mapping.get(old_num, old_num)
        para    = doc.paragraphs[pi]
        for run in para.runs:
            if f"{old_num}." in run.text:
                run.text = run.text.replace(f"{old_num}.", f"{new_num}.", 1)
                break

    # Re-sort paragraphs if sort order given
    if bib_sorted and len(bib) > 1:
        anchor = all_p[bib[0][0]]
        for entry in bib_sorted:
            elem = all_p[entry[0]]
            anchor.addprevious(elem)


def renumber_citations_alpha(docx_bytes):
    """
    Renumbers citations alphabetically by first author last name (A=1, B=2...).
    Parses the bibliography, sorts A-Z, assigns new numbers, updates
    inline superscripts, and re-sorts the bibliography.
    Returns (fixed_bytes, {old_num: new_num})
    """
    doc     = Document(io.BytesIO(docx_bytes))
    ref_pat = re.compile(r'^\s*(\d+)\.\s+(.+)')
    bib     = [(pi, int(m.group(1)), para.text.strip())
               for pi, para in enumerate(doc.paragraphs)
               if (m := ref_pat.match(para.text.strip()))]

    if not bib:
        return docx_bytes, {}

    def sort_key(text):
        text       = re.sub(r'^\d+\.\s+', '', text).strip()
        first_auth = re.split(r';|,\s+[A-Z]', text)[0].strip()
        last_name  = first_auth.split(',')[0].strip().split()[-1] if first_auth else text[:20]
        return last_name.lower()

    bib_sorted = sorted(bib, key=lambda x: sort_key(x[2]))
    mapping    = {entry[1]: idx + 1 for idx, entry in enumerate(bib_sorted)}

    _apply_superscript_mapping(doc, mapping)
    _renumber_bib_paras(doc, bib, mapping, bib_sorted=bib_sorted)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read(), mapping


def renumber_citations_appearance(docx_bytes):
    """
    Renumbers citations by order of first appearance in the text (1, 2, 3...).
    Scans inline superscripts left-to-right, assigns new sequential numbers,
    updates all superscripts, and re-orders the bibliography to match.
    Returns (fixed_bytes, {old_num: new_num})
    """
    doc     = Document(io.BytesIO(docx_bytes))
    ref_pat = re.compile(r'^\s*(\d+)\.\s+(.+)')

    # Build appearance-order mapping by scanning superscripts
    seen    = {}
    for para in doc.paragraphs:
        for run in para.runs:
            rpr = run._r.find(qn('w:rPr'))
            if rpr is None: continue
            va  = rpr.find(qn('w:vertAlign'))
            if va is None or va.get(qn('w:val')) != 'superscript': continue
            for ps in re.split(r'[,;]', run.text.strip()):
                if ps.strip().isdigit():
                    num = int(ps.strip())
                    if num not in seen: seen[num] = len(seen) + 1

    if not seen:
        return docx_bytes, {}

    mapping = seen  # old_num -> new_num

    _apply_superscript_mapping(doc, mapping)

    # Re-order bibliography to match appearance order
    bib = [(pi, int(m.group(1)), para.text.strip())
           for pi, para in enumerate(doc.paragraphs)
           if (m := ref_pat.match(para.text.strip()))]

    if bib:
        # Sort bib paragraphs by their new number (appearance order)
        bib_by_new = sorted(bib, key=lambda x: mapping.get(x[1], x[1]))
        _renumber_bib_paras(doc, bib, mapping, bib_sorted=bib_by_new)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read(), mapping

# ─────────────────────────────────────────────────────────────────────────────
# APP 6 LOGIC — FIGURE INVENTORY
# ─────────────────────────────────────────────────────────────────────────────
def scan_figures(docx_bytes):
    doc=Document(io.BytesIO(docx_bytes)); items=[]
    cap_pats=[(re.compile(r'^\s*Fig(?:ure)?\.?\s*(\d+[\w\-\.]*)',re.IGNORECASE),'Figure'),
              (re.compile(r'^\s*Table\s*(\d+[\w\-\.]*)',re.IGNORECASE),'Table'),
              (re.compile(r'^\s*Box\s*(\d+[\w\-\.]*)',re.IGNORECASE),'Box'),
              (re.compile(r'^\s*Plate\s*(\d+[\w\-\.]*)',re.IGNORECASE),'Plate'),
              (re.compile(r'^\s*Video\s*(\d+[\w\-\.]*)',re.IGNORECASE),'Video'),
              (re.compile(r'^\s*Appendix\s*(\d+[\w\-\.]*)',re.IGNORECASE),'Appendix')]
    DN='http://schemas.openxmlformats.org/drawingml/2006/main'
    for pi,para in enumerate(doc.paragraphs):
        text=para.text.strip()
        has_img=para._p.find(f'.//{{{DN}}}blip') is not None
        style=para.style.name if para.style else ''
        itype=None; inum=None
        for pat,pt in cap_pats:
            m=pat.match(text)
            if m: itype=pt; inum=m.group(1); break
        if itype or has_img or 'caption' in style.lower():
            items.append(dict(para_idx=pi,type=itype or ('Image' if has_img else 'Caption'),
                              number=inum or '',caption=text[:250],has_image=has_img,style=style))
    return items

def cross_ref_excel(items, excel_bytes):
    import openpyxl
    try:
        wb=openpyxl.load_workbook(io.BytesIO(excel_bytes)); ws=wb.active
        rows=list(ws.iter_rows(values_only=True))
        if not rows: return [(i,'no_data','') for i in items]
        hdrs=[str(c or '').lower().strip() for c in rows[0]]
        nc=next((i for i,h in enumerate(hdrs) if any(w in h for w in ['num','#'])),None)
        lc=next((i for i,h in enumerate(hdrs) if any(w in h for w in ['name','caption','title','new','label'])),None)
        if lc is None: return [(i,'no_label_col','') for i in items]
        lookup={}
        for row in rows[1:]:
            if not row: continue
            num=str(row[nc] or '').strip() if nc is not None else ''
            label=str(row[lc] or '').strip()
            if num and label: lookup[num.lower()]=label
        results=[]
        for item in items:
            exp=lookup.get(item['number'].lower(),'')
            if not exp: status='not_in_excel'
            elif exp.lower() in item['caption'].lower(): status='match'
            else: status='mismatch'
            results.append((item,status,exp))
        return results
    except Exception as e: return [(i,f'error','') for i in items]

# ─────────────────────────────────────────────────────────────────────────────
# APP 7 LOGIC — PUBMED SEARCH
# ─────────────────────────────────────────────────────────────────────────────
def pubmed_search_full(query,date_from='',date_to='',journal_filter='',max_results=20):
    try:
        term=query
        if date_from and date_to: term+=f' AND {date_from}:{date_to}[dp]'
        elif date_from: term+=f' AND {date_from}:3000[dp]'
        if journal_filter: term+=f' AND "{journal_filter}"[ta]'
        r=requests.get(PUBMED_ESEARCH,params={'db':'pubmed','term':term,
            'retmax':max_results,'retmode':'json','sort':'relevance'},timeout=10)
        ids=r.json().get('esearchresult',{}).get('idlist',[])
        if not ids: return []
        r2=requests.get('https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi',
            params={'db':'pubmed','id':','.join(ids),'retmode':'xml','rettype':'abstract'},timeout=15)
        root=ET.fromstring(r2.content); results=[]
        for art in root.findall('.//PubmedArticle'):
            pmid=art.findtext('.//PMID') or ''
            title=art.findtext('.//ArticleTitle') or ''
            abstract=' '.join(t.text or '' for t in art.findall('.//AbstractText'))
            journal=art.findtext('.//Journal/Title') or art.findtext('.//ISOAbbreviation') or ''
            year=art.findtext('.//PubDate/Year') or (art.findtext('.//PubDate/MedlineDate') or '')[:4]
            volume=art.findtext('.//Volume') or ''; issue=art.findtext('.//Issue') or ''
            pages=art.findtext('.//MedlinePgn') or ''
            doi=next((a.text for a in art.findall('.//ArticleId') if a.get('IdType')=='doi'),'')
            pmc=next((a.text for a in art.findall('.//ArticleId') if a.get('IdType')=='pmc'),'')
            authors=[]
            for auth in art.findall('.//Author')[:6]:
                last=auth.findtext('LastName') or ''; fore=auth.findtext('Initials') or ''
                if last: authors.append(f"{last} {fore}".strip())
            results.append(dict(pmid=pmid,title=title,authors=authors,journal=journal,
                year=year,volume=volume,issue=issue,pages=pages,
                abstract=abstract[:600],doi=doi,
                pubmed_url=f'https://pubmed.ncbi.nlm.nih.gov/{pmid}/',
                pmc_url=f'https://www.ncbi.nlm.nih.gov/pmc/articles/{pmc}/' if pmc else '',
                doi_url=f'https://doi.org/{doi}' if doi else ''))
        return results
    except: return []

def results_to_xml(results):
    import html as hm
    xml='<?xml version="1.0" encoding="UTF-8"?>\n<xml>\n  <records>\n'
    for i,r in enumerate(results,1):
        xml+=f'    <record>\n      <rec-number>{i}</rec-number>\n      <ref-type name="Journal Article">17</ref-type>\n'
        if r['authors']:
            xml+='      <contributors>\n        <authors>\n'
            for a in r['authors']: xml+=f'          <author>{hm.escape(a)}</author>\n'
            xml+='        </authors>\n      </contributors>\n'
        xml+=f'      <titles>\n        <title>{hm.escape(r["title"])}</title>\n        <secondary-title>{hm.escape(r["journal"])}</secondary-title>\n      </titles>\n'
        xml+=f'      <dates><year>{hm.escape(r["year"])}</year></dates>\n'
        for k,t in [('volume','volume'),('issue','number'),('pages','pages')]:
            if r[k]: xml+=f'      <{t}>{hm.escape(r[k])}</{t}>\n'
        if r['abstract']: xml+=f'      <abstract>{hm.escape(r["abstract"])}</abstract>\n'
        if r['doi']: xml+=f'      <electronic-resource-number>{hm.escape(r["doi"])}</electronic-resource-number>\n'
        xml+=f'      <urls><related-urls><url>{hm.escape(r["pubmed_url"])}</url></related-urls></urls>\n'
        xml+='    </record>\n'
    xml+='  </records>\n</xml>\n'; return xml

# ─────────────────────────────────────────────────────────────────────────────
# APP 8 LOGIC — BATCH RENAME
# ─────────────────────────────────────────────────────────────────────────────
def load_rename_pairs(excel_bytes):
    import openpyxl
    try:
        wb=openpyxl.load_workbook(io.BytesIO(excel_bytes)); ws=wb.active
        rows=list(ws.iter_rows(values_only=True))
        if not rows: return []
        hdrs=[str(c or '').lower().strip() for c in rows[0]]
        oc=next((i for i,h in enumerate(hdrs) if any(w in h for w in ['old','current','find','original','from'])),0)
        nc=next((i for i,h in enumerate(hdrs) if any(w in h for w in ['new','replace','final','to','updated'])),1)
        pairs=[]
        for row in rows[1:]:
            if not row or len(row)<=max(oc,nc): continue
            old=str(row[oc] or '').strip(); new=str(row[nc] or '').strip()
            if old and new and old!=new: pairs.append((old,new))
        return pairs
    except: return []

def batch_rename(docx_bytes,pairs,match_case=False,whole_word=False):
    import html as hm
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        raw=z.read('word/document.xml').decode('utf-8')
        all_files={n:z.read(n) for n in z.namelist()}
    report=[]; fixed=raw
    for old,new in pairs:
        oe=hm.escape(old); ne=hm.escape(new)
        flags=0 if match_case else re.IGNORECASE
        pat=re.escape(oe)
        if whole_word: pat=r'(?<![a-zA-Z])'+pat+r'(?![a-zA-Z])'
        count=len(re.findall(pat,fixed,flags=flags))
        if count>0: fixed=re.sub(pat,ne,fixed,flags=flags); report.append({'old':old,'new':new,'count':count,'status':'replaced'})
        else: report.append({'old':old,'new':new,'count':0,'status':'not_found'})
    all_files['word/document.xml']=fixed.encode('utf-8')
    buf=io.BytesIO()
    with zipfile.ZipFile(buf,'w',zipfile.ZIP_DEFLATED) as zout:
        for n,d in all_files.items(): zout.writestr(n,d)
    buf.seek(0); return buf.read(),report

# ─────────────────────────────────────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────────────────────────────────────
defaults=dict(flagged=[],current_idx=0,decisions=[],doc_obj=None,repair_done=False,
              refs=[],vec=None,mat=None,
              fix_stage=1,fix_analysis=None,fix_raw_xml=None,fix_docx_bytes=None,
              fix_after_stage1=None,fix_after_stage2=None,fix_karol_db_id=None,
              fix_karol_rec_nums={},fix_missing_refs=[],fix_doc_name="document.docx",
              comp_result=None,comp_usage={},comp_labels=("",""),comp_refs=([],[]))
for k,v in defaults.items():
    if k not in st.session_state: st.session_state[k]=v

# ─────────────────────────────────────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown('''
    <div style="padding:1.4rem 0.5rem 0.6rem;">
      <div class="sidebar-logo"><em>Citation</em> Toolkit</div>
      <div class="sidebar-sub">Tachdjian's Pediatric Orthopaedics</div>
    </div>
    <div style="height:1px;background:var(--border);margin-bottom:1rem;"></div>
    <div class="sidebar-rule">Search & Import</div>
    ''', unsafe_allow_html=True)

    tool = st.radio("", [
        "01 · PubMed Search",
    ], label_visibility="collapsed", key="_nav_s1")

    st.markdown('<div class="sidebar-rule">Citation Repair</div>', unsafe_allow_html=True)
    tool2 = st.radio("", [
        "02 · Broken Citation Fixer",
        "03 · Document Merger",
        "04 · Citation Repair",
    ], label_visibility="collapsed", key="_nav_s2")

    st.markdown('<div class="sidebar-rule">Audit & Compare</div>', unsafe_allow_html=True)
    tool3 = st.radio("", [
        "05 · Bibliography Auditor",
        "06 · Reference Comparator",
        "11 · RecNum Inspector",
    ], label_visibility="collapsed", key="_nav_s3")

    st.markdown('<div class="sidebar-rule">Finalise</div>', unsafe_allow_html=True)
    tool4 = st.radio("", [
        "07 · Citation Renumbering",
        "08 · Figure Inventory",
        "09 · Document Health Check",
        "10 · Batch Rename",
    ], label_visibility="collapsed", key="_nav_s4")

    # Resolve which tool is active from the four groups
    _sel_map = {
        "01 · PubMed Search":         "App 1 — PubMed Search",
        "02 · Broken Citation Fixer":  "App 2 — Broken Citation Fixer",
        "03 · Document Merger":        "App 3 — Document Merger",
        "04 · Citation Repair":        "App 4 — Citation Repair",
        "05 · Bibliography Auditor":   "App 5 — Bibliography Auditor",
        "06 · Reference Comparator":   "App 6 — Reference Comparator",
        "07 · Citation Renumbering":   "App 7 — Citation Renumbering",
        "08 · Figure Inventory":       "App 8 — Figure Inventory",
        "09 · Document Health Check":  "App 9 — Document Health Check",
        "10 · Batch Rename":           "App 10 — Batch Rename",
        "11 · RecNum Inspector":       "App 11 — RecNum Inspector",
    }
    # Last-clicked radio determines active tool
    if "_last_tool" not in st.session_state:
        st.session_state["_last_tool"] = "02 · Broken Citation Fixer"
    for _r, _k in [(tool,"_nav_s1"),(tool2,"_nav_s2"),(tool3,"_nav_s3"),(tool4,"_nav_s4")]:
        if _r and st.session_state.get(_k) and _r != st.session_state.get("_prev_"+_k):
            st.session_state["_last_tool"] = _r
            st.session_state["_prev_"+_k] = _r
    tool = _sel_map.get(st.session_state.get("_last_tool","02 · Broken Citation Fixer"),
                        "App 2 — Broken Citation Fixer")

    st.markdown('<div style="height:1.5rem"></div>', unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# APP 2 UI — BROKEN CITATION FIXER
# ─────────────────────────────────────────────────────────────────────────────
if tool == "App 2 — Broken Citation Fixer":
    st.markdown('<div class="app-label">A practical tool &nbsp;02</div>', unsafe_allow_html=True)
    st.markdown("## Broken Citation Fixer")
    st.markdown('<div class="instruction-box">Use this when EndNote only recognizes some citations even though the bibliography shows all references — or when citations show as "Traveling Library" instead of your EndNote library.</div>', unsafe_allow_html=True)

    # Step 0 — Extract traveling library
    with st.expander("📥 Step 0 — Extract traveling library references (start here)", expanded=True):
        st.markdown('<div class="step-desc">If EndNote cannot find your refs, extract them from the Word file and import into your library.</div>', unsafe_allow_html=True)
        tl_file=st.file_uploader("Word document",type=["docx"],key="tl_doc")
        if tl_file:
            if st.button("Extract references",type="primary",key="tl_run"):
                with st.spinner("Extracting..."):
                    try:
                        tl_xml,tl_count=extract_traveling_library_xml(tl_file.read())
                        st.success(f"Extracted {tl_count} references.")
                        st.download_button(f"⬇ Download EndNote XML ({tl_count} refs)",
                            data=tl_xml.encode('utf-8'),
                            file_name=Path(tl_file.name).stem+"_traveling_library.xml",
                            mime="application/xml",type="primary")
                        st.markdown('<div class="instruction-box">Import into your library: <b>File → Import → File</b> → select XML → Import Option: <code>EndNote XML</code> → Duplicates: <code>Discard Duplicates</code> → Import</div>',unsafe_allow_html=True)
                    except Exception as e: st.error(f"Error: {e}")

    st.divider()

    # Step 0b — Remap traveling library citations
    with st.expander("🔄 Step 0b — Remap citations still showing as Traveling Library"):
        st.markdown('<div class="step-desc">After importing the traveling library, some citations may still show as "Traveling Library" because their RecNums don\'t match your library\'s record IDs. This remaps them by author+year+title matching.</div>',unsafe_allow_html=True)
        col1,col2=st.columns(2)
        with col1: remap_doc=st.file_uploader("Word document",type=["docx"],key="remap_doc")
        with col2: remap_enl=st.file_uploader("EndNote library (.enl)",type=["enl"],key="remap_enl")
        if remap_doc and remap_enl:
            if st.button("Find and remap",type="primary",key="remap_run"):
                with st.spinner("Comparing against your library..."):
                    try:
                        fixed_bytes,report=remap_traveling_citations(remap_doc.read(),remap_enl.read())
                        remapped=[r for r in report if r['status']=='remapped']
                        not_found=[r for r in report if r['status']=='not_found']
                        col1,col2,col3=st.columns(3)
                        col1.metric("Checked",len(report)); col2.metric("Remapped",len(remapped)); col3.metric("Not found",len(not_found))
                        if remapped:
                            st.success(f"Remapped {len(remapped)} citations.")
                            st.download_button("⬇ Download remapped document",data=fixed_bytes,
                                file_name=Path(remap_doc.name).stem+"_remapped.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",type="primary")
                        if not_found:
                            st.warning(f"{len(not_found)} citations not found in your library — add them manually.")
                            with st.expander("Unmatched citations"):
                                for r in not_found:
                                    st.markdown(f'<div class="ref-item error"><b>RecNum {r["old_rec_num"]}</b> — {r["author"]} ({r["year"]}) {r["title"]}</div>',unsafe_allow_html=True)
                    except Exception as e: st.error(f"Error: {e}")

    st.divider()

    # Stage 1
    stage1_done=st.session_state.fix_analysis is not None
    st.markdown(f'<div class="step-card {"done" if stage1_done else "active"}"><div class="step-header"><span class="step-num {"done" if stage1_done else ""}">{"✓" if stage1_done else "STEP 1"}</span><span class="step-title">Repair broken citation field codes</span></div><div class="step-desc">Recovers citation XML data from the document\'s internal backup storage.</div></div>',unsafe_allow_html=True)
    if not stage1_done:
        doc_file=st.file_uploader("Word document (.docx)",type=["docx"],key="fix_doc_upload")
        if doc_file and st.button("Analyze & repair",type="primary"):
            with st.spinner("Scanning..."):
                docx_bytes=doc_file.read()
                analysis=analyze_docx_citations(docx_bytes)
                fixed_xml,n_fixed=fix_broken_fields(analysis['raw'])
                st.session_state.fix_analysis=analysis; st.session_state.fix_docx_bytes=docx_bytes
                st.session_state.fix_after_stage1=fixed_xml; st.session_state.fix_raw_xml=fixed_xml
                st.session_state.fix_doc_name=doc_file.name
            st.rerun()
    else:
        a=st.session_state.fix_analysis

        # Key metrics
        col1,col2,col3,col4 = st.columns(4)
        col1.metric("Total citation fields", a['total_fields'])
        col2.metric("EndNote currently sees",
                    a.get('endnote_sees', a['working']),
                    help="Unique references EndNote can read right now")
        col3.metric("Hidden in broken fields",
                    a.get('endnote_misses', a['broken_empty']),
                    delta=f"-{a.get('endnote_misses', a['broken_empty'])}" if a.get('endnote_misses', a['broken_empty']) > 0 else None,
                    delta_color="inverse",
                    help="References locked in broken field codes — EndNote cannot see these, causing bibliography undercount")
        col4.metric("Recovered by this fix", a['broken_empty'])

        if a.get('endnote_misses', 0) > 0:
            bib = a.get('bib_count', 0)
            sees = a.get('endnote_sees', a['working'])
            misses = a.get('endnote_misses', a['broken_empty'])
            st.error(
                f"⚠ **EndNote is undercounting your references.** "
                f"It can currently see **{sees}** unique references, but **{misses}** more "
                f"are locked in broken citation field codes and invisible to EndNote. "
                f"This is why your bibliography shows fewer entries than expected. "
                f"Download the Stage 1 result below — this fix restores all {misses} hidden references."
            )
        elif a['broken_empty'] > 0:
            st.success(f"✓ {a['broken_empty']} broken field(s) recovered.")
        else:
            st.success("✓ All citation fields are intact — no broken fields found.")
        stage1_bytes=build_fixed_docx(st.session_state.fix_docx_bytes,st.session_state.fix_after_stage1)
        st.download_button("⬇ Download Stage 1 result",data=stage1_bytes,
            file_name=Path(st.session_state.fix_doc_name).stem+"_stage1.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Orphaned superscript removal — run on the Stage 1 output
        st.markdown("---")
        st.markdown("**Optional: Remove orphaned superscripts**")
        st.markdown(
            "After a merge, plain-text superscript citation numbers sometimes remain "
            "next to working EndNote field codes (e.g. `135.40,47`). "
            "These are invisible to EndNote but visible to the reader. Click below to strip them."
        )
        if st.button("Remove orphaned plain superscripts"):
            with st.spinner("Scanning for orphaned superscripts..."):
                cleaned_bytes, n_removed = remove_orphan_superscripts(stage1_bytes)
            if n_removed == 0:
                st.info("No orphaned superscripts found — document is clean.")
            else:
                st.success(f"✓ Removed {n_removed} orphaned superscript run(s).")
                st.download_button(
                    "⬇ Download cleaned document",
                    data=cleaned_bytes,
                    file_name=Path(st.session_state.fix_doc_name).stem+"_cleaned.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )

        if st.button("↺ Start over"):
            for k in ['fix_analysis','fix_raw_xml','fix_docx_bytes','fix_after_stage1','fix_after_stage2','fix_karol_db_id','fix_karol_rec_nums','fix_missing_refs']:
                st.session_state[k]=defaults[k]
            st.rerun()

    st.divider()

    # Stage 2
    stage2_done=st.session_state.fix_karol_db_id is not None
    st.markdown(f'<div class="step-card {"done" if stage2_done else ("active" if stage1_done else "waiting")}"><div class="step-header"><span class="step-num {"done" if stage2_done else ""}">{"✓" if stage2_done else "STEP 2"}</span><span class="step-title">Get your EndNote library fingerprint (db-id)</span></div><div class="step-desc">Export any refs from your library as XML (File → Export → XML), then upload below. The app extracts the library fingerprint automatically.</div></div>',unsafe_allow_html=True)
    if stage1_done and not stage2_done:
        col_a,col_b=st.columns(2)
        with col_a:
            xml_exp=st.file_uploader("EndNote XML export (gets db-id)",type=["xml"],key="fix_xml_export")
        with col_b:
            enl_f=st.file_uploader("EndNote library (.enl) — optional, checks for missing refs",type=["enl"],key="fix_enl")
        if xml_exp:
            db_id=extract_karol_db_id(xml_exp.read())
            if db_id:
                st.session_state.fix_karol_db_id=db_id
                st.success(f"✓ Library fingerprint found: `{db_id[:20]}...`"); st.rerun()
            else: st.error("No db-id found. Make sure you exported from EndNote as XML format.")
        if enl_f:
            with st.spinner("Reading library..."):
                rns=get_karol_rec_nums(enl_f.read())
                st.session_state.fix_karol_rec_nums=rns
                missing=check_missing_from_karol(st.session_state.fix_raw_xml or "",set(rns.keys()))
                st.session_state.fix_missing_refs=missing
            st.success(f"✓ {len(rns):,} refs in your library.")
            if missing: st.warning(f"⚠ {len(missing)} ref(s) not in your library: RecNums {', '.join(missing)}")
        with st.expander("Enter db-id manually"):
            mid=st.text_input("db-id",placeholder="e.g. s5pa559ekdxfr0esvw85...")
            if st.button("Use this db-id") and mid.strip():
                st.session_state.fix_karol_db_id=mid.strip(); st.rerun()
    elif stage2_done:
        st.success(f"✓ Library db-id: `{st.session_state.fix_karol_db_id[:20]}...`")

    st.divider()

    # Stage 3
    stage3_done=st.session_state.fix_after_stage2 is not None
    stage3_ready=stage1_done and stage2_done
    st.markdown(f'<div class="step-card {"done" if stage3_done else ("active" if stage3_ready else "waiting")}"><div class="step-header"><span class="step-num {"done" if stage3_done else ""}">{"✓" if stage3_done else "STEP 3"}</span><span class="step-title">Apply full fix and generate files</span></div></div>',unsafe_allow_html=True)
    if stage3_ready and not stage3_done:
        missing=st.session_state.fix_missing_refs
        proceed=True
        if missing:
            st.warning(f"⚠ {len(missing)} ref(s) not in your library — add them manually first (EndNote → References → New Reference).")
            for rn in missing: st.markdown(f"- RecNum **{rn}**")
            proceed=st.checkbox("I've added the missing refs (or will add them later)")
        if proceed and st.button("Apply full fix",type="primary"):
            with st.spinner("Patching db-ids..."):
                patched,_=patch_db_ids(st.session_state.fix_raw_xml,
                                       st.session_state.fix_analysis['db_ids'],
                                       st.session_state.fix_karol_db_id)
                st.session_state.fix_after_stage2=patched; st.rerun()
    elif stage3_done:
        patched=st.session_state.fix_after_stage2
        final_bytes=build_fixed_docx(st.session_state.fix_docx_bytes,patched)
        working_after=len(re.findall(r'&lt;EndNote&gt;',patched))
        st.success("✓ Full fix applied.")
        col1,col2=st.columns(2)
        col1.metric("Citations now linked",working_after)
        col_d1,col_d2=st.columns(2)
        with col_d1:
            st.download_button("⬇ Download fixed document",data=final_bytes,
                file_name=Path(st.session_state.fix_doc_name).stem+"_fully_fixed.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",type="primary")
        with col_d2:
            macro=generate_vba_macro(st.session_state.fix_doc_name)
            st.download_button("⬇ Download VBA macro",data=macro.encode(),
                file_name="RelinkEndNoteCitations.bas",mime="text/plain")
        st.markdown('<div class="instruction-box"><b>Final step in Word:</b><br>1. Open the fixed document with your EndNote library connected<br>2. Open the .bas file in Notepad, copy all<br>3. In Word: Alt+F11 → Insert → Module → paste → Alt+F8 → RelinkAllCitations → Run<br>4. EndNote tab → Update Citations and Bibliography</div>',unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# APP 1 UI — CITATION REPAIR
# ─────────────────────────────────────────────────────────────────────────────
elif tool == "App 4 — Citation Repair":
    st.markdown('<div class="app-label">A practical tool &nbsp;04</div>', unsafe_allow_html=True)
    st.markdown("## Citation Repair")
    st.markdown('<div class="step-desc">Find missing citation placeholders and match them to your EndNote library.</div>',unsafe_allow_html=True)
    col1,col2=st.columns(2)
    with col1: doc_file=st.file_uploader("Word document (.docx)",type=["docx"],key="doc_upload")
    with col2: lib_file=st.file_uploader("EndNote library (.xml)",type=["xml"],key="lib_upload")
    cfg1,cfg2=st.columns(2)
    with cfg1: mode=st.selectbox("Mode",["Interactive Review","Auto-Insert","Report Only"])
    with cfg2: use_pubmed=st.toggle("PubMed fallback",value=False)
    with st.expander("Custom citation markers"):
        custom=st.text_input("Additional markers (comma-separated)",placeholder="e.g. ??, [TBD]")
        if custom:
            extras=[re.escape(m.strip()) for m in custom.split(",") if m.strip()]
            CITATION_MARKERS=re.compile('|'.join(MISSING_PATTERNS+extras),re.IGNORECASE)
    run_btn=st.button("Scan Document",type="primary",disabled=not(doc_file and lib_file))
    if run_btn and doc_file and lib_file:
        for k in ["flagged","current_idx","decisions","doc_obj","refs","vec","mat","repair_done"]:
            st.session_state[k]=defaults[k]
        with st.spinner("Parsing..."):
            db=doc_file.read(); flagged,doc_obj=extract_flagged(db)
            refs=parse_endnote_xml_bytes(lib_file.read())
            if refs:
                corpora=tuple(r["corpus"] for r in refs); vec,mat=build_tfidf(corpora)
                st.session_state.update(dict(flagged=flagged,doc_obj=doc_obj,refs=refs,
                    vec=vec,mat=mat,current_idx=0,decisions=[],repair_done=False))
        if not refs: st.error("No references found in XML.")
        elif not flagged: st.warning("No citation markers found.")
        else: st.success(f"Found **{len(flagged)}** missing citation(s) across **{len(refs)}** refs.")
    flagged=st.session_state.flagged
    if flagged and not st.session_state.repair_done:
        st.divider()
        idx=st.session_state.current_idx; total=len(flagged); done=len(st.session_state.decisions)
        st.markdown(f'<div class="progress-outer"><div class="progress-inner" style="width:{int(done/total*100)}%"></div></div>',unsafe_allow_html=True)
        st.caption(f"{done} of {total} reviewed")
        if idx<total:
            item=flagged[idx]
            cands=match_sentence(item["sentence"],st.session_state.vec,st.session_state.mat,st.session_state.refs)
            best=cands[0]["score"] if cands else 0
            if mode=="Auto-Insert" and best>=TFIDF_THRESHOLD:
                label=author_label(cands[0]["ref"]); para=st.session_state.doc_obj.paragraphs[item["para_idx"]]
                insert_superscript(para,item["marker"],label)
                st.session_state.decisions.append({**item,"action":"accepted","ref":cands[0]["ref"],"score":best,"candidates":cands})
                st.session_state.current_idx+=1; st.rerun()
            elif mode=="Report Only":
                st.session_state.decisions.append({**item,"action":"skipped","candidates":cands})
                st.session_state.current_idx+=1; st.rerun()
            else:
                st.markdown(f'<div class="match-card"><span class="match-marker">{item["marker"]}</span><div class="match-sentence">"{item["sentence"][:280]}"</div><div style="font-size:0.72rem;color:#3a4a5a">Para {item["para_idx"]+1}</div></div>',unsafe_allow_html=True)
                st.markdown('<div class="section-label">Top Matches</div>',unsafe_allow_html=True)
                chosen_idx=None
                for j,c in enumerate(cands):
                    ca,cb=st.columns([1,8])
                    with ca:
                        if st.button("Use",key=f"pick_{idx}_{j}"): chosen_idx=j
                    with cb:
                        sc=score_class(c["score"])
                        st.markdown(f'<div style="display:flex;align-items:center;gap:8px;padding:5px 0"><span class="score-pill {sc}">{c["score"]:.3f}</span><span style="font-size:0.83rem;color:#c8d0db">{fmt_ref(c["ref"])}</span></div>',unsafe_allow_html=True)
                    if chosen_idx==j:
                        label=author_label(cands[chosen_idx]["ref"])
                        para=st.session_state.doc_obj.paragraphs[item["para_idx"]]
                        insert_superscript(para,item["marker"],label)
                        st.session_state.decisions.append({**item,"action":"accepted","ref":cands[chosen_idx]["ref"],"score":cands[chosen_idx]["score"],"candidates":cands})
                        st.session_state.current_idx+=1; st.rerun()
                ba,bb=st.columns(2)
                with ba:
                    if st.button("Skip",key=f"skip_{idx}"):
                        st.session_state.decisions.append({**item,"action":"skipped","candidates":cands})
                        st.session_state.current_idx+=1; st.rerun()
                with bb:
                    if done>0 and st.button("Finish",key=f"fin_{idx}"):
                        st.session_state.repair_done=True; st.rerun()
        else:
            st.session_state.repair_done=True; st.rerun()
    if st.session_state.repair_done and st.session_state.decisions:
        st.divider()
        decisions=st.session_state.decisions
        accepted=sum(1 for d in decisions if d["action"]=="accepted")
        skipped=sum(1 for d in decisions if d["action"]=="skipped")
        col1,col2,col3=st.columns(3)
        col1.metric("Total",len(decisions)); col2.metric("Accepted",accepted); col3.metric("Skipped",skipped)
        cd1,cd2=st.columns(2)
        with cd1:
            st.download_button("⬇ Repaired document",data=doc_to_bytes(st.session_state.doc_obj),
                file_name="manuscript_repaired.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",type="primary")
        with cd2:
            st.download_button("⬇ Decision report",data=write_repair_report(decisions),
                file_name="citation_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ─────────────────────────────────────────────────────────────────────────────
# APP 3 UI — REFERENCE COMPARATOR
# ─────────────────────────────────────────────────────────────────────────────
elif tool == "App 6 — Reference Comparator":
    st.markdown("## Reference List Comparator")
    col1,col2=st.columns(2)
    with col1:
        st.markdown("**List A**")
        file_a=st.file_uploader("List A",type=["xml","docx","txt"],key="comp_a",label_visibility="collapsed")
    with col2:
        st.markdown("**List B**")
        file_b=st.file_uploader("List B",type=["xml","docx","txt"],key="comp_b",label_visibility="collapsed")
    st.markdown("**Manuscript (optional)**")
    ms_file=st.file_uploader("Manuscript",type=["docx"],key="comp_ms",label_visibility="collapsed")
    if st.button("Compare Lists",type="primary",disabled=not(file_a and file_b)):
        with st.spinner("Comparing..."):
            refs_a,label_a=load_ref_file(file_a); refs_b,label_b=load_ref_file(file_b)
            if refs_a and refs_b:
                all_c=[r["corpus"] for r in refs_a]+[r["corpus"] for r in refs_b]
                vec=TfidfVectorizer(ngram_range=(1,2),sublinear_tf=True,max_features=50000)
                vec.fit(all_c)
                emb_a=vec.transform([r["corpus"] for r in refs_a])
                emb_b=vec.transform([r["corpus"] for r in refs_b])
                matrix=cosine_similarity(emb_a,emb_b)
                matched,only_a,fuzzy,matched_b=[],[],[],set()
                for i,ra in enumerate(refs_a):
                    bj=int(matrix[i].argmax()); bs=float(matrix[i][bj])
                    if bs>=MATCH_THRESHOLD: matched.append((ra,refs_b[bj],bs)); matched_b.add(bj)
                    elif bs>=FUZZY_THRESHOLD: fuzzy.append((ra,refs_b[bj],bs))
                    else: only_a.append(ra)
                only_b=[refs_b[j] for j in range(len(refs_b)) if j not in matched_b]
                result=dict(matched=matched,only_in_a=only_a,only_in_b=only_b,fuzzy=fuzzy)
                usage={}
                if ms_file:
                    ms_doc=Document(io.BytesIO(ms_file.read()))
                    ms_paras=[p.text for p in ms_doc.paragraphs if p.text.strip()]
                    for ref in only_a+only_b+[x[0] for x in fuzzy]:
                        words=[w for w in ref["title"].split() if len(w)>5][:5]
                        found=[p[:200] for p in ms_paras if sum(1 for w in words if w.lower() in p.lower())>=min(3,len(words))]
                        usage[ref.get("id","")]=found
                st.session_state.update(dict(comp_result=result,comp_usage=usage,
                    comp_labels=(label_a,label_b),comp_refs=(refs_a,refs_b)))
    if st.session_state.comp_result:
        result=st.session_state.comp_result; usage=st.session_state.comp_usage
        label_a,label_b=st.session_state.comp_labels
        st.divider()
        col1,col2,col3,col4=st.columns(4)
        col1.metric("Matched",len(result["matched"])); col2.metric("Only in A",len(result["only_in_a"]))
        col3.metric("Only in B",len(result["only_in_b"])); col4.metric("Review",len(result["fuzzy"]))
        tabs=st.tabs([f"Only in A ({len(result['only_in_a'])})",f"Only in B ({len(result['only_in_b'])})",
                      f"Review ({len(result['fuzzy'])})",f"Matched ({len(result['matched'])})"])
        def ref_card(ref,color,locs=None):
            loc=""
            if locs: loc=f'<div style="font-size:0.75rem;color:#4fc3f7;margin-top:4px">Found in {len(locs)} paragraph(s)</div>'+"".join(f'<div style="font-size:0.74rem;color:#607080;font-style:italic">{l[:150]}...</div>' for l in locs[:2])
            st.markdown(f'<div class="ref-item {color}">{fmt_ref(ref)}{loc}</div>',unsafe_allow_html=True)
        with tabs[0]:
            if result["only_in_a"]:
                st.caption(f"In **{label_a}** but not **{label_b}**")
                for ref in result["only_in_a"]: ref_card(ref,"missing",usage.get(ref.get("id",""),[]))
            else: st.success("None")
        with tabs[1]:
            if result["only_in_b"]:
                st.caption(f"In **{label_b}** but not **{label_a}**")
                for ref in result["only_in_b"]: ref_card(ref,"warn",usage.get(ref.get("id",""),[]))
            else: st.success("None")
        with tabs[2]:
            for ra,rb,score in result["fuzzy"]:
                st.markdown(f'<div class="ref-item warning"><b>[{score:.3f}]</b> A:{fmt_ref(ra,True)}<br>B:{fmt_ref(rb,True)}</div>',unsafe_allow_html=True)
        with tabs[3]:
            for ra,rb,score in result["matched"]:
                st.markdown(f'<div class="ref-item ok">[{score:.3f}] {fmt_ref(ra,True)}</div>',unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# APP 4 UI — DOCUMENT MERGER
# ─────────────────────────────────────────────────────────────────────────────
elif tool == "App 3 — Document Merger":
    st.markdown('<div class="app-label">A practical tool &nbsp;03</div>', unsafe_allow_html=True)
    st.markdown("## Document Merger")

    merge_mode = st.radio(
        "What do you need to do?",
        [
            "Restore lost citations — merge citation-intact OLD doc with text-edited NEW doc",
            "Repair already-merged document — fix broken field codes after a merge",
        ]
    )
    st.divider()

    if "Restore lost citations" in merge_mode:
        st.markdown('''<div class="instruction-box">
        <b>Use this when:</b> A merge broke or removed inline citations from the new document,
        but you still have the older version with all citations intact.<br><br>
        The app matches paragraphs between the two documents by text content and copies
        citation field codes from the old document into the new one wherever they are missing.
        All text edits in the new document are preserved.
        </div>''', unsafe_allow_html=True)

        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**New document** — has text edits but missing/broken citations")
            new_file = st.file_uploader("New .docx", type=["docx"], key="merge_new")
        with col2:
            st.markdown("**Old document** — has all citations intact")
            old_file = st.file_uploader("Old .docx", type=["docx"], key="merge_old")

        if new_file and old_file:
            new_bytes_m = new_file.read()
            old_bytes_m = old_file.read()
            new_fname   = new_file.name

            new_analysis = analyze_merge_damage(new_bytes_m)
            col1, col2, col3 = st.columns(3)
            col1.metric("Bibliography entries",  new_analysis.get("bib_count", 0))
            col2.metric("Cited in text (new)",   new_analysis.get("cited_count", 0))
            col3.metric("Citations lost",         new_analysis.get("lost_in_merge", 0),
                        delta=f'-{new_analysis.get("lost_in_merge",0)}' if new_analysis.get("lost_in_merge",0) else None,
                        delta_color="inverse")

            if st.button("Restore citations from old document", type="primary"):
                with st.spinner("Matching paragraphs and restoring citations..."):
                    try:
                        merged_bytes, rpt = safe_merge_documents(new_bytes_m, old_bytes_m)
                        st.success(
                            f"Done. Matched {rpt['matched']} paragraphs. "
                            f"Restored citations in {rpt['citations_restored']} paragraph(s). "
                            f"{rpt['unmatched']} paragraph(s) could not be matched to the old document."
                        )
                        col1,col2,col3,col4 = st.columns(4)
                        col1.metric("Matched",           rpt["matched"])
                        col2.metric("Citations restored",rpt["citations_restored"])
                        col3.metric("Already had cites", rpt["already_had_cites"])
                        col4.metric("Unmatched",         rpt["unmatched"])
                        if rpt["unmatched"] > 0:
                            st.info(
                                f"{rpt['unmatched']} paragraph(s) are new content not in the old document "
                                f"— these could not have citations restored automatically."
                            )
                        st.download_button(
                            "⬇ Download merged document",
                            data=merged_bytes,
                            file_name=Path(new_fname).stem + "_citations_restored.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary"
                        )
                        st.markdown('''<div class="instruction-box">
                        <b>After downloading:</b><br>
                        1. Open in Word with your EndNote library connected<br>
                        2. EndNote tab → Update Citations and Bibliography<br>
                        3. If any citations still unlinked → use App 2 Remap tool<br>
                        4. Check any unmatched paragraphs manually
                        </div>''', unsafe_allow_html=True)
                        if rpt["details"]:
                            with st.expander(f"Paragraphs with restored citations ({len(rpt['details'])})"):
                                for d in rpt["details"]:
                                    st.markdown(
                                        f'<div class="ref-item ok">Para {d["para_idx"]+1} — {d["cites_added"]} citation run(s) restored<br><span style="font-size:0.8rem;font-style:italic">{d["text_preview"]}</span></div>',
                                        unsafe_allow_html=True
                                    )
                    except Exception as e:
                        st.error(f"Merge failed: {e}")
                        st.exception(e)

    else:
        st.markdown('''<div class="instruction-box">
        <b>When to use:</b> You merged two Word documents and EndNote no longer
        recognizes the citations. This tool accepts tracked changes safely
        (rescuing any citations inside deleted text) and repairs broken citation field codes.<br><br>
        <b>Best practice for future merges:</b> Before using Word's Compare, go to
        EndNote tab → Convert Citations → Convert to Unformatted Citations.
        This turns field codes into plain text like {Hall, 1997 #18} which survives
        merging perfectly. Then after accepting changes, use Update Citations and Bibliography.
        </div>''', unsafe_allow_html=True)

        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Post-merge document** — the merged file with broken citations")
            merged_file = st.file_uploader("Merged .docx", type=["docx"], key="merge_merged")
        with col2:
            st.markdown("**Original document** (optional) — used to detect lost citations")
            orig_file = st.file_uploader("Original .docx", type=["docx"], key="merge_orig")

        if merged_file:
            if st.button("Analyze & repair", type="primary"):
                with st.spinner("Analyzing citation damage..."):
                    merged_bytes = merged_file.read()
                    orig_bytes   = orig_file.read() if orig_file else None
                    analysis     = analyze_merge_damage(merged_bytes)

                st.markdown("### Damage report")
                bib_count   = analysis.get("bib_count", 0)
                cited_count = analysis.get("cited_count", 0)
                lost_count  = analysis.get("lost_in_merge", 0)
                field_count = analysis["total_en"]
                sees        = analysis.get("endnote_sees", analysis["with_data"])
                misses      = analysis.get("endnote_misses", analysis["empty_cite"])

                col1,col2,col3,col4 = st.columns(4)
                col1.metric("Bibliography entries",    bib_count,
                            help="Number of references in the numbered reference list")
                col2.metric("EndNote currently sees",  sees,
                            help="Unique references EndNote can read from working field codes")
                col3.metric("Hidden in broken fields", misses,
                            delta=f"-{misses}" if misses else None,
                            delta_color="inverse" if misses else "off",
                            help="References locked in broken field codes — EndNote cannot count these")
                col4.metric("Lost during merge",       lost_count,
                            delta=f"-{lost_count}" if lost_count else None,
                            delta_color="inverse" if lost_count else "off",
                            help="References whose inline citations were removed by the merge")

                if misses > 0:
                    st.error(
                        f"⚠ **EndNote is undercounting your references.** "
                        f"It currently sees **{sees}** unique references but **{misses}** more "
                        f"are locked in broken citation field codes and invisible to EndNote — "
                        f"causing your bibliography to show fewer entries than expected. "
                        f"This is fixed automatically below."
                    )
                if lost_count > 0:
                    st.warning(
                        f"⚠ {lost_count} citation(s) were lost during the merge. "
                        f"Use **Restore lost citations** mode with the original document to recover them."
                    )
                if misses == 0 and lost_count == 0 and bib_count > 0:
                    st.success(f"✓ All {bib_count} bibliography references are intact and visible to EndNote.")
                if not analysis["balanced"]:
                    st.warning(
                        f"⚠ Unbalanced field markers "
                        f"(begin:{analysis['begins']}/sep:{analysis['separates']}/end:{analysis['ends']}) "
                        f"— some citation fields were split during merge."
                    )
                elif field_count > 0:
                    st.success("✓ Citation field boundaries are intact.")
                if len(analysis["db_ids"]) > 1:
                    st.info(f"Multiple library fingerprints found ({len(analysis['db_ids'])}) — citations from different libraries.")

                with st.spinner("Repairing..."):
                    fixed_bytes, rpt = repair_post_merge_citations(merged_bytes, orig_bytes)
                st.markdown("### Results")
                col1,col2,col3 = st.columns(3)
                col1.metric("Citations before", rpt["citations_before"])
                col2.metric("Citations after",  rpt["citations_after"])
                col3.metric("Steps applied",    len(rpt["steps"]))
                for step in rpt["steps"]:
                    if step == "track_changes_accepted":
                        st.markdown("- ✓ Tracked changes accepted safely")
                    elif "restored" in step:
                        n = step.split("_")[1]
                        st.markdown(f"- ✓ {n} broken field(s) restored from backup data")
                    elif "lost" in step:
                        n = step.split("_")[0]
                        st.markdown(f"- ⚠ {n} citation(s) lost in merge — use Restore mode to recover")
                if rpt.get("lost_rec_nums"):
                    st.warning(f"⚠ {len(rpt['lost_rec_nums'])} citation(s) from original not found after merge.")
                st.download_button(
                    "⬇ Download repaired document",
                    data=fixed_bytes,
                    file_name=Path(merged_file.name).stem + "_repaired.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )
                st.markdown('''<div class="instruction-box">
                <b>After downloading:</b><br>
                1. Open in Word with your EndNote library connected<br>
                2. EndNote tab → Update Citations and Bibliography<br>
                3. Still issues? → Use App 2 Remap tool
                </div>''', unsafe_allow_html=True)

elif tool == "App 7 — Citation Renumbering":
    st.markdown('<div class="app-label">A practical tool &nbsp;07</div>', unsafe_allow_html=True)
    st.markdown("## Citation Renumbering")
    st.markdown('''<div class="instruction-box">
    <b>When to use:</b> After editing, citation superscript numbers and the bibliography
    are out of order. This tool renumbers them consistently using your chosen method.<br><br>
    <b>Before using:</b> Make sure EndNote has already formatted the bibliography so
    citations appear as plain superscript numbers in the text — not live field codes.
    </div>''', unsafe_allow_html=True)

    ren_file = st.file_uploader(
        "Word document (.docx) with formatted citations",
        type=["docx"], key="ren_doc"
    )

    method = st.radio(
        "Numbering method",
        [
            "Alphabetical — A=1, B=2... (sort bibliography by author last name)",
            "Order of appearance — first cited in text = 1, second = 2...",
        ]
    )

    if ren_file:
        if st.button("Renumber citations", type="primary"):
            with st.spinner("Scanning and renumbering..."):
                raw_bytes = ren_file.read()
                if "Alphabetical" in method:
                    fixed_bytes, mapping = renumber_citations_alpha(raw_bytes)
                    method_label = "alphabetically (A-Z by author)"
                else:
                    fixed_bytes, mapping = renumber_citations_appearance(raw_bytes)
                    method_label = "by order of appearance"

            if not mapping:
                st.warning(
                    "No superscript citation numbers or numbered bibliography found. "
                    "Make sure EndNote has formatted the bibliography first."
                )
            else:
                changed = {k: v for k, v in mapping.items() if k != v}
                st.success(
                    f"Done — renumbered {method_label}. "
                    f"{len(mapping)} unique citations, {len(changed)} numbers changed."
                )
                col1, col2, col3 = st.columns(3)
                col1.metric("Unique citations",  len(mapping))
                col2.metric("Numbers changed",   len(changed))
                col3.metric("Already in order",  len(mapping) - len(changed))

                st.download_button(
                    "⬇ Download renumbered document",
                    data=fixed_bytes,
                    file_name=Path(ren_file.name).stem + "_renumbered.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )

                if changed:
                    with st.expander(f"Renumbering map ({len(changed)} changes)"):
                        col1, col2 = st.columns(2)
                        col1.markdown("**Old #**")
                        col2.markdown("**New #**")
                        for old in sorted(mapping):
                            new = mapping[old]
                            if old != new:
                                col1.markdown(str(old))
                                col2.markdown(str(new))

# ─────────────────────────────────────────────────────────────────────────────
# APP 6 UI — FIGURE INVENTORY
# ─────────────────────────────────────────────────────────────────────────────
elif tool == "App 8 — Figure Inventory":
    st.markdown("## Figure & Table Inventory")
    st.markdown('<div class="instruction-box">Scans your document for all figures, tables, boxes, plates, and videos. Optionally cross-references against your Excel naming sheet to flag mismatches and missing items.</div>',unsafe_allow_html=True)
    col1,col2=st.columns(2)
    with col1: fig_doc=st.file_uploader("Word document (.docx)",type=["docx"],key="fig_doc")
    with col2: fig_excel=st.file_uploader("Excel naming sheet (optional)",type=["xlsx","xls"],key="fig_excel",help="Columns: Number | Name/Caption")
    if fig_doc:
        if st.button("Scan document",type="primary"):
            with st.spinner("Scanning..."):
                fig_bytes=fig_doc.read(); items=scan_figures(fig_bytes)
            if not items:
                st.warning("No captioned items found. Captions must start with Figure/Table/Box/Plate/Video followed by a number.")
            else:
                type_counts={}
                for item in items: type_counts[item['type']]=type_counts.get(item['type'],0)+1
                cols=st.columns(min(len(type_counts),6))
                for i,(t,c) in enumerate(sorted(type_counts.items())): cols[i%len(cols)].metric(t+"s",c)
                st.markdown(f"**{len(items)} total items**")
                if fig_excel:
                    excel_bytes=fig_excel.read(); results=cross_ref_excel(items,excel_bytes)
                    matched=[r for r in results if r[1]=='match']
                    mismatch=[r for r in results if r[1]=='mismatch']
                    not_found=[r for r in results if r[1]=='not_in_excel']
                    col1,col2,col3,col4=st.columns(4)
                    col1.metric("Matched",len(matched)); col2.metric("Mismatch",len(mismatch))
                    col3.metric("Not in Excel",len(not_found))
                    tabs=st.tabs([f"Mismatches ({len(mismatch)})",f"Not in Excel ({len(not_found)})",
                                  f"Matched ({len(matched)})",f"All ({len(results)})"])
                    def show_res(res_list):
                        for item,status,expected in res_list:
                            color='ok' if status=='match' else 'missing' if status=='not_in_excel' else 'warn'
                            exp=""
                            if expected and status=='mismatch': exp=f'<br><span style="font-size:0.75rem;color:#ffb300">Expected: {expected}</span>'
                            st.markdown(f'<div class="ref-item {color}"><b>{item["type"]} {item["number"]}</b>{"🖼" if item["has_image"] else ""} — Para {item["para_idx"]+1}<br><span style="font-size:0.82rem">{item["caption"][:180]}</span>{exp}</div>',unsafe_allow_html=True)
                    with tabs[0]: show_res(mismatch) if mismatch else st.success("No mismatches.")
                    with tabs[1]: show_res(not_found) if not_found else st.success("All found in Excel.")
                    with tabs[2]: show_res(matched) if matched else st.info("No confirmed matches.")
                    with tabs[3]: show_res(results)
                else:
                    type_filter=st.selectbox("Filter by type",["All"]+sorted(type_counts.keys()))
                    for item in items:
                        if type_filter!="All" and item['type']!=type_filter: continue
                        st.markdown(f'<div class="match-card"><span class="match-marker">{item["type"]} {item["number"]}</span>{"🖼" if item["has_image"] else ""}<span style="font-size:0.72rem;color:#3a4a5a"> Para {item["para_idx"]+1}</span><div class="match-sentence">{item["caption"][:200]}</div></div>',unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# APP 7 UI — PUBMED SEARCH
# ─────────────────────────────────────────────────────────────────────────────
elif tool == "App 1 — PubMed Search":
    st.markdown("## PubMed Literature Search")
    st.markdown('<div class="instruction-box">Search PubMed for relevant articles. Results include abstracts and links to free full text where available (PubMed Central). Export results as EndNote XML to import directly into your EndNote library.</div>',unsafe_allow_html=True)
    query=st.text_input("Search query",placeholder="e.g. arthrogryposis clubfoot Ponseti treatment",help="Supports PubMed syntax: AND, OR, NOT, [MeSH], [ti], [au]")
    col1,col2,col3=st.columns(3)
    with col1: date_from=st.text_input("Year from",placeholder="2019")
    with col2: date_to=st.text_input("Year to",placeholder="2026")
    with col3: max_res=st.slider("Max results",5,50,20)
    journal_filter=st.text_input("Limit to journal (optional)",placeholder="e.g. J Pediatr Orthop")
    if st.button("Search PubMed",type="primary",disabled=not query.strip()):
        with st.spinner(f"Searching PubMed: {query}..."):
            results=pubmed_search_full(query.strip(),date_from.strip(),date_to.strip(),journal_filter.strip(),max_res)
        if not results: st.warning("No results found. Try broadening your search terms.")
        else:
            st.success(f"Found {len(results)} results.")
            xml_export=results_to_xml(results)
            st.download_button(f"⬇ Export all {len(results)} refs as EndNote XML",
                data=xml_export.encode('utf-8'),file_name="pubmed_results.xml",mime="application/xml")
            st.divider()
            for i,r in enumerate(results,1):
                authors_str='; '.join(r['authors'][:3])+(' et al.' if len(r['authors'])>3 else '')
                cit=f"{authors_str} ({r['year']}). *{r['journal']}*"
                if r['volume']: cit+=f" {r['volume']}"
                if r['issue']:  cit+=f"({r['issue']})"
                if r['pages']:  cit+=f":{r['pages']}"
                with st.expander(f"**{i}.** {r['title'][:100]}{'...' if len(r['title'])>100 else ''}"):
                    st.markdown(cit)
                    lc=st.columns(3)
                    with lc[0]: st.markdown(f"[PubMed]({r['pubmed_url']})")
                    with lc[1]:
                        if r['pmc_url']: st.markdown(f"[Free full text (PMC)]({r['pmc_url']})")
                        elif r['doi_url']: st.markdown(f"[DOI]({r['doi_url']})")
                        else: st.markdown("*No free full text*")
                    with lc[2]:
                        if r['doi']: st.caption(f"DOI: {r['doi']}")
                    if r['abstract']: st.markdown("**Abstract:**"); st.markdown(r['abstract'])
                    single=results_to_xml([r])
                    st.download_button("⬇ Export this ref",data=single.encode('utf-8'),
                        file_name=f"ref_{r['pmid']}.xml",mime="application/xml",key=f"exp_{r['pmid']}")

# ─────────────────────────────────────────────────────────────────────────────
# APP 8 UI — BATCH RENAME
# ─────────────────────────────────────────────────────────────────────────────
elif tool == "App 10 — Batch Rename":
    st.markdown('<div class="app-label">A practical tool &nbsp;10</div>', unsafe_allow_html=True)
    st.markdown("## Batch Rename")
    st.markdown('<div class="instruction-box">Apply bulk find-and-replace across your Word document using an Excel naming sheet. Use for final editing to update chapter names, author names, figure labels, section titles, etc.<br><br><b>Excel format:</b> Two columns — <code>Old Name</code> (find) and <code>New Name</code> (replace). One row per replacement.</div>',unsafe_allow_html=True)
    col1,col2=st.columns(2)
    with col1: ren_doc=st.file_uploader("Word document (.docx)",type=["docx"],key="batch_doc")
    with col2: ren_excel=st.file_uploader("Excel naming sheet (.xlsx)",type=["xlsx","xls"],key="batch_excel")
    col3,col4=st.columns(2)
    with col3: match_case=st.toggle("Match case",value=False)
    with col4: whole_word=st.toggle("Whole word only",value=False)
    pairs=[]
    if ren_excel:
        with st.spinner("Reading naming sheet..."):
            pairs=load_rename_pairs(ren_excel.read())
        if pairs:
            st.success(f"✓ {len(pairs)} rename pairs loaded.")
            with st.expander("Preview"):
                col1,col2=st.columns(2)
                col1.markdown("**Find**"); col2.markdown("**Replace**")
                for old,new in pairs[:20]: col1.markdown(old); col2.markdown(new)
                if len(pairs)>20: st.caption(f"...and {len(pairs)-20} more")
        else: st.error("Could not read pairs. Check Excel has 'Old Name' and 'New Name' columns.")
    if ren_doc and ren_excel and pairs:
        if st.button("Apply batch rename",type="primary"):
            with st.spinner("Applying..."):
                fixed_bytes,report=batch_rename(ren_doc.read(),pairs,match_case,whole_word)
            replaced=[r for r in report if r['status']=='replaced']
            not_found=[r for r in report if r['status']=='not_found']
            total_changes=sum(r['count'] for r in replaced)
            col1,col2,col3=st.columns(3)
            col1.metric("Total pairs",len(report)); col2.metric("Replaced",len(replaced)); col3.metric("Not found",len(not_found))
            st.success(f"✓ {total_changes} replacement(s) across {len(replaced)} pairs.")
            st.download_button("⬇ Download renamed document",data=fixed_bytes,
                file_name=Path(ren_doc.name).stem+"_renamed.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",type="primary")
            if replaced:
                with st.expander(f"Replacements made ({len(replaced)})"):
                    for r in replaced:
                        st.markdown(f'<div class="ref-item ok"><b>{r["old"]}</b> → {r["new"]} <span style="float:right;font-size:0.75rem">{r["count"]}x</span></div>',unsafe_allow_html=True)
            if not_found:
                with st.expander(f"Not found ({len(not_found)})"):
                    for r in not_found:
                        st.markdown(f'<div class="ref-item error"><b>{r["old"]}</b> — not found in document</div>',unsafe_allow_html=True)

elif tool == "App 5 — Bibliography Auditor":
    st.markdown('<div class="app-label">A practical tool &nbsp;05</div>', unsafe_allow_html=True)
    st.markdown("## Bibliography Auditor")
    st.markdown('''<div class="instruction-box">
    Cross-reference your Word document bibliography against the published PDF — for a specific chapter.
    Specify page ranges so the app only reads the relevant pages rather than the whole book.<br><br>
    <b>Finds:</b>
    <ul style="margin:0.4rem 0 0 1rem">
      <li>Refs in the published PDF missing from the Word doc</li>
      <li>Refs in the Word doc not in the published PDF (new additions)</li>
      <li>Published PDF refs not cited in the published body text</li>
      <li>Published PDF refs not cited in the Word doc body text</li>
      <li>Generates an importable EndNote file (.enw) for any missing refs</li>
    </ul>
    </div>''', unsafe_allow_html=True)

    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**Word document (.docx)**")
        audit_doc = st.file_uploader("Word .docx", type=["docx"], key="audit_doc")
    with col2:
        st.markdown("**Published PDF (full book or chapter)**")
        audit_pdf = st.file_uploader("Published PDF", type=["pdf"], key="audit_pdf")

    if audit_pdf:
        import pdfplumber as _plumber
        with _plumber.open(audit_pdf) as _pp:
            _total_pages = len(_pp.pages)
        st.caption(f"PDF has {_total_pages} pages total")

        st.markdown("**Specify page ranges** — leave blank to use the whole PDF")
        pr1, pr2 = st.columns(2)
        with pr1:
            st.markdown("Body text pages (chapter content)")
            body_p1 = st.number_input("From page", min_value=1, max_value=_total_pages,
                                       value=1, key="body_p1")
            body_p2 = st.number_input("To page",   min_value=1, max_value=_total_pages,
                                       value=_total_pages, key="body_p2")
        with pr2:
            st.markdown("Bibliography pages")
            bib_p1 = st.number_input("From page", min_value=1, max_value=_total_pages,
                                      value=max(1, _total_pages-5), key="bib_p1")
            bib_p2 = st.number_input("To page",   min_value=1, max_value=_total_pages,
                                      value=_total_pages, key="bib_p2")

    if audit_doc and audit_pdf:
        if st.button("Run bibliography audit", type="primary"):
            import pdfplumber as _plumber, re as _re
            from docx import Document as _Doc
            from lxml import etree as _et

            with st.spinner("Extracting bibliography from PDF pages "
                            f"{bib_p1}–{bib_p2}..."):
                # ── Extract PDF bibliography (specified pages only) ──────────
                pdf_bib_text = ""
                with _plumber.open(audit_pdf) as _pdf:
                    for pg in _pdf.pages[bib_p1-1:bib_p2]:
                        t = pg.extract_text()
                        if t: pdf_bib_text += t + "\n"

                pdf_refs = {}
                # Two-column layout: refs appear as "N.\t..." anywhere on line
                for m in _re.finditer(
                    r'(?:^|\s{2,})(\d+[a-z]?)\.\s*\x07?([^\n]{10,})',
                    pdf_bib_text, _re.MULTILINE):
                    num  = m.group(1)
                    text = _re.sub(r'\s+', ' ', m.group(2)).strip()
                    if len(text) > 15 and num not in pdf_refs:
                        pdf_refs[num] = text

            with st.spinner("Extracting body text from PDF pages "
                            f"{body_p1}–{body_p2}..."):
                # ── Extract PDF body text (specified pages only) ─────────────
                pdf_body_text = ""
                with _plumber.open(audit_pdf) as _pdf:
                    for pg in _pdf.pages[body_p1-1:body_p2]:
                        t = pg.extract_text()
                        if t: pdf_body_text += t + "\n"

                # Find all citation numbers in PDF body text
                # Superscripts in published text appear as plain numbers after words
                pdf_body_cited = set()
                for m in _re.finditer(
                    r'(?<=[a-z\.\,\)])(\d+[a-z]?)(?=[,\.\s\n]|$)',
                    pdf_body_text, _re.MULTILINE | _re.IGNORECASE):
                    val = m.group(1)
                    if val.rstrip('a-z').isdigit():
                        pdf_body_cited.add(val)

            with st.spinner("Reading Word document..."):
                # ── Extract Word bibliography ────────────────────────────────
                wdoc = _Doc(audit_doc)
                ref_pat = _re.compile(r'^\s*(\d+[a-z]?)[\.)\s]\s+(.+)')
                word_bib = {}
                for p in wdoc.paragraphs:
                    m2 = ref_pat.match(p.text.strip())
                    if m2: word_bib[m2.group(1)] = p.text.strip()

                # ── Word body superscript citation numbers ───────────────────
                with zipfile.ZipFile(io.BytesIO(audit_doc.getvalue())) as _z:
                    doc_xml = _z.read('word/document.xml').decode('utf-8')
                W_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                root_elem = _et.fromstring(doc_xml.encode('utf-8'))
                word_body_cited = set()
                for r in root_elem.iter(f'{{{W_ns}}}r'):
                    rpr = r.find(f'{{{W_ns}}}rPr')
                    if rpr is None: continue
                    va = rpr.find(f'{{{W_ns}}}vertAlign')
                    if va is None or va.get(f'{{{W_ns}}}val') != 'superscript': continue
                    t_el = r.find(f'{{{W_ns}}}t')
                    txt  = (t_el.text or '') if t_el is not None else ''
                    for part in _re.split(r'[,;\s]+', txt):
                        p2 = part.strip().rstrip('.')
                        if p2.isdigit(): word_body_cited.add(p2)

                cited_rns = set(_re.findall(r'&lt;RecNum&gt;(\d+)&lt;/RecNum&gt;', doc_xml))
                total_field_rns = len(cited_rns)

            # ── Compare bibliographies ───────────────────────────────────────
            def ref_key(text):
                text = _re.sub(r'^\d+[a-z]?[\.)\s]\s*', '', text).strip()
                parts = _re.split(r',|;', text)
                auth  = parts[0].strip().split()[-1].lower() if parts[0].strip() else ''
                yr_m  = _re.search(r'\b(19|20)\d{2}\b', text)
                yr    = yr_m.group(0) if yr_m else ''
                return f"{auth} {yr}"

            word_keys = {ref_key(v): k for k, v in word_bib.items()}
            pdf_keys  = {ref_key(v): k for k, v in pdf_refs.items()}

            missing_from_word = [(num, pdf_refs[num]) for key, num in pdf_keys.items()
                                  if key not in word_keys]
            extra_in_word     = [(num, word_bib[num])  for key, num in word_keys.items()
                                  if key not in pdf_keys]

            pdf_nums  = set(pdf_refs.keys())
            not_in_pdf_body  = {n for n in pdf_nums if n not in pdf_body_cited}
            not_in_word_body = {n for n in pdf_nums if n not in word_body_cited}

            # ── Results ──────────────────────────────────────────────────────
            st.divider()
            st.markdown("### Section 1 — Bibliography comparison")
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Published PDF refs",     len(pdf_refs))
            c2.metric("Word doc refs",          len(word_bib))
            c3.metric("Missing from Word",      len(missing_from_word),
                      delta=f"-{len(missing_from_word)}" if missing_from_word else None,
                      delta_color="inverse" if missing_from_word else "off")
            c4.metric("EndNote field codes",    total_field_rns)

            if missing_from_word:
                st.error(f"⚠ **{len(missing_from_word)} ref(s) in the published PDF are missing from the Word document.**")
                with st.expander(f"Missing from Word ({len(missing_from_word)})", expanded=True):
                    for num, text in sorted(missing_from_word, key=lambda x: int(x[0]) if x[0].isdigit() else 999):
                        st.markdown(f'<div class="ref-item error">#{num} — {text[:100]}</div>',
                                    unsafe_allow_html=True)
                # Generate ENW
                enw_lines = []
                for num, text in missing_from_word:
                    auth_m = _re.match(
                        r'([A-Za-z\'\-]+)\s+(?:[A-Z]+\.?\s+)?(.+?)\.\s+(.+?)\.\s+(\d{4})', text)
                    if auth_m:
                        enw_lines += ['%0 Journal Article',
                                      f'%A {auth_m.group(1)}',
                                      f'%T {auth_m.group(2).strip()}',
                                      f'%J {auth_m.group(3).strip()}',
                                      f'%D {auth_m.group(4)}', '']
                    else:
                        enw_lines += ['%0 Journal Article', f'%T {text[:120]}', '']
                if enw_lines:
                    st.download_button("⬇ Download missing refs as EndNote import (.enw)",
                                       data='\n'.join(enw_lines),
                                       file_name="missing_refs.enw", mime="text/plain")

            if extra_in_word:
                st.info(f"ℹ {len(extra_in_word)} ref(s) in Word doc not in published PDF (new additions):")
                for num, text in extra_in_word[:15]:
                    st.markdown(f"  • #{num} — {text[:80]}")

            if not missing_from_word and not extra_in_word:
                st.success("✓ Word doc and published PDF bibliographies match perfectly.")

            st.divider()
            st.markdown("### Section 2 — Citations in published PDF body text")
            st.caption(f"Checked pages {body_p1}–{body_p2}")
            c5, c6 = st.columns(2)
            c5.metric("PDF refs cited in published body", len(pdf_nums) - len(not_in_pdf_body))
            c6.metric("PDF refs NOT cited in published body", len(not_in_pdf_body),
                      delta=str(len(not_in_pdf_body)) if not_in_pdf_body else None,
                      delta_color="inverse" if not_in_pdf_body else "off",
                      help="In the bibliography but no citation number found in the chapter body text")

            if not_in_pdf_body:
                with st.expander(f"Not cited in published body ({len(not_in_pdf_body)})", expanded=True):
                    st.caption("These appear in the published bibliography but no citation number "
                               "was detected in the published body text pages specified.")
                    for n in sorted(not_in_pdf_body, key=lambda x: int(x) if x.isdigit() else 999):
                        ref_text = pdf_refs.get(n, 'Unknown ref')
                        st.markdown(f'<div class="ref-item warning">#{n} — {ref_text[:100]}</div>',
                                    unsafe_allow_html=True)
            else:
                st.success("✓ All published PDF refs appear to be cited in the published body text.")

            st.divider()
            st.markdown("### Section 3 — Citations in Word document body text")
            c7, c8 = st.columns(2)
            c7.metric("PDF refs cited in Word body", len(pdf_nums) - len(not_in_word_body))
            c8.metric("PDF refs NOT cited in Word body", len(not_in_word_body),
                      delta=str(len(not_in_word_body)) if not_in_word_body else None,
                      delta_color="inverse" if not_in_word_body else "off",
                      help="In the published bibliography but no inline superscript found in the Word body text")

            if not_in_word_body:
                with st.expander(f"Not cited in Word body ({len(not_in_word_body)})", expanded=True):
                    st.caption("These appear in the published bibliography but their reference "
                               "number was not found as an inline superscript in the Word document. "
                               "May indicate citations lost during merge, or renumbering issues.")
                    for n in sorted(not_in_word_body, key=lambda x: int(x) if x.isdigit() else 999):
                        ref_text = pdf_refs.get(n, 'Unknown ref')
                        st.markdown(f'<div class="ref-item error">#{n} — {ref_text[:100]}</div>',
                                    unsafe_allow_html=True)
            else:
                st.success("✓ All published PDF refs appear as inline citations in the Word document.")

elif tool == "App 9 — Document Health Check":
    st.markdown('<div class="app-label">A practical tool &nbsp;09</div>', unsafe_allow_html=True)
    st.markdown("## Document Health Check")
    st.markdown('''<div class="instruction-box">
    Scans your Word document for common problems and presents each issue as a
    click-through review — accept individual fixes or accept all at once.<br><br>
    <b>Checks:</b> orphaned superscripts · garbled figure numbers (3334 → 33) ·
    broken citation fields · duplicate w:id values · orphan comment markup ·
    plain-text bibliography entries that should be EndNote fields
    </div>''', unsafe_allow_html=True)

    health_file = st.file_uploader("Word document (.docx)", type=["docx"], key="health_doc")

    if health_file:
        if st.button("Run health check", type="primary"):
            with st.spinner("Scanning document..."):
                from lxml import etree as _et
                from docx import Document as _Doc
                import base64 as _b64

                raw_bytes = health_file.read()
                with zipfile.ZipFile(io.BytesIO(raw_bytes)) as z:
                    doc_xml   = z.read('word/document.xml').decode('utf-8')
                    all_zfiles = {n: z.read(n) for n in z.namelist()}

                doc = _Doc(io.BytesIO(raw_bytes))
                W   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

                issues = []

                # ── 1. Broken citation fields ──────────────────────────────
                broken_a = len(re.findall(r'<w:instrText[^>]*> ADDIN EN\.CITE </w:instrText>', doc_xml))
                broken_b = len(re.findall(r'<w:instrText[^>]*>\s*ADDIN EN\.CITE\.DATA\s*</w:instrText>', doc_xml))
                if broken_a + broken_b > 0:
                    issues.append({
                        'id': 'broken_fields',
                        'severity': 'error',
                        'title': f'{broken_a + broken_b} broken citation field(s)',
                        'detail': f'{broken_a} empty instrText + {broken_b} ADDIN EN.CITE.DATA — EndNote cannot see these refs, causing bibliography undercount.',
                        'fix_label': 'Fix broken fields',
                        'count': broken_a + broken_b,
                    })

                # ── 2. Hidden refs in fldData ──────────────────────────────
                working_rns = set(re.findall(r'&lt;RecNum&gt;(\d+)&lt;/RecNum&gt;', doc_xml))
                fld_rns = set()
                for b64r in re.findall(r'<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>', doc_xml):
                    b64 = b64r.replace('\r','').replace('\n','').replace(' ','')
                    pad = (4-len(b64)%4)%4
                    try:
                        dec = _b64.b64decode(b64+'='*pad).decode('utf-8',errors='replace').replace('\x00','')
                        for rn in re.findall(r'<RecNum>(\d+)</RecNum>', dec): fld_rns.add(rn)
                    except: pass
                hidden = fld_rns - working_rns
                if hidden:
                    issues.append({
                        'id': 'hidden_refs',
                        'severity': 'error',
                        'title': f'{len(hidden)} ref(s) hidden in fldData',
                        'detail': f'RecNums {sorted(int(x) for x in hidden)[:10]}{"..." if len(hidden)>10 else ""} — locked in backup blobs, invisible to EndNote.',
                        'fix_label': 'Restore hidden refs',
                        'count': len(hidden),
                    })

                # ── 3. Orphaned superscripts ───────────────────────────────
                orphan_count = 0
                for para in doc.paragraphs:
                    has_field = any('EN.CITE' in (i.text or '') for i in para._p.findall(f'.//{{{W}}}instrText'))
                    if not has_field: continue
                    fd = 0
                    for run in para._p.findall(f'{{{W}}}r'):
                        for fc in run.findall(f'.//{{{W}}}fldChar'):
                            ft = fc.get(f'{{{W}}}fldCharType','')
                            if ft=='begin': fd+=1
                            elif ft=='end': fd=max(0,fd-1)
                        if fd>0: continue
                        rpr = run.find(f'{{{W}}}rPr')
                        if rpr is None: continue
                        va = rpr.find(f'{{{W}}}vertAlign')
                        if va is None or va.get(f'{{{W}}}val')!='superscript': continue
                        t = run.find(f'{{{W}}}t')
                        text = (t.text or '') if t is not None else ''
                        if text.strip() and re.match(r'[\d,;\s\.]+$', text.strip()):
                            orphan_count += 1
                if orphan_count:
                    issues.append({
                        'id': 'orphan_superscripts',
                        'severity': 'warning',
                        'title': f'{orphan_count} orphaned superscript(s)',
                        'detail': 'Plain-text citation numbers sitting next to EndNote field codes — visible to reader but invisible to EndNote.',
                        'fix_label': 'Remove orphaned superscripts',
                        'count': orphan_count,
                    })

                # ── 4. Garbled figure numbers (33+34 split) ────────────────
                fig_count = 0
                for para in doc.paragraphs:
                    runs = list(para.runs)
                    for i in range(len(runs)-1):
                        if (runs[i].text or '').endswith('33') and (runs[i+1].text or '').startswith('34'):
                            fig_count += 1
                if fig_count:
                    issues.append({
                        'id': 'garbled_figs',
                        'severity': 'warning',
                        'title': f'{fig_count} garbled figure number(s)',
                        'detail': 'Chapter numbers got concatenated during merge (e.g. "Fig. 3334.1" instead of "Fig. 33.1"). Each occurrence will have the extra "34" removed.',
                        'fix_label': 'Fix figure numbers',
                        'count': fig_count,
                    })

                # ── 5. Duplicate w:id values ───────────────────────────────
                from collections import Counter as _Ctr
                dup_ids = len({k for k,v in _Ctr(re.findall(r'\bw:id="(\d+)"', doc_xml)).items() if v>1})
                if dup_ids:
                    issues.append({
                        'id': 'dup_ids',
                        'severity': 'error',
                        'title': f'{dup_ids} duplicate w:id value(s)',
                        'detail': 'Internal ID conflicts caused by copying paragraph XML between documents. Causes Word to show "unreadable content" or repair dialogs.',
                        'fix_label': 'Deduplicate IDs',
                        'count': dup_ids,
                    })

                # ── 6. Orphan comment markup ───────────────────────────────
                com_markup = len(re.findall(r'<w:comment(?:RangeStart|RangeEnd|Reference)\b', doc_xml))
                has_comments_xml = 'word/comments.xml' in all_zfiles
                if com_markup > 0 and not has_comments_xml:
                    issues.append({
                        'id': 'orphan_comments',
                        'severity': 'warning',
                        'title': f'{com_markup} orphaned comment markup element(s)',
                        'detail': 'Comment reference markers exist in the document body but no comments.xml file exists. Causes Word repair dialog on open.',
                        'fix_label': 'Remove orphaned comment markup',
                        'count': com_markup,
                    })

                # ── 7. Unbalanced field codes ──────────────────────────────
                begins = doc_xml.count('fldCharType="begin"')
                ends   = doc_xml.count('fldCharType="end"')
                seps   = doc_xml.count('fldCharType="separate"')
                if begins != ends or begins != seps:
                    issues.append({
                        'id': 'unbalanced_fields',
                        'severity': 'error',
                        'title': f'Unbalanced field markers (begin:{begins}/sep:{seps}/end:{ends})',
                        'detail': 'Field code begin/separate/end markers are mismatched. This causes citation formatting failures and can make Word reject the file.',
                        'fix_label': 'Cannot auto-fix — manual repair needed',
                        'count': abs(begins - ends),
                    })

                st.session_state['health_issues']    = issues
                st.session_state['health_raw_bytes'] = raw_bytes
                st.session_state['health_doc_name']  = health_file.name

        if 'health_issues' in st.session_state and st.session_state['health_issues'] is not None:
            issues    = st.session_state['health_issues']
            raw_bytes = st.session_state['health_raw_bytes']
            doc_name  = st.session_state['health_doc_name']

            if not issues:
                st.success("✓ No issues found — document is clean.")
            else:
                errors   = [i for i in issues if i['severity']=='error']
                warnings = [i for i in issues if i['severity']=='warning']
                c1, c2, c3 = st.columns(3)
                c1.metric("Issues found", len(issues))
                c2.metric("Errors",   len(errors),   delta=str(len(errors))   if errors   else None, delta_color="inverse")
                c3.metric("Warnings", len(warnings), delta=str(len(warnings)) if warnings else None, delta_color="inverse")

                # Accept All button
                fixable = [i for i in issues if 'Cannot' not in i['fix_label']]
                if st.button(f"✓ Accept all {len(fixable)} fixable issue(s)", type="primary"):
                    st.session_state['health_accept_all'] = True

                st.divider()
                st.markdown("### Review each issue")

                # Per-issue accept checkboxes
                accepted = {}
                for issue in issues:
                    fixable_issue = 'Cannot' not in issue['fix_label']
                    icon = '🔴' if issue['severity']=='error' else '🟡'
                    with st.expander(f"{icon} {issue['title']}", expanded=True):
                        st.markdown(issue['detail'])
                        if fixable_issue:
                            accepted[issue['id']] = st.checkbox(
                                f"Accept fix: {issue['fix_label']}",
                                value=st.session_state.get('health_accept_all', False),
                                key=f"accept_{issue['id']}"
                            )
                        else:
                            st.warning(issue['fix_label'])

                if st.button("Apply accepted fixes"):
                    with st.spinner("Applying fixes..."):
                        with zipfile.ZipFile(io.BytesIO(raw_bytes)) as z:
                            doc_xml   = z.read('word/document.xml').decode('utf-8')
                            all_zfiles = {n: z.read(n) for n in z.namelist()}

                        applied = []

                        if accepted.get('broken_fields') or accepted.get('hidden_refs'):
                            doc_xml, n = fix_broken_fields(doc_xml)
                            applied.append(f"Fixed {n} broken field(s)")

                        if accepted.get('orphan_comments'):
                            before = len(re.findall(r'<w:comment(?:RangeStart|RangeEnd|Reference)\b', doc_xml))
                            doc_xml = re.sub(r'<w:commentRangeStart\b[^/]*/>', '', doc_xml)
                            doc_xml = re.sub(r'<w:commentRangeEnd\b[^/]*/>', '', doc_xml)
                            doc_xml = re.sub(r'<w:commentReference\b[^/]*/>', '', doc_xml)
                            applied.append(f"Removed {before} orphaned comment element(s)")

                        if accepted.get('dup_ids'):
                            from collections import Counter as _Ctr2
                            _ids = re.findall(r'\bw:id="(\d+)"', doc_xml)
                            _max = max(int(x) for x in _ids)
                            _nxt = [_max+1]; _seen = set()
                            def _fid(m):
                                v=m.group(2)
                                if v in _seen:
                                    nw=str(_nxt[0]); _nxt[0]+=1; return f'{m.group(1)}{nw}{m.group(3)}'
                                _seen.add(v); return m.group(0)
                            doc_xml = re.sub(r'(w:id=")(\d+)(")', _fid, doc_xml)
                            applied.append("Deduplicated w:id values")

                        if accepted.get('garbled_figs'):
                            from docx import Document as _Doc2
                            all_zfiles['word/document.xml'] = doc_xml.encode('utf-8')
                            buf2 = io.BytesIO()
                            with zipfile.ZipFile(buf2,'w',zipfile.ZIP_DEFLATED) as zt:
                                for nn,dd in all_zfiles.items(): zt.writestr(nn,dd)
                            buf2.seek(0)
                            d2 = _Doc2(buf2)
                            W2 = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                            n_fig = 0
                            for para in d2.paragraphs:
                                runs = list(para.runs)
                                for i in range(len(runs)-1):
                                    t1=(runs[i].text or ''); t2=(runs[i+1].text or '')
                                    if t1.endswith('33') and t2.startswith('34'):
                                        nt = t2[2:]
                                        te = runs[i+1]._r.find(f'{{{W2}}}t')
                                        if te is not None:
                                            te.text = nt
                                            if not nt: runs[i+1]._r.getparent().remove(runs[i+1]._r)
                                            n_fig += 1
                            buf3 = io.BytesIO(); d2.save(buf3); buf3.seek(0)
                            with zipfile.ZipFile(buf3) as zt2:
                                doc_xml = zt2.read('word/document.xml').decode('utf-8')
                                for nn in zt2.namelist():
                                    all_zfiles[nn] = zt2.read(nn)
                            applied.append(f"Fixed {n_fig} garbled figure number(s)")

                        if accepted.get('orphan_superscripts'):
                            all_zfiles['word/document.xml'] = doc_xml.encode('utf-8')
                            buf4 = io.BytesIO()
                            with zipfile.ZipFile(buf4,'w',zipfile.ZIP_DEFLATED) as zt:
                                for nn,dd in all_zfiles.items(): zt.writestr(nn,dd)
                            buf4.seek(0)
                            fixed_bytes, n_orp = remove_orphan_superscripts(buf4.read())
                            with zipfile.ZipFile(io.BytesIO(fixed_bytes)) as zt3:
                                doc_xml = zt3.read('word/document.xml').decode('utf-8')
                                for nn in zt3.namelist():
                                    all_zfiles[nn] = zt3.read(nn)
                            applied.append(f"Removed {n_orp} orphaned superscript(s)")

                        # Save final
                        all_zfiles['word/document.xml'] = doc_xml.encode('utf-8')
                        out_buf = io.BytesIO()
                        with zipfile.ZipFile(out_buf,'w',zipfile.ZIP_DEFLATED) as zout:
                            for nn,dd in all_zfiles.items(): zout.writestr(nn,dd)
                        out_buf.seek(0)
                        final_bytes = out_buf.read()

                    if applied:
                        for msg in applied:
                            st.success(f"✓ {msg}")
                        st.download_button(
                            "⬇ Download repaired document",
                            data=final_bytes,
                            file_name=Path(doc_name).stem + "_repaired.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary"
                        )
                        # Re-run health check on fixed doc
                        if st.button("Re-run health check on repaired document"):
                            st.session_state['health_issues'] = None
                            st.rerun()
                    else:
                        st.info("No fixes were accepted.")

        # Initialise state
        if 'health_issues' not in st.session_state:
            st.session_state['health_issues'] = None
        if 'health_accept_all' not in st.session_state:
            st.session_state['health_accept_all'] = False

elif tool == "App 11 — RecNum Inspector":
    st.markdown('<div class="app-label">A practical tool &nbsp;11</div>', unsafe_allow_html=True)
    st.markdown("## RecNum Inspector")
    st.markdown('''<div class="instruction-box">
    Reads a Word document that is connected to your EndNote library and shows you every
    reference's RecNum (Record Number), author, year, and title — so you can identify
    the correct RecNums to use when fixing field codes.<br><br>
    <b>Also:</b> if you have field codes using the wrong RecNums (e.g. placeholder RecNums
    500-506 we assigned before knowing your library's actual numbers), paste a remapping
    table and this tool will fix all field codes in one step and download the corrected document.
    </div>''', unsafe_allow_html=True)

    rn_file = st.file_uploader(
        "Word document (.docx) — must be library-connected (have working EN.CITE fields)",
        type=["docx"], key="rn_doc")

    if rn_file:
        raw_bytes = rn_file.read()
        with zipfile.ZipFile(io.BytesIO(raw_bytes)) as z:
            rn_xml = z.read('word/document.xml').decode('utf-8')

        # Extract all unique RecNum -> author/year/title from instrText
        import html as _html
        rn_data = {}
        for m in re.finditer(
                r'ADDIN EN\.CITE (&lt;EndNote&gt;[\s\S]+?)(?=</w:instrText>)', rn_xml):
            decoded = _html.unescape(m.group(1))
            rns     = re.findall(r'<RecNum>(\d+)</RecNum>', decoded)
            authors = re.findall(r'<Author>([^<]+)</Author>', decoded)
            years   = re.findall(r'<Year>([^<]+)</Year>', decoded)
            titles  = re.findall(r'<title>([^<]+)</title>', decoded)
            db_ids  = re.findall(r'db-id="([^"]*)"', decoded)
            for i, rn in enumerate(rns):
                if rn not in rn_data:
                    auth  = authors[i].split(',')[0] if i < len(authors) else (authors[0].split(',')[0] if authors else '?')
                    yr    = years[i] if i < len(years) else (years[0] if years else '?')
                    title = titles[0][:60] if titles else ''
                    db    = db_ids[i] if i < len(db_ids) else (db_ids[0] if db_ids else '')
                    rn_data[rn] = {'author': auth, 'year': yr, 'title': title, 'db_id': db}

        # Also check fldData for hidden refs
        import base64 as _b64
        for b64r in re.findall(r'<w:fldData[^>]*>([\s\S+?]+?)</w:fldData>', rn_xml):
            b64 = b64r.replace('\r','').replace('\n','').replace(' ','')
            pad = (4-len(b64)%4)%4
            try:
                dec = _b64.b64decode(b64+'='*pad).decode('utf-8',errors='replace').replace('\x00','')
                if '<EndNote>' not in dec: continue
                rns    = re.findall(r'<RecNum>(\d+)</RecNum>', dec)
                auths  = re.findall(r'<Author>([^<]+)</Author>', dec)
                yrs    = re.findall(r'<Year>([^<]+)</Year>', dec)
                titles = re.findall(r'<title>([^<]+)</title>', dec)
                db_ids = re.findall(r'db-id="([^"]*)"', dec)
                for i, rn in enumerate(rns):
                    if rn not in rn_data:
                        auth  = auths[i].split(',')[0] if i < len(auths) else (auths[0].split(',')[0] if auths else '?')
                        yr    = yrs[i] if i < len(yrs) else (yrs[0] if yrs else '?')
                        title = titles[0][:60] if titles else ''
                        db    = db_ids[i] if i < len(db_ids) else (db_ids[0] if db_ids else '')
                        rn_data[rn] = {'author': auth, 'year': yr, 'title': title,
                                       'db_id': db, 'hidden': True}
            except: pass

        correct_db = '9veea52shxtee2e2wsbpwvf89wz55atsf52s'
        st.metric("Unique RecNums found", len(rn_data))

        # ── Search / filter ───────────────────────────────────────────────────
        search = st.text_input("Search by author, year, or title", placeholder="e.g. Guirguis or 2017")

        rows = []
        for rn, info in sorted(rn_data.items(), key=lambda x: int(x[0]) if x[0].isdigit() else 9999):
            if search:
                needle = search.lower()
                haystack = f"{info['author']} {info['year']} {info['title']}".lower()
                if needle not in haystack:
                    continue
            db_ok = info['db_id'] == correct_db if info['db_id'] else False
            rows.append({
                'RecNum':  rn,
                'Author':  info['author'],
                'Year':    info['year'],
                'Title':   info['title'][:55],
                'DB-ID ✓': '✓' if db_ok else ('— empty' if not info['db_id'] else '✗ wrong'),
                'Hidden':  '⚠ fldData only' if info.get('hidden') else '',
            })

        if rows:
            import pandas as _pd
            df = _pd.DataFrame(rows)
            st.dataframe(df, use_container_width=True, height=400,
                         column_config={
                             'RecNum':  st.column_config.TextColumn('RecNum', width=80),
                             'Author':  st.column_config.TextColumn('Author', width=120),
                             'Year':    st.column_config.TextColumn('Year', width=60),
                             'Title':   st.column_config.TextColumn('Title'),
                             'DB-ID ✓': st.column_config.TextColumn('DB-ID', width=80),
                             'Hidden':  st.column_config.TextColumn('Status', width=110),
                         })
            # Download as CSV
            csv_data = _pd.DataFrame([{
                'RecNum': r['RecNum'], 'Author': r['Author'],
                'Year': r['Year'], 'Title': r['Title']} for r in rows])
            st.download_button("⬇ Download RecNum list as CSV",
                               data=csv_data.to_csv(index=False),
                               file_name="recnum_list.csv", mime="text/csv")
        else:
            st.info("No matching references found.")

        # ── RecNum remapping ──────────────────────────────────────────────────
        st.divider()
        st.markdown("### Fix wrong RecNums in another document")
        st.markdown(
            "If a document uses placeholder RecNums (e.g. 500, 501…) that don't exist "
            "in your library, enter the remapping below — then upload the document to fix "
            "and download a corrected version."
        )
        st.markdown("**Remapping format** — one per line: `old_recnum -> new_recnum`  "
                    "e.g. `500 -> 79`")

        remap_text = st.text_area("RecNum remapping", height=200,
                                   placeholder="500 -> 79\n501 -> 112\n502 -> 116\n503 -> 413")

        fix_file = st.file_uploader(
            "Document to fix (.docx)", type=["docx"], key="rn_fix_doc")

        if remap_text.strip() and fix_file and st.button("Apply remapping", type="primary"):
            # Parse remap table
            remap = {}
            errors = []
            for line in remap_text.strip().splitlines():
                line = line.strip()
                if not line or line.startswith('#'): continue
                parts = re.split(r'\s*->\s*', line)
                if len(parts) != 2:
                    errors.append(f"Bad line: '{line}'")
                    continue
                old_rn, new_rn = parts[0].strip(), parts[1].strip()
                if not old_rn.isdigit() or not new_rn.isdigit():
                    errors.append(f"Non-numeric RecNum in: '{line}'")
                    continue
                remap[old_rn] = new_rn

            if errors:
                for e in errors: st.error(e)
            elif not remap:
                st.warning("No valid remapping lines found.")
            else:
                fix_bytes = fix_file.read()
                with zipfile.ZipFile(io.BytesIO(fix_bytes)) as z:
                    fix_xml   = z.read('word/document.xml').decode('utf-8')
                    fix_files = {n: z.read(n) for n in z.namelist()}

                applied = []
                for old_rn, new_rn in remap.items():
                    before = fix_xml.count(f'&lt;RecNum&gt;{old_rn}&lt;/RecNum&gt;')
                    if before == 0:
                        st.warning(f"RecNum {old_rn} not found in document — skipped")
                        continue
                    fix_xml = fix_xml.replace(
                        f'&lt;RecNum&gt;{old_rn}&lt;/RecNum&gt;',
                        f'&lt;RecNum&gt;{new_rn}&lt;/RecNum&gt;')
                    fix_xml = fix_xml.replace(
                        f'&lt;rec-number&gt;{old_rn}&lt;/rec-number&gt;',
                        f'&lt;rec-number&gt;{new_rn}&lt;/rec-number&gt;')
                    fix_xml = re.sub(
                        rf'(&lt;key[^&]*&gt;){old_rn}(&lt;/key&gt;)',
                        rf'\g<1>{new_rn}\2', fix_xml)
                    after = fix_xml.count(f'&lt;RecNum&gt;{new_rn}&lt;/RecNum&gt;')
                    applied.append(f"{old_rn} → {new_rn} ({before} citation(s) updated)")

                fix_files['word/document.xml'] = fix_xml.encode('utf-8')
                out_buf = io.BytesIO()
                with zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
                    for n, d in fix_files.items(): zout.writestr(n, d)
                out_buf.seek(0)
                fixed_bytes = out_buf.read()

                new_rns = set(re.findall(r'&lt;RecNum&gt;(\d+)&lt;/RecNum&gt;', fix_xml))
                st.success(f"✓ Applied {len(applied)} remapping(s). Document now has {len(new_rns)} unique RecNums.")
                for msg in applied:
                    st.markdown(f"  • {msg}")
                st.download_button(
                    "⬇ Download remapped document",
                    data=fixed_bytes,
                    file_name=Path(fix_file.name).stem + "_remapped.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )
