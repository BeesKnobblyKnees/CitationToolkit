#!/usr/bin/env python3
"""
Citation Repair Tool — App 1
For pediatric orthopedics manuscript editing workflow.

Usage:
    python citation_repair.py --doc manuscript.docx --library library.xml
    python citation_repair.py --doc manuscript.docx --library library.xml --mode auto
    python citation_repair.py --doc manuscript.docx --library library.xml --pubmed
    python citation_repair.py --doc manuscript.docx --library library.xml --semantic

Modes:
    review  (default) — interactive: shows top matches, you accept/skip
    auto    — inserts best match above threshold automatically, flags rest
    report  — no insertions, just produces a match report as Word doc

Matching:
    Default: TF-IDF (fast, offline, no download required).
    --semantic: sentence-transformers (better accuracy, ~90MB download first run).
"""

import argparse, re, sys, xml.etree.ElementTree as ET
from copy import deepcopy
from datetime import datetime
from pathlib import Path

import requests
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

TOP_N           = 5
TFIDF_THRESHOLD = 0.15   # auto-accept for TF-IDF mode
SEM_THRESHOLD   = 0.82   # auto-accept for semantic mode
PUBMED_MAX      = 5

MISSING_PATTERNS = [
    r'\[CITATION\]', r'\[REF\]', r'\[ref\]', r'\[\?\]',
    r'\[citation needed\]', r'\bXXX\b', r'\[#\]',
    r'<citation>', r'\[ *\]',
]
CITATION_MARKERS = re.compile('|'.join(MISSING_PATTERNS), re.IGNORECASE)

PUBMED_ESEARCH = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
PUBMED_ESUM    = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi"

# ── EndNote XML Parser ──────────────────────────────────────────────────────

def _xml_text(elem, path):
    node = elem.find(path)
    if node is None: return ""
    parts = [node.text or ""] + [c.text or "" for c in node] + [c.tail or "" for c in node]
    return " ".join(p.strip() for p in parts if p.strip())

def parse_endnote_xml(path):
    tree = ET.parse(path); root = tree.getroot()
    refs = []
    for rec in root.iter("record"):
        authors = []
        for a in rec.findall(".//contributors/authors/author"):
            parts = [a.text or ""] + [c.text or "" for c in a]
            name = " ".join(p.strip() for p in parts if p.strip())
            if name: authors.append(name)

        title    = _xml_text(rec, ".//titles/title")
        journal  = _xml_text(rec, ".//periodical/full-title") or _xml_text(rec, ".//periodical/abbr-1")
        year     = _xml_text(rec, ".//dates/year")
        abstract = _xml_text(rec, ".//abstract")
        rec_num  = _xml_text(rec, "rec-number")
        keywords = [" ".join((kw.text or "").split()) for kw in rec.findall(".//keywords/keyword")]

        if not title: continue
        corpus = " ".join(filter(None, [title, abstract, journal,
                                         " ".join(a.split(",")[0] for a in authors),
                                         year, " ".join(keywords)]))
        refs.append(dict(rec_number=rec_num, authors=authors, title=title,
                         journal=journal, year=year, abstract=abstract,
                         keywords=keywords, corpus=corpus))
    print(f"  Loaded {len(refs)} references.")
    return refs

def fmt(ref, short=False):
    aa = ref.get("authors", [])
    if aa:
        a = aa[0].split(",")[0] if len(aa)==1 else \
            f"{aa[0].split(',')[0]} & {aa[1].split(',')[0]}" if len(aa)==2 else \
            f"{aa[0].split(',')[0]} et al."
    else: a = "Unknown"
    y = ref.get("year","n.d.")
    t = ref["title"][:80] + ("…" if len(ref["title"])>80 else "")
    return f"{a} ({y}) — {t}" if short else f"{a} ({y}). {t}. {ref.get('journal','')}"

# ── Flagged Sentence Scanner ────────────────────────────────────────────────

def extract_flagged(doc):
    flagged = []
    for pi, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip(): continue
        for m in CITATION_MARKERS.finditer(text):
            sents = re.split(r'(?<=[.!?])\s+', text)
            cum, target = 0, text
            for s in sents:
                if cum + len(s) >= m.start(): target = s; break
                cum += len(s) + 1
            flagged.append(dict(para_idx=pi, para_text=text, sentence=target,
                                marker=m.group(), marker_start=m.start(), marker_end=m.end()))
    print(f"  Found {len(flagged)} flagged positions.")
    return flagged

# ── Matching Engines ────────────────────────────────────────────────────────

class TfidfMatcher:
    def __init__(self, refs):
        self.refs = refs
        print("  Building TF-IDF index…")
        self.vec = TfidfVectorizer(ngram_range=(1,2), sublinear_tf=True, max_features=50000)
        self.mat = self.vec.fit_transform([r["corpus"] for r in refs])
        print("  Index ready.")

    def match(self, sentence, top_n=TOP_N):
        sv = self.vec.transform([sentence])
        sims = cosine_similarity(sv, self.mat)[0]
        idx  = sims.argsort()[::-1][:top_n]
        return [dict(ref=self.refs[i], score=float(sims[i]), index=int(i)) for i in idx]

class SemanticMatcher:
    def __init__(self, refs):
        try:
            from sentence_transformers import SentenceTransformer
        except ImportError:
            print("ERROR: pip install sentence-transformers"); sys.exit(1)
        self.refs = refs
        print("  Loading semantic model (downloads ~90MB first time)…")
        self.model = SentenceTransformer("all-MiniLM-L6-v2")
        self.embs  = self.model.encode([r["corpus"] for r in refs], show_progress_bar=True)

    def match(self, sentence, top_n=TOP_N):
        sv   = self.model.encode([sentence])
        sims = cosine_similarity(sv, self.embs)[0]
        idx  = sims.argsort()[::-1][:top_n]
        return [dict(ref=self.refs[i], score=float(sims[i]), index=int(i)) for i in idx]

# ── PubMed Fallback ─────────────────────────────────────────────────────────

def pubmed_search(query, max_results=PUBMED_MAX):
    try:
        r  = requests.get(PUBMED_ESEARCH, params=dict(db="pubmed", term=query,
                          retmax=max_results, retmode="json", sort="relevance"), timeout=10)
        ids = r.json().get("esearchresult",{}).get("idlist",[])
        if not ids: return []
        r2 = requests.get(PUBMED_ESUM, params=dict(db="pubmed", id=",".join(ids),
                          retmode="json"), timeout=10)
        data = r2.json().get("result",{})
        return [dict(pmid=pmid,
                     authors=[a.get("name","") for a in data.get(pmid,{}).get("authors",[])],
                     title=data.get(pmid,{}).get("title",""),
                     journal=data.get(pmid,{}).get("fulljournalname",""),
                     year=data.get(pmid,{}).get("pubdate","")[:4],
                     url=f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/") for pmid in ids]
    except Exception as e:
        print(f"  PubMed error: {e}"); return []

# ── Citation Inserter ───────────────────────────────────────────────────────

def insert_superscript(para, marker, label):
    if marker not in para.text: return False
    combined, run_map = "", []
    for run in para.runs:
        s = len(combined); combined += run.text
        run_map.append((s, s+len(run.text), run))
    pos = combined.find(marker)
    if pos == -1: return False
    for (s, e, run) in run_map:
        if s <= pos < e:
            before, after = run.text[:pos-s], run.text[pos-s+len(marker):]
            run.text = before
            rPr = OxmlElement("w:rPr")
            va  = OxmlElement("w:vertAlign"); va.set(qn("w:val"),"superscript"); rPr.append(va)
            nr  = OxmlElement("w:r"); nr.append(deepcopy(rPr))
            t   = OxmlElement("w:t"); t.text = f"[{label}]"
            t.set("{http://www.w3.org/XML/1998/namespace}space","preserve"); nr.append(t)
            run._r.addnext(nr)
            if after:
                tr = OxmlElement("w:r"); tt = OxmlElement("w:t")
                tt.text = after; tt.set("{http://www.w3.org/XML/1998/namespace}space","preserve")
                tr.append(tt); nr.addnext(tr)
            return True
    return False

def author_label(ref):
    aa = ref.get("authors",[])
    last = aa[0].split(",")[0].strip().split()[-1] if aa else "Ref"
    return f"{last} {ref.get('year','')}".strip()

# ── Report Writer ───────────────────────────────────────────────────────────

def write_report(results, out_path):
    from docx.shared import RGBColor
    doc = Document()
    doc.add_heading("Citation Repair Report", 0)
    doc.add_paragraph(f"Generated: {datetime.now():%Y-%m-%d %H:%M}")
    accepted = [r for r in results if r["action"]=="accepted"]
    skipped  = [r for r in results if r["action"]=="skipped"]
    pm_list  = [r for r in results if r["action"]=="pubmed"]
    doc.add_paragraph(f"Total: {len(results)} | Accepted: {len(accepted)} | "
                      f"Skipped: {len(skipped)} | PubMed: {len(pm_list)}")

    if accepted:
        doc.add_heading("Accepted Citations", 1)
        for it in accepted:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"Marker: {it['marker']}\n").bold = True
            p.add_run(f"Context: {it['sentence'][:200]}\n")
            p.add_run(f"Inserted: {fmt(it['ref'])}\n")
            p.add_run(f"Score: {it['score']:.3f}")

    if skipped:
        doc.add_heading("Skipped — Manual Review", 1)
        for it in skipped:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run("NEEDS REVIEW\n"); r.font.color.rgb = RGBColor(0xC0,0,0)
            p.add_run(f"Marker: {it['marker']}\n")
            p.add_run(f"Context: {it['sentence'][:200]}\n")
            for i,c in enumerate(it.get("candidates",[])[:3],1):
                p.add_run(f"  {i}. [{c['score']:.3f}] {fmt(c['ref'],short=True)}\n")

    if pm_list:
        doc.add_heading("PubMed Suggestions", 1)
        for it in pm_list:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"Context: {it['sentence'][:200]}\n")
            for pm in it.get("pubmed_results",[]):
                a = pm["authors"][0] if pm["authors"] else "Unknown"
                p.add_run(f"  • {a} ({pm['year']}). {pm['title']}. {pm['journal']}\n")
                p.add_run(f"    PMID {pm['pmid']} — {pm['url']}\n")

    doc.save(out_path)
    print(f"  Report saved: {out_path}")

# ── Interactive Review Loop ─────────────────────────────────────────────────

def review_loop(flagged, matcher, doc, use_pubmed, mode, threshold):
    results = []
    for i, item in enumerate(flagged):
        sent  = item["sentence"]
        mark  = item["marker"]
        para  = doc.paragraphs[item["para_idx"]]
        cands = matcher.match(sent)
        best  = cands[0]["score"] if cands else 0

        if mode == "auto" and best >= threshold:
            label = author_label(cands[0]["ref"])
            insert_superscript(para, mark, label)
            results.append({**item,"action":"accepted","ref":cands[0]["ref"],
                             "score":best,"candidates":cands})
            print(f"  [{i+1}/{len(flagged)}] AUTO: {label} ({best:.3f})")
            continue

        if mode == "report":
            results.append({**item,"action":"skipped","candidates":cands}); continue

        print(f"\n{'─'*65}")
        print(f"[{i+1}/{len(flagged)}]  Para {item['para_idx']+1}")
        print(f"SENTENCE : {sent[:300]}")
        print(f"MARKER   : {mark}\n")
        for j,c in enumerate(cands,1):
            bar = "█"*int(c["score"]*40)
            print(f"  {j}. [{c['score']:.3f}] {bar}")
            print(f"     {fmt(c['ref'],short=True)}")
        print()
        if use_pubmed and best < 0.10: print("  [p] available — low confidence, try PubMed")
        print("  Accept [1-5] | Skip [s] | PubMed [p] | Quit [q]")

        while True:
            ch = input("  > ").strip().lower()
            if ch == "q":
                print("\nExiting…"); return results
            elif ch == "s":
                results.append({**item,"action":"skipped","candidates":cands}); break
            elif ch == "p" and use_pubmed:
                words = [w for w in re.sub(r'[^\w\s]','',sent).split() if len(w)>4][:8]
                q = " ".join(words)
                print(f"\n  PubMed: '{q}'…")
                pm = pubmed_search(q)
                if pm:
                    for k,r in enumerate(pm,1):
                        a = r["authors"][0] if r["authors"] else "Unknown"
                        print(f"    PM{k}. {a} ({r['year']}). {r['title']}")
                        print(f"         {r['journal']} — {r['url']}")
                    results.append({**item,"action":"pubmed","pubmed_results":pm,"candidates":cands})
                else:
                    print("  No results."); results.append({**item,"action":"skipped","candidates":cands})
                break
            elif ch.isdigit() and 1 <= int(ch) <= len(cands):
                chosen = cands[int(ch)-1]
                label  = author_label(chosen["ref"])
                ok = insert_superscript(para, mark, label)
                if ok:
                    print(f"  ✓ Inserted [{label}]")
                    results.append({**item,"action":"accepted","ref":chosen["ref"],
                                    "score":chosen["score"],"candidates":cands})
                else:
                    print("  ✗ Could not insert (marker spans multiple runs). Flagged.")
                    results.append({**item,"action":"skipped","candidates":cands})
                break
            else:
                print("  Enter 1-5, s, p, or q.")
    return results

# ── Main ────────────────────────────────────────────────────────────────────

def main():
    ap = argparse.ArgumentParser(description="Citation Repair Tool")
    ap.add_argument("--doc",       required=True)
    ap.add_argument("--library",   required=True)
    ap.add_argument("--mode",      default="review", choices=["review","auto","report"])
    ap.add_argument("--pubmed",    action="store_true")
    ap.add_argument("--semantic",  action="store_true", help="Use sentence-transformers (more accurate)")
    ap.add_argument("--threshold", type=float, default=None)
    ap.add_argument("--out",       default=None)
    ap.add_argument("--report",    default=None)
    args = ap.parse_args()

    doc_p = Path(args.doc); lib_p = Path(args.library)
    if not doc_p.exists(): print(f"ERROR: {doc_p} not found"); sys.exit(1)
    if not lib_p.exists(): print(f"ERROR: {lib_p} not found"); sys.exit(1)

    out_p = args.out    or str(doc_p.with_stem(doc_p.stem + "_repaired"))
    rep_p = args.report or str(doc_p.with_stem(doc_p.stem + "_report"))
    thr   = args.threshold or (SEM_THRESHOLD if args.semantic else TFIDF_THRESHOLD)

    print(f"\n{'═'*65}\n  CITATION REPAIR TOOL\n{'═'*65}")
    print(f"  Doc      : {doc_p}\n  Library  : {lib_p}")
    print(f"  Mode     : {args.mode}  |  Matcher: {'Semantic' if args.semantic else 'TF-IDF'}")
    print(f"  Threshold: {thr}  |  PubMed: {'on' if args.pubmed else 'off'}\n{'═'*65}\n")

    doc  = Document(str(doc_p))
    refs = parse_endnote_xml(str(lib_p))
    if not refs: print("ERROR: No refs found in XML."); sys.exit(1)

    flagged = extract_flagged(doc)
    if not flagged:
        print("No markers found. Add placeholders to MISSING_PATTERNS if needed."); sys.exit(0)

    matcher = SemanticMatcher(refs) if args.semantic else TfidfMatcher(refs)
    results = review_loop(flagged, matcher, doc, args.pubmed, args.mode, thr)

    if args.mode != "report":
        doc.save(out_p); print(f"\nRepaired doc: {out_p}")
    write_report(results, rep_p)

    a = sum(1 for r in results if r["action"]=="accepted")
    s = sum(1 for r in results if r["action"]=="skipped")
    p = sum(1 for r in results if r["action"]=="pubmed")
    print(f"\n{'═'*65}\n  DONE  Accepted:{a}  Skipped:{s}  PubMed:{p}\n{'═'*65}\n")

if __name__ == "__main__":
    main()
